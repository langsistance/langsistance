"""Patent document text extraction — PDF (pypdf + OCR), DOCX, XML, and routing.

Shared between celery_worker.py (pipeline Phase 2 download) and
api_routes/core.py (file upload text extraction).
"""

import io
import os as _os
import shutil as _shutil

from sources.logger import Logger

_logger = Logger("text_extractor.log")

# ── OCR image pre-processing ────────────────────────────────────────────────
# Resize images whose longest side exceeds this value before OCR.
# A4 at 300 DPI ≈ 3508 px. Capping at 2400 px gives ~200 DPI equivalent
# — good balance between OCR speed and accuracy, especially for CJK text.
OCR_MAX_DIMENSION = 2400

# ── Preferred download format order ───────────────────────────────────────
# DOCX first (best for text), then PDF (widely available, with OCR fallback),
# then XML last (USPTO xmlarchive is ZIP-wrapped binary, not plain text).
USPTO_PREFERRED_MIME_ORDER = (
    "MS_WORD",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "PDF",
    "application/pdf",
    "XML",
    "application/xml",
    "text/xml",
)


def get_download_url_from_doc(doc: dict) -> str:
    """Extract the best download URL from a USPTO documentBag entry.

    Prefers DOCX/MS_WORD and XML formats over PDF for text extraction.
    The URL may be at the top level or nested in downloadOptionBag[].downloadUrl.
    """
    options = doc.get("downloadOptionBag", [])
    if isinstance(options, list) and options:
        # Collect all download options with their mime types
        url_by_mime: dict[str, str] = {}
        for opt in options:
            if not isinstance(opt, dict):
                continue
            url = ""
            for key in ("downloadUrl", "url"):
                val = opt.get(key, "")
                if val:
                    url = val
                    break
            if url:
                mime = str(opt.get("mimeTypeIdentifier", "")).upper()
                # If no mime, infer from URL extension
                if not mime:
                    url_lower = url.lower()
                    if url_lower.endswith(".docx"):
                        mime = "MS_WORD"
                    elif url_lower.endswith(".xml"):
                        mime = "XML"
                    elif url_lower.endswith(".pdf"):
                        mime = "PDF"
                if mime not in url_by_mime:
                    url_by_mime[mime] = url

        if url_by_mime:
            # Pick the best option based on preference order
            for preferred in USPTO_PREFERRED_MIME_ORDER:
                mime_key = preferred.upper()
                if mime_key in url_by_mime:
                    return url_by_mime[mime_key]
            # Fallback: return any available URL
            return next(iter(url_by_mime.values()))

    # Top-level fallback
    for key in ("downloadUrl", "documentUrl", "url"):
        val = doc.get(key, "")
        if val:
            return val
    return ""


def extract_text_from_pdf(
    content: bytes,
    on_progress: "Callable[[int, int], None] | None" = None,
) -> str | None:
    """Extract text from a PDF binary using pdftotext (poppler-utils).

    1. Scan first 3 pages.
    2. If >2000 chars across first 3 pages → extract all pages.
    3. If ≤2000 chars → return None (likely scanned PDF → vision).

    Args:
        content: PDF binary data.
        on_progress: Ignored (kept for API compatibility).
    """
    import pypdfium2 as pdfium

    try:
        pdf = pdfium.PdfDocument(content)
        total = len(pdf)

        # Scan first 3 pages
        scan_parts: list[str] = []
        for i in range(min(3, total)):
            text = pdf[i].get_text()
            if text:
                scan_parts.append(text)
        scan_text = "\n\n".join(scan_parts).strip()

        if len(scan_text) <= 2000:
            _logger.info(
                f"pdfium_scan_short — pages_scanned={min(3, total)}, "
                f"chars={len(scan_text)}, likely scanned PDF"
            )
            pdf.close()
            return None

        # Full extraction
        _logger.info(
            f"pdfium_scan_ok — chars={len(scan_text)}, extracting all {total} pages"
        )
        parts: list[str] = []
        for i in range(total):
            text = pdf[i].get_text()
            if text:
                parts.append(text)
        pdf.close()
        extracted = "\n\n".join(parts).strip()
        _logger.info(f"pdfium_extracted — pages={total}, chars={len(extracted)}")
        return extracted if extracted else None
    except Exception as e:
        _logger.warning(f"pdfium_failed — {e}")
        try:
            pdf.close()
        except Exception:
            pass
        return None
    finally:
        try:
            _os.unlink(tmp_path)
        except OSError:
            pass


def _ocr_from_pdf_bytes(
    content: bytes,
    on_progress: "Callable[[int, int], None] | None" = None,
) -> str | None:
    """OCR fallback for image-based PDFs — pdf2image rendering + tesseract.

    Args:
        content: PDF binary data.
        on_progress: Optional callback(current_page, total_pages).
    """
    tesseract_bin = _find_tesseract_bin()
    if not tesseract_bin:
        _logger.warning("ocr_skipped — tesseract binary not found")
        return None

    try:
        import pytesseract
    except ImportError:
        _logger.warning("ocr_skipped — pytesseract not installed")
        return None

    try:
        from PIL import Image
    except ImportError:
        _logger.warning("ocr_skipped — Pillow not installed")
        return None

    pytesseract.pytesseract.tesseract_cmd = tesseract_bin

    from pdf2image import convert_from_bytes

    try:
        pil_images = convert_from_bytes(content, dpi=200)
    except Exception as e:
        _logger.warning(f"pdf2image_render_failed — {e}")
        return None

    page_count = len(pil_images)
    all_text: list[str] = []
    ocr_failures = 0

    for i, pil_img in enumerate(pil_images):
        try:
            if pil_img.mode not in ("RGB", "L", "1"):
                pil_img = pil_img.convert("RGB")
            text = pytesseract.image_to_string(pil_img, lang="eng")
            if text and text.strip():
                all_text.append(text.strip())
            else:
                ocr_failures += 1
        except Exception as e:
            _logger.warning(f"ocr_page_{i+1}_failed — {e}")
            ocr_failures += 1

        if on_progress:
            try:
                on_progress(i + 1, page_count)
            except Exception:
                pass

    extracted = "\n\n".join(all_text).strip()
    successful = page_count - ocr_failures

    if extracted and len(extracted) > 100:
        _logger.info(
            f"ocr_text_extracted — pages={page_count}, successful={successful}, "
            f"failed={ocr_failures}, chars={len(extracted)}"
        )
        return extracted

    _logger.warning(
        f"ocr_text_empty_or_short — pages={page_count}, successful={successful}, "
        f"failed={ocr_failures}, chars={len(extracted)}"
    )
    return None


# Backward-compatible alias used by celery_worker
_ocr_from_pdf_reader = _ocr_from_pdf_bytes


def _find_tesseract_bin() -> str | None:
    """Find the tesseract binary on the system.

    Checks PATH first, then common installation directories.
    """
    # 1. Check PATH
    path = _shutil.which("tesseract")
    if path:
        return path

    # 2. Check common installation directories
    candidates = [
        # Windows
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
        # Linux / macOS
        "/usr/bin/tesseract",
        "/usr/local/bin/tesseract",
        "/opt/homebrew/bin/tesseract",
    ]
    for candidate in candidates:
        if _os.path.isfile(candidate) and _os.access(candidate, _os.X_OK):
            return candidate

    return None


def extract_text_from_docx(content: bytes) -> str | None:
    """Extract text from a DOCX binary using python-docx."""
    from docx import Document

    try:
        doc = Document(io.BytesIO(content))
        parts: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        extracted = "\n".join(parts).strip()
        if extracted and len(extracted) > 100:
            _logger.info(f"docx_text_extracted — chars={len(extracted)}")
            return extracted
        _logger.warning(f"docx_text_empty_or_short — chars={len(extracted)}")
        return None
    except Exception as e:
        _logger.warning(f"docx_extract_failed — {e}")
        return None


def extract_text_from_binary(
    content: bytes,
    content_type: str = "",
    filename: str = "",
    on_progress: "Callable[[int, int], None] | None" = None,
) -> str | None:
    """Route binary content to the appropriate text extractor.

    Args:
        content: Raw file bytes.
        content_type: MIME type (e.g. 'application/pdf'), optional.
        filename: Original filename used for extension detection, optional.
        on_progress: Optional callback(current_page, total_pages) for OCR.

    Returns:
        Extracted text, or None if extraction fails.
    """
    ct = content_type.lower()
    fn_lower = filename.lower()

    # Check for DOCX
    is_docx = (
        "vnd.openxmlformats-officedocument.wordprocessingml" in ct
        or "msword" in ct
        or fn_lower.endswith(".docx")
        or content[:4] == b"PK\x03\x04"  # DOCX is a ZIP; check magic bytes
    )
    if is_docx:
        return extract_text_from_docx(content)

    # Check for XML (plain-text XML only, not USPTO xmlarchive ZIP files)
    is_xml = (
        ("xml" in ct and "vnd.openxmlformats" not in ct)
        or fn_lower.endswith(".xml")
        or content[:100].lstrip().startswith(b"<?xml")
    )
    if is_xml:
        try:
            text = content.decode("utf-8", errors="replace")
            if len(text.strip()) > 100:
                _logger.info(f"xml_text — chars={len(text)}")
                return text
        except Exception as e:
            _logger.warning(f"xml_decode_failed — {e}")

    # Default: treat as PDF
    is_likely_pdf = (
        "pdf" in ct
        or "octet-stream" in ct
        or fn_lower.endswith(".pdf")
        or content[:5] == b"%PDF-"
    )
    if is_likely_pdf:
        return extract_text_from_pdf(content, on_progress=on_progress)

    # Last resort: try UTF-8 decode
    try:
        text = content.decode("utf-8", errors="replace")
        if text.strip():
            return text
    except Exception:
        pass
    return None
