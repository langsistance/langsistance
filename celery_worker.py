"""Celery worker entry point for long-running patent analysis tasks."""

import asyncio
import json
import os
import sys
import warnings

# Suppress Celery's superuser warning in Docker (container runs as root)
warnings.filterwarnings('ignore', message='.*running the worker with superuser privileges.*')

# Ensure the project root is on sys.path so `sources.*` imports work
# both under docker-compose (volume mount at /app) and bare Docker runs.
_project_root = os.path.dirname(os.path.abspath(__file__))
if _project_root not in sys.path:
    sys.path.insert(0, _project_root)

from celery import Celery
from sources.logger import Logger
from sources.analytics import track_event

_pipeline_logger = Logger("long_task_pipeline.log")

REDIS_HOST = os.getenv('REDIS_HOST', 'localhost')
REDIS_PORT = int(os.getenv('REDIS_PORT', '6379'))
REDIS_URL = f"redis://{REDIS_HOST}:{REDIS_PORT}/0"

app = Celery('patent_tasks', broker=REDIS_URL, backend=REDIS_URL)

app.conf.update(
    task_serializer='json',
    result_serializer='json',
    accept_content=['json'],
    task_track_started=True,
    task_acks_late=True,
    task_reject_on_worker_lost=True,
    worker_prefetch_multiplier=1,
    broker_connection_retry_on_startup=True,
    broker_transport_options={
        # After a worker dies HARD (kill -9 / power loss), unacked tasks
        # become visible to other workers after this many seconds.
        #
        # MUST be >= time_limit (3600 s) to prevent dual execution.
        # With task_acks_late=True, the task is NOT acknowledged until it
        # completes.  If visibility_timeout < actual runtime, Redis will
        # redeliver the task to another worker while the first worker is
        # still running — causing TWO workers to execute the same task.
        #
        # Graceful shutdown (docker stop / SIGTERM) is handled by
        # worker_shutdown_timeout + stop_grace_period — the worker
        # rejects unacked tasks within 25 s, so they return to the
        # queue immediately.  visibility_timeout is only for hard crashes.
        'visibility_timeout': 7200,
    },
    # Allow running tasks enough time to reject cleanly during warm shutdown
    # before Docker sends SIGKILL.
    worker_shutdown_timeout=25,
)


@app.task(bind=True, max_retries=3, default_retry_delay=30, time_limit=3600, soft_time_limit=3540)
def execute_patent_analysis(self, task_id: str, params: dict):
    """Batch patent analysis -- 4-phase serial pipeline with checkpointing."""
    retry_count = self.request.retries
    _pipeline_logger.info(
        f"[task={task_id}] START — "
        f"query={params.get('query', '')[:120]}, "
        f"patent_source={params.get('patent_source', 'cnipa')}, "
        f"session_id={params.get('session_id', '')}, "
        f"scene_id={params.get('scene_id', '')}, "
        f"retry={retry_count}/{self.max_retries}"
    )
    # Belt-and-suspenders: if Celery's retry tracking is broken, hard-stop here
    if retry_count >= self.max_retries:
        _pipeline_logger.error(
            f"[task={task_id}] HARD_STOP — retry_count={retry_count} >= {self.max_retries}"
        )
        user_id = params.get('user_id', '')
        if user_id:
            from sources.long_task.user_queue import complete_user_task
            try:
                complete_user_task(str(user_id), task_id)
            except Exception:
                pass
        return {'status': 'failed', 'task_id': task_id,
                'error': f'Max retries ({self.max_retries}) exceeded'}
    from sources.long_task.status_manager import (
        update_task_status, set_task_completed, set_task_failed,
        save_checkpoint, load_checkpoint,
    )
    from sources.long_task.config import (
        get_long_task_config,
        DEFAULT_VISION_PROVIDER,
        DEFAULT_VISION_MODEL,
    )
    from sources.long_task.patent_analyzer import (
        generate_table_columns, download_patent_document,
        analyze_single_patent, generate_patent_summary, build_failed_row,
    )
    from sources.long_task.report_generator import (
        generate_report_outline, generate_report_section,
        generate_executive_summary,
    )
    from sources.long_task.storage import create_storage
    from sources.llm_provider import Provider

    ltc = get_long_task_config()
    model_family = ltc['provider_family']
    max_patents = ltc['max_patents']
    max_patents_cnipa = ltc.get('max_patents_cnipa', 10)
    max_patents_uspto = ltc.get('max_patents_uspto', 50)
    vision_enabled = ltc.get('vision_enabled', True)

    # ---- Input dedup + source-based truncation ----
    patent_ids = sorted(set(params.get('patent_ids', [])))
    # Auto-detect USPTO: pure 8-digit numbers are USPTO application IDs
    if patent_ids and params.get('patent_source', '') != 'uspto':
        all_uspto = all(
            pid.isdigit() and len(pid) == 8 for pid in patent_ids
        )
        if all_uspto:
            params['patent_source'] = 'uspto'
            _pipeline_logger.info(
                f"[task={task_id}] CONFIG auto_detect_uspto — "
                f"all {len(patent_ids)} IDs are 8-digit, "
                f"overriding patent_source to 'uspto'"
            )
    source_max = _get_max_patents_for_source(
        params.get('patent_source', ''), max_patents,
        max_patents_cnipa, max_patents_uspto,
    )
    patent_ids = patent_ids[:source_max]
    total = len(patent_ids)

    _pipeline_logger.info(
        f"[task={task_id}] CONFIG — "
        f"model_family={model_family}, max_patents={source_max}, "
        f"patent_ids_count={len(patent_ids)}, "
        f"patent_ids={patent_ids[:10]}{'...' if len(patent_ids) > 10 else ''}"
    )

    batch_lang = params.get('lang', 'zh')

    # ── Immediate progress update so frontend shows feedback right away ──
    _pipeline_logger.info(f'[task={task_id}] PRE_STATUS_UPDATE — writing preparing status to Redis')
    try:
        update_task_status(task_id, 'preparing', 1,
                           _t('preparing', batch_lang, total=total))
        _pipeline_logger.info(f'[task={task_id}] POST_STATUS_UPDATE — Redis write ok')
    except Exception as e:
        _pipeline_logger.error(f'[task={task_id}] STATUS_UPDATE_FAILED — Redis write error: {e}')
        _pipeline_logger.info(f'[task={task_id}] CONTINUING_WITHOUT_REDIS — pipeline proceeding despite status update failure')

    # ---- Provider setup ----
    if model_family == 'minimax':
        flash_provider = Provider(provider_name='minimax', model='MiniMax-M2.7-highspeed',
                                  server_address='', is_local=False)
        # Single patent: use M2.7 for text analysis
        # Batch: use M3
        if total == 1:
            pro_provider = Provider(provider_name='minimax', model='MiniMax-M2.7-highspeed',
                                    server_address='', is_local=False)
        else:
            pro_provider = Provider(provider_name='minimax', model='MiniMax-M3',
                                    server_address='', is_local=False)
    else:
        flash_provider = Provider(provider_name='deepseek', model='deepseek-v4-flash',
                                  server_address='', is_local=False)
        pro_provider = Provider(provider_name='deepseek', model='deepseek-v4-pro',
                                server_address='', is_local=False)

    # Vision provider — configured via config.ini [LONG_TASK]
    #   vision_provider / vision_model — independent of provider_family
    vision_provider = None
    if vision_enabled:
        vision_cfg_provider = ltc.get('vision_provider', DEFAULT_VISION_PROVIDER)
        vision_cfg_model = ltc.get('vision_model', DEFAULT_VISION_MODEL)
        vision_provider = Provider(
            provider_name=vision_cfg_provider,
            model=vision_cfg_model,
            server_address='',
            is_local=False,
        )

    try:
        # ---- Run the 4-phase pipeline using event-loop-safe pattern ----
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            result = loop.run_until_complete(_run_pipeline(
                task_id=task_id,
                params=params,
                patent_ids=patent_ids,
                total=total,
                max_patents=max_patents,
                max_patents_cnipa=max_patents_cnipa,
                max_patents_uspto=max_patents_uspto,
                batch_lang=batch_lang,
                flash_provider=flash_provider,
                pro_provider=pro_provider,
                vision_provider=vision_provider if vision_enabled else None,
                update_task_status=update_task_status,
                set_task_completed=set_task_completed,
                set_task_failed=set_task_failed,
                save_checkpoint=save_checkpoint,
                load_checkpoint=load_checkpoint,
                generate_table_columns=generate_table_columns,
                download_patent_document=download_patent_document,
                analyze_single_patent=analyze_single_patent,
                generate_patent_summary=generate_patent_summary,
                build_failed_row=build_failed_row,
                generate_report_outline=generate_report_outline,
                generate_report_section=generate_report_section,
                generate_executive_summary=generate_executive_summary,
                create_storage=create_storage,
            ))
            _pipeline_logger.info(
                f"[task={task_id}] PIPELINE_DONE — status={result.get('status')}"
            )
            return result
        finally:
            loop.close()
    except Exception as e:
        _pipeline_logger.error(
            f"[task={task_id}] FAILED — error={e}"
        )
        set_task_failed(task_id, str(e))
        user_id_for_analytics = params.get('user_id', '')
        if user_id_for_analytics:
            track_event("long_task:fail", user_id=user_id_for_analytics,
                        task_id=task_id,
                        extra={"error": str(e)[:100]})
        try:
            raise self.retry(exc=e)
        except self.MaxRetriesExceededError:
            # Permanent failure — clear user's running key so queued tasks proceed.
            # Delete the running key directly first (safety net), then try
            # complete_user_task to dispatch the next queued task.
            _pipeline_logger.error(
                f"[task={task_id}] MAX_RETRIES_EXCEEDED — clearing user queue lock"
            )
            user_id = params.get('user_id', '')
            if user_id:
                try:
                    from sources.knowledge.knowledge import get_redis_connection
                    r = get_redis_connection()
                    running_key = f"lt:user:{user_id}:running"
                    r.delete(running_key)
                except Exception:
                    pass
                try:
                    from sources.long_task.user_queue import complete_user_task
                    next_id = complete_user_task(str(user_id), task_id)
                    if next_id:
                        _pipeline_logger.info(
                            f"[task={task_id}] QUEUE_DISPATCHED_AFTER_FAILURE — "
                            f"next_task_id={next_id}"
                        )
                except Exception as qe:
                    _pipeline_logger.warning(
                        f"[task={task_id}] QUEUE_CLEANUP_FAILED — {qe}"
                    )
            raise


def _update_mysql_progress(task_id: str, current_phase: str, progress: int, result_summary: str = None) -> None:
    """Update the long_tasks MySQL row with current progress and optional report text."""
    try:
        from sources.knowledge.knowledge import get_db_connection
        conn = get_db_connection()
        try:
            with conn.cursor() as cur:
                if result_summary is not None:
                    cur.execute(
                        """UPDATE long_tasks
                           SET status = 'running', current_phase = %s, progress = %s,
                               result_summary = %s
                           WHERE task_id = %s""",
                        (current_phase, progress, result_summary, task_id))
                else:
                    cur.execute(
                        """UPDATE long_tasks
                           SET status = 'running', current_phase = %s, progress = %s
                           WHERE task_id = %s""",
                        (current_phase, progress, task_id))
                conn.commit()
        finally:
            conn.close()
    except Exception:
        pass  # Non-fatal: MySQL update failure should not break the pipeline


# ── Status message translations ──
_STATUS_MSGS = {
    "zh": {
        "preparing": "正在准备专利分析（{total} 个专利）...",
        "extracting_file": "正在解析上传文件（{current}/{total}）...",
        "extracting_by_name": "正在解析：{name}（{current}/{total}）...",
        "ocr_page": "OCR识别：{name}（{current}/{total}页）...",
        "tool_select": "已发现 {count} 个场景工具，正在选择检索方案...",
        "tool_search": "正在检索专利：{reason}",
        "searching": "正在搜索 {focus} 的相关专利（USPTO）...",
        "searching_page": "正在搜索 {focus} 的相关专利（第{page}页）...",
        "searching_resolve": "正在获取第{page}页专利详情（{resolved}/{on_page}）...",
        "analyzing": "正在分析第 {current}/{total} 个专利...",
        "analyzing_progress": "分析进度: {current}/{total}",
        "generating_framework": "正在生成分析框架（{total} 个专利）...",
        "downloading_patent": "正在下载专利文件（{current}/{total}）...",
        "analysis_completed_count": "已完成 {current}/{total} 个专利分析",
        "writing_summary": "正在撰写核心审查洞察...",
        "generating_outline": "正在规划报告结构...",
        "writing_section": "正在撰写：{heading}",
        "section_failed": "（{heading} 生成失败）",
        "report_writing_complete": "报告撰写完成",
        "default_report_title": "专利分析报告",
        "default_section_heading": "分析结果",
        "default_exec_summary_heading": "核心审查洞察",
        "default_analysis_table": "审查文件分析数据表",
        "search_complete": "检索到 {count} 个专利，开始分析",
        "no_patents_found": "未找到匹配的专利",
        "generating_word": "正在生成 Word 文件...",
        "generating_pdf": "正在从 Word 生成 PDF 文件...",
        "fetching_uspto": "正在获取USPTO文件列表...",
        "prosecution_no_docs": "USPTO未返回专利申请 {app_number} 的任何文件。可能原因：申请号不存在、无权访问、或尚未公开。",
        "prosecution_no_analyzable": "未找到可分析的审查文件（无 Office Action、Response 或 Amendment）。可能该专利尚未进入实质审查阶段。",
        "prosecution_framework": "正在生成分析框架（{total} 个审查文件）...",
        "prosecution_downloading": "正在下载审查文件（{current}/{total}）...",
        "prosecution_ocr": "正在OCR识别（{current}/{total}）...",
        "prosecution_analyzing": "正在分析（{current}/{total}）：{desc}",
        "prosecution_all_failed": "所有审查文件处理失败。文件可能是扫描件或加密PDF。",
        "prosecution_report_complete": "报告撰写完成",
        "family_lookup": "正在查询EPO同族专利...",
        "china_resolving": "正在解析中国专利申请号...",
        "china_fetching_review": "正在获取中国审查决定数据...",
        "china_analyzing_decisions": "正在分析审查决定（{current}/{total}）...",
        "china_generating_timeline": "正在生成审查时间线...",
        "china_no_cn_member": "未找到该专利的中国同族成员，无法进行中国审查历史分析。",
        "china_no_review_data": "未找到该中国专利的审查决定数据。",
        "china_report_title": "中国专利 {patent_id} 审查历史分析报告",
        "china_family_overview": "查到 {input_id} 的中国同族申请号 {cn_app}，同族覆盖 {n_juris} 个司法辖区: {juris}",
        "china_fetching_for_family": "正在获取中国同族专利 {cn_app} 的审查数据...",
        "family_cn_done": "已获取 {events} 条审查决定 / {claims} 项权利要求",
        "family_cn_basic": "仅有著录项目数据",
        "japan_fetching_for_family": "正在获取日本同族专利 {jp_app} 的审查数据...",
        "family_jp_done": "已获取 {events} 条审查进度事件",
        "family_jp_basic": "仅有著录项目数据",
        "family_ep_done": "已获取 {steps} 个审查步骤 ({status})",
        "family_ep_basic": "仅有著录项目数据",
        "epo_fetching_for_family": "正在获取欧洲同族专利 EP{ep_app} 的审查数据...",
        "epo_no_ep_member": "未找到该专利的欧洲同族成员，无法进行EPO审查历史分析。",
        "epo_no_review_data": "未找到该欧洲专利的审查数据。",
        "epo_report_title": "欧洲专利 {patent_id} 审查历史分析报告",
        "japan_no_review_data": "未找到该日本专利的审查数据。",
        "japan_report_title": "日本专利 {patent_id} 审查历史分析报告",
    },
    "en": {
        "preparing": "Preparing patent analysis ({total} patents)...",
        "extracting_file": "Parsing uploaded files ({current}/{total})...",
        "extracting_by_name": "Parsing: {name} ({current}/{total})...",
        "ocr_page": "OCR: {name} (page {current}/{total})...",
        "tool_select": "Found {count} scene tools, selecting search strategy...",
        "tool_search": "Searching patents: {reason}",
        "searching": "Searching {focus} patents (USPTO)...",
        "searching_page": "Searching {focus} patents (page {page})...",
        "searching_resolve": "Fetching patent details page {page} ({resolved}/{on_page})...",
        "analyzing": "Analyzing patent {current}/{total}...",
        "analyzing_progress": "Analysis progress: {current}/{total}",
        "generating_framework": "Building analysis framework ({total} patents)...",
        "downloading_patent": "Downloading patent files ({current}/{total})...",
        "analysis_completed_count": "Completed {current}/{total} patent analyses",
        "writing_summary": "Writing Key Prosecution Insights...",
        "generating_outline": "Planning report structure...",
        "writing_section": "Writing: {heading}",
        "section_failed": "({heading} generation failed)",
        "report_writing_complete": "Report writing complete",
        "default_report_title": "Patent Analysis Report",
        "default_section_heading": "Analysis Results",
        "default_exec_summary_heading": "Key Prosecution Insights",
        "default_analysis_table": "Prosecution Document Analysis Table",
        "search_complete": "Found {count} patents, starting analysis",
        "no_patents_found": "No matching patents found",
        "generating_word": "Generating Word document...",
        "generating_pdf": "Converting DOCX to PDF...",
        "fetching_uspto": "Fetching USPTO file list...",
        "prosecution_no_docs": "USPTO returned no documents for application {app_number}. The application may not exist, may not be accessible, or may not yet be published.",
        "prosecution_no_analyzable": "No analyzable prosecution documents found (no Office Actions, Responses, or Amendments). The patent may not have entered substantive examination.",
        "prosecution_framework": "Building analysis framework ({total} prosecution documents)...",
        "prosecution_downloading": "Downloading documents ({current}/{total})...",
        "prosecution_ocr": "OCR processing {current}/{total}...",
        "prosecution_analyzing": "Analyzing {current}/{total}: {desc}",
        "prosecution_all_failed": "All prosecution documents failed processing. Files may be scanned images or encrypted PDFs.",
        "prosecution_report_complete": "Report writing complete",
        "family_lookup": "Looking up patent family via EPO...",
        "china_resolving": "Resolving Chinese patent application number...",
        "china_fetching_review": "Fetching Chinese examination review data...",
        "china_analyzing_decisions": "Analyzing examination decisions ({current}/{total})...",
        "china_generating_timeline": "Building examination timeline...",
        "china_no_cn_member": "No Chinese family member found for this patent.",
        "china_no_review_data": "No examination review data found for this Chinese patent.",
        "china_report_title": "Chinese Patent {patent_id} Examination History Report",
        "china_family_overview": "Found CN application {cn_app} for {input_id}, family spans {n_juris} jurisdictions: {juris}",
        "china_fetching_for_family": "Fetching CN family member {cn_app} examination data...",
        "family_cn_done": "Fetched {events} decisions / {claims} claims",
        "family_cn_basic": "Bibliographic data only",
        "japan_fetching_for_family": "Fetching JP family member {jp_app} examination data...",
        "family_jp_done": "Fetched {events} progress events",
        "family_jp_basic": "Bibliographic data only",
        "family_ep_done": "Fetched {steps} procedural steps ({status})",
        "family_ep_basic": "Bibliographic data only",
        "epo_fetching_for_family": "Fetching EP family member EP{ep_app} examination data...",
        "epo_no_ep_member": "No European family member found for this patent.",
        "epo_no_review_data": "No examination data found for this European patent.",
        "epo_report_title": "European Patent {patent_id} Examination History Report",
        "japan_no_review_data": "No examination data found for this Japanese patent.",
        "japan_report_title": "Japanese Patent {patent_id} Examination History Report",
    },
}

def _t(key, lang="zh", **kwargs):
    msgs = _STATUS_MSGS.get(lang, _STATUS_MSGS["zh"])
    msg = msgs.get(key, key)
    return msg.format(**kwargs) if kwargs else msg


async def _run_pipeline(
    task_id: str,
    params: dict,
    patent_ids: list,
    total: int,
    max_patents: int,
    flash_provider,
    pro_provider,
    vision_provider,
    update_task_status,
    set_task_completed,
    set_task_failed,
    save_checkpoint,
    load_checkpoint,
    generate_table_columns,
    download_patent_document,
    analyze_single_patent,
    generate_patent_summary,
    build_failed_row,
    generate_report_outline,
    generate_report_section,
    generate_executive_summary,
    create_storage,
    max_patents_cnipa: int = 10,
    max_patents_uspto: int = 50,
    batch_lang: str = 'zh',
) -> dict:
    """Internal async pipeline orchestrator."""

    # ---- Crash recovery / resume: load checkpoint ----
    checkpoint = load_checkpoint(task_id)
    _pipeline_logger.info(
        f"[task={task_id}] CHECKPOINT_LOAD — "
        f"found={checkpoint is not None}, "
        f"has_pending={bool(checkpoint and checkpoint.get('pending'))}, "
        f"completed_rows={len(checkpoint.get('completed_rows', [])) if checkpoint else 0}, "
        f"pending_count={len(checkpoint.get('pending', [])) if checkpoint and checkpoint.get('pending') else 0}"
    )
    resume_columns = None
    if checkpoint and checkpoint.get('pending'):
        table_rows = checkpoint.get('completed_rows', [])
        pending = checkpoint['pending']
        # Restore columns from checkpoint (only when resuming a paused task)
        if checkpoint.get('columns'):
            resume_columns = checkpoint['columns']
        elif table_rows:
            # Fallback: extract columns from first completed row
            first = table_rows[0]
            resume_columns = [k for k in first.keys() if not k.startswith('_')]
    elif checkpoint and checkpoint.get('completed_rows'):
        # All patents already analyzed — skip PHASE0-2, jump directly to PHASE3
        table_rows = checkpoint['completed_rows']
        pending = []
        resume_columns = checkpoint.get('columns')
        if not resume_columns and table_rows:
            first = table_rows[0]
            resume_columns = [k for k in first.keys() if not k.startswith('_')]
        _pipeline_logger.info(
            f"[task={task_id}] CHECKPOINT_ALL_DONE — "
            f"all {len(table_rows)} patents already analyzed, "
            f"skipping PHASE0-2, jumping to PHASE3"
        )
    else:
        table_rows = []
        pending = patent_ids

    # ---- Mode detection ----
    patent_texts = params.get('patent_texts', {}) or {}
    patent_file_refs = params.get('patent_file_refs', []) or []

    # Mode 2 (file upload): extract text from saved files asynchronously
    if patent_file_refs and not patent_texts:
        total_files = len(patent_file_refs)
        _pipeline_logger.info(
            f"[task={task_id}] FILE_EXTRACT — extracting text from "
            f"{total_files} uploaded files"
        )
        update_task_status(task_id, 'extracting_text', 0,
                           _t('extracting_file', batch_lang, current=0, total=total_files))

        from sources.long_task.text_extractor import extract_text_from_binary

        for idx, ref in enumerate(patent_file_refs):
            pid = ref['filename'].rsplit('.', 1)[0]
            update_task_status(
                task_id, 'extracting_text',
                5 + int((idx / total_files) * 15),
                _t('extracting_by_name', batch_lang, name=ref['filename'], current=idx+1, total=total_files),
            )

            # Per-page OCR progress callback (updates Redis)
            ocr_total_pages = [0]

            def _ocr_callback(current: int, total: int):
                ocr_total_pages[0] = total
                file_pct = 5 + int(((idx + current / total) / total_files) * 15)
                update_task_status(
                    task_id, 'extracting_text', file_pct,
                    _t('ocr_page', batch_lang, name=ref['filename'], current=current, total=total),
                )

            try:
                with open(ref['path'], 'rb') as fh:
                    content = fh.read()
                text = extract_text_from_binary(
                    content, ref.get('content_type', ''), ref['filename'],
                    on_progress=_ocr_callback,
                )
                if text and len(text) > 100:
                    patent_texts[pid] = text
                    _pipeline_logger.info(
                        f"[task={task_id}] FILE_EXTRACT — {ref['filename']}: "
                        f"{len(text)} chars"
                    )
                else:
                    _pipeline_logger.warning(
                        f"[task={task_id}] FILE_EXTRACT — {ref['filename']}: "
                        f"extraction failed ({len(text) if text else 0} chars)"
                    )
            except Exception as e:
                _pipeline_logger.error(
                    f"[task={task_id}] FILE_EXTRACT — {ref['filename']}: {e}"
                )

    is_file_upload_mode = bool(patent_texts)
    is_direct_id_mode = bool(patent_ids) and not is_file_upload_mode
    if is_file_upload_mode:
        _pipeline_logger.info(
            f"[task={task_id}] MODE=file_upload — patent_count={len(patent_ids)}, "
            f"patent_ids={patent_ids[:10]}{'...' if len(patent_ids) > 10 else ''}"
        )
    elif is_direct_id_mode:
        _pipeline_logger.info(
            f"[task={task_id}] MODE=direct_ids — patent_count={len(patent_ids)}, "
            f"patent_ids={patent_ids[:10]}{'...' if len(patent_ids) > 10 else ''}"
        )
    else:
        _pipeline_logger.info(
            f"[task={task_id}] MODE=search_extract — "
            f"patent_ids_from_params={len(patent_ids)}"
        )

    scene_id = params.get('scene_id')
    # Load scene knowledge+tools once for use across all phases
    scene_candidates = None
    if scene_id:
        from sources.long_task.scene_tools import get_scene_knowledge_tools
        try:
            scene_candidates = get_scene_knowledge_tools(scene_id)
        except Exception as e:
            _pipeline_logger.warning(
                f'[task={task_id}] PHASE0 scene_tools_failed — '
                f'scene_id={scene_id}, error={e}'
            )
            scene_candidates = None
        _pipeline_logger.info(
            f"[task={task_id}] PHASE0 — scene_id={scene_id}, "
            f"scene_candidates_count={len(scene_candidates) if scene_candidates else 0}"
        )
    id_url_map = {}          # patent_id → document_url from search results
    id_pid_map = {}          # patent_id → pid from CNIPA search results

    # ==== Phase 0-prep: Extract patent IDs from query + conversation via LLM ====
    if not is_file_upload_mode and not is_direct_id_mode and not patent_ids:
        query_text = params.get('query', '')
        conv_history = params.get('conversation_history', [])
        scenario = params.get('scenario', '')
        # Only search conversation history for IDs when the user is explicitly
        # referencing previous results (conversation_refs). For standalone queries
        # (direct_ids or search), limit to the query text only.
        if scenario == 'conversation_refs':
            combined_parts = [query_text] if query_text else []
            for msg in (conv_history or []):
                content = msg.get('content', '') if isinstance(msg, dict) else ''
                if content and content not in combined_parts:
                    combined_parts.append(content)
            combined_text = "\n".join(combined_parts)
        else:
            combined_text = query_text

        if combined_text and len(combined_text) > 20:
            _pipeline_logger.info(
                f"[task={task_id}] PHASE0 extract_patent_ids_from_text — "
                f"text_length={len(combined_text)}, "
                f"sources=query({len(query_text)})+conv_msgs({len(conv_history)})"
            )
            EXTRACT_ALL_PROMPT = (
                "You are a patent ID extractor. "
                "Extract ALL patent APPLICATION NUMBERS from the provided text. "
                "\n\n"
                "CRITICAL: The user may have typed IDs directly (e.g. "
                "'17429113、18012525、18331482。这是3个专利申请号'). "
                "When the user mentions '专利申请号', 'application number', "
                "or 'patent application', the adjacent numbers ARE the IDs. "
                "\n\n"
                "—— HOW TO IDENTIFY THE SOURCE COUNTRY ——\n"
                "CNIPA (China / 中国国家知识产权局) application numbers:\n"
                "  - Format: YYYY + 8 digits + optional '.' + 1 check digit\n"
                "  - Examples: 202310123456.7, 2023101234567, 202410567890.1\n"
                "  - ALWAYS start with the year (2018-2026), 13+ digits total\n"
                "  - May have a dot before the last digit (check digit)\n\n"
                "USPTO (United States / 美国专利商标局) application numbers:\n"
                "  - Format: PURE 8-DIGIT NUMBERS or 2/6 slash format\n"
                "  - Examples: 18331482, 17429113, 18/333482\n"
                "  - Exactly 8 consecutive digits, or 2 digits + '/' + 6 digits\n"
                "  - Do NOT confuse with publication numbers (US20230310100A1)\n"
                "  - Do NOT confuse with granted patent numbers (e.g. 10299867)\n\n"
                "DECISION RULE: count the digits.\n"
                "  - Starts with year (20XX) + 9+ more digits → CNIPA\n"
                "  - Pure 8-digit number → USPTO\n"
                "  - If unsure, return 'unknown'\n\n"
                'Return JSON: {"patent_ids": ["id1", ...], '
                '"source": "uspto" or "cnipa" or "unknown"}'
            )
            EXTRACT_ONE_PROMPT = (
                "Extract ALL patent APPLICATION NUMBERS from this text. "
                "Application numbers can be CNIPA (starts with year like 2023, "
                "13+ digits, e.g. 202310123456.7) or USPTO (pure 8-digit, "
                "e.g. 18331482). "
                "Look for numbers labeled '申请号' or 'Application Number', "
                "or bare numbers that the user refers to as patent/application IDs. "
                "DO NOT extract publication numbers (US-prefixed like US20230310100A1) "
                "or granted patent numbers. "
                "If none found, return empty list. "
                'Return JSON: {"patent_ids": ["id1", ...], '
                '"source": "uspto" or "cnipa" or "unknown"}'
            )

            import re as _re2
            normalized = []
            patent_source = 'cnipa'

            # Short text: one-shot LLM call.  Long text: split into sections.
            if len(combined_text) < 100000:
                _pipeline_logger.info(
                    f"[task={task_id}] PHASE0 one_shot_extraction "
                    f"({len(combined_text)} chars)"
                )
                try:
                    extract_result = await flash_provider.complete_json(
                        EXTRACT_ALL_PROMPT, combined_text,
                    )
                    if extract_result and isinstance(extract_result, dict):
                        extracted = extract_result.get('patent_ids', [])
                        if isinstance(extracted, list):
                            if extract_result.get('source') == 'uspto':
                                patent_source = 'uspto'
                            for pid in extracted:
                                pid = str(pid).strip().replace(',', '').replace('/', '')
                                if pid.upper().startswith('US') and len(pid) > 2:
                                    pid = pid[2:]
                                pid = _re2.sub(r'[AB]\d$', '', pid)
                                if pid and pid.isdigit() and len(pid) == 8:
                                    normalized.append(pid)
                except Exception as e:
                    _pipeline_logger.warning(
                        f"[task={task_id}] PHASE0 one_shot LLM failed: {e}"
                    )
            else:
                # Large text: split into per-patent sections
                import re as _re3
                split_patterns = [
                    r'\n(?=---\n)',
                    r'\n(?=### )',
                    r'\n(?=\d+\.\s+\*\*)',
                    r'\n\n(?=\*\*[A-Z])',
                    r'\n\n(?=\d+\.\s)',
                ]
                best_sections = []
                for pat in split_patterns:
                    candidates = _re3.split(pat, combined_text)
                    candidates = [s.strip() for s in candidates if len(s.strip()) > 200]
                    if len(candidates) > len(best_sections):
                        best_sections = candidates
                sections = best_sections
                if len(sections) <= 1:
                    sections = []
                    start = 0
                    while start < len(combined_text):
                        end = min(start + 8000, len(combined_text))
                        chunk = combined_text[start:end].strip()
                        if len(chunk) > 200:
                            sections.append(chunk)
                        start = end - 500
                _pipeline_logger.info(
                    f"[task={task_id}] PHASE0 split_extraction — "
                    f"sections={len(sections)}, "
                    f"lengths={[len(s) for s in sections]}"
                )
                for i, section in enumerate(sections):
                    try:
                        result = await flash_provider.complete_json(
                            EXTRACT_ONE_PROMPT, section[:8000],
                        )
                        if result and isinstance(result, dict):
                            extracted = result.get('patent_ids', [])
                            if isinstance(extracted, list) and extracted:
                                if result.get('source') == 'uspto':
                                    patent_source = 'uspto'
                                for pid in extracted:
                                    pid = str(pid).strip().replace(',', '').replace('/', '')
                                    if pid.upper().startswith('US') and len(pid) > 2:
                                        pid = pid[2:]
                                    pid = _re2.sub(r'[AB]\d$', '', pid)
                                    if pid and pid.isdigit() and len(pid) == 8:
                                        normalized.append(pid)
                            _pipeline_logger.info(
                                f"[task={task_id}] PHASE0 section[{i}] — "
                                f"found={extracted}"
                            )
                    except Exception as e:
                        _pipeline_logger.warning(
                            f"[task={task_id}] PHASE0 section[{i}] LLM failed: {e}"
                        )

            patent_ids = sorted(set(normalized))
            # If all extracted IDs are pure 8-digit → almost certainly USPTO
            if patent_ids:
                all_uspto_like = all(
                    pid.isdigit() and len(pid) == 8 for pid in patent_ids
                )
                if all_uspto_like and patent_source != 'uspto':
                    _pipeline_logger.info(
                        f"[task={task_id}] PHASE0 auto_detect_uspto — "
                        f"all {len(patent_ids)} IDs are 8-digit, "
                        f"overriding source from '{patent_source}' to 'uspto'"
                    )
                    patent_source = 'uspto'
                params['patent_source'] = patent_source
            source_max = _get_max_patents_for_source(
                patent_source, max_patents,
                max_patents_cnipa, max_patents_uspto,
            )
            patent_ids = patent_ids[:source_max]
            _pipeline_logger.info(
                f"[task={task_id}] PHASE0 llm_extracted_patent_ids — "
                f"count={len(patent_ids)}, "
                f"source={params.get('patent_source', 'cnipa')}, "
                f"patent_ids={patent_ids}"
            )

        # Fallback: when conversation_refs yields no IDs from text (e.g. the
        # previous query was a large-list summary with download-only output),
        # try to retrieve patent IDs stored from the generated Excel/CSV artifacts.
        if not patent_ids and scenario == 'conversation_refs':
            user_id = params.get('user_id', '')
            if user_id:
                try:
                    from sources.knowledge.knowledge import get_redis_connection
                    r = get_redis_connection()
                    key = f"lt:conv:{user_id}:patent_ids"
                    stored = r.get(key)
                    if stored:
                        stored_ids = json.loads(stored)
                        if isinstance(stored_ids, list):
                            patent_ids = [str(pid).strip() for pid in stored_ids
                                          if str(pid).strip()]
                        _pipeline_logger.info(
                            f"[task={task_id}] PHASE0 fallback_from_stored_ids — "
                            f"count={len(patent_ids)}, key={key}"
                        )
                    else:
                        _pipeline_logger.info(
                            f"[task={task_id}] PHASE0 stored_ids_empty — "
                            f"key={key} not found in Redis"
                        )
                except Exception as e:
                    _pipeline_logger.warning(
                        f"[task={task_id}] PHASE0 stored_ids_lookup_failed — {e}"
                    )

        if patent_ids:
            total = len(patent_ids)
            # Only reset state for fresh runs — when resuming from checkpoint,
            # pending and table_rows were already restored above.
            if not (checkpoint and checkpoint.get('pending')):
                pending = patent_ids
                table_rows = []
            else:
                _pipeline_logger.info(
                    f"[task={task_id}] CHECKPOINT_RESUME — "
                    f"keeping checkpoint state: table_rows={len(table_rows)}, "
                    f"pending={len(pending)}"
                )

    # ==== Phase 0: Search patents via scene tools (if no patent_ids provided) ====
    if not patent_ids and scene_candidates:
        from sources.long_task.scene_tools import (
            select_tool,
            execute_tool,
            extract_patent_ids,
            extract_patent_id_url_map,
        )
        patent_source = params.get('patent_source', 'cnipa')
        update_task_status(task_id, 'searching_patents', 0,
                           _t('tool_select', batch_lang, count=len(scene_candidates)))
        selected = await select_tool(
            'search patents',
            f"专利来源: {patent_source}\n用户查询: {params['query']}",
            scene_candidates, flash_provider,
        )
        _pipeline_logger.info(
            f"[task={task_id}] PHASE0 select_tool — "
            f"selected={selected is not None}, "
            f"tool_title={selected.get('tool', {}).title if selected else 'N/A'}, "
            f"tool_url={selected.get('tool', {}).url if selected else 'N/A'}, "
            f"llm_params={json.dumps(selected.get('params', {}), ensure_ascii=False) if selected else 'N/A'}, "
            f"reason={selected.get('reason', '') if selected else ''}"
        )
        if selected:
            update_task_status(task_id, 'searching_patents', 2,
                               _t('tool_search', batch_lang, reason=selected.get("reason", "")))
            result = await execute_tool(selected['tool'], selected['params'])
            raw_items = result.get('raw_items', []) or []
            patent_ids = extract_patent_ids(raw_items)
            id_url_map = extract_patent_id_url_map(raw_items)
            from sources.long_task.scene_tools import extract_patent_id_pid_map
            id_pid_map = extract_patent_id_pid_map(raw_items)
            _pipeline_logger.info(
                f"[task={task_id}] PHASE0 search_result — "
                f"raw_items_count={len(raw_items)}, "
                f"patent_ids_found={len(patent_ids)}, "
                f"patent_ids={patent_ids[:10]}{'...' if len(patent_ids) > 10 else ''}, "
                f"id_url_map_size={len(id_url_map)}, "
                f"id_pid_map_size={len(id_pid_map)}"
            )
            params['patent_source'] = _infer_source_from_tool(
                selected['tool'], params.get('patent_source', 'cnipa'),
            )
        if patent_ids:
            source_max = _get_max_patents_for_source(
                params.get('patent_source', 'cnipa'), max_patents,
                max_patents_cnipa, max_patents_uspto,
            )
            if len(patent_ids) > source_max:
                _pipeline_logger.info(
                    f"[task={task_id}] PHASE0 truncating — "
                    f"source={params.get('patent_source', 'cnipa')}, "
                    f"max={source_max}, found={len(patent_ids)}, "
                    f"truncated_to={source_max}"
                )
                patent_ids = patent_ids[:source_max]
            total = len(patent_ids)
            # Only reset state for fresh runs — when resuming from checkpoint,
            # pending and table_rows were already restored above.
            if not (checkpoint and checkpoint.get('pending')):
                pending = patent_ids
                table_rows = []
            else:
                _pipeline_logger.info(
                    f"[task={task_id}] CHECKPOINT_RESUME_SCENE — "
                    f"keeping checkpoint state: table_rows={len(table_rows)}, "
                    f"pending={len(pending)}"
                )
            update_task_status(task_id, 'searching_patents', 5,
                               _t('search_complete', batch_lang, count=len(patent_ids)),
                               patent_ids=patent_ids)
        else:
            set_task_failed(task_id, _t('no_patents_found', batch_lang))
            user_id_for_analytics = params.get('user_id', '')
            if user_id_for_analytics:
                track_event("long_task:fail", user_id=user_id_for_analytics,
                            task_id=task_id,
                            extra={"error": "no_patents_found"})
            return {'status': 'failed', 'task_id': task_id,
                    'error': 'No patents found matching the search criteria'}

    # ==== Phase 1: Generate columns (Flash) ====
    if resume_columns:
        # Resuming a paused task — reuse columns from checkpoint
        columns = resume_columns
        _pipeline_logger.info(
            f"[task={task_id}] PHASE1 skipped — resuming with {len(columns)} "
            f"columns from checkpoint: {columns}"
        )
    else:
        update_task_status(task_id, 'generating_columns', 5,
                           _t('generating_framework', batch_lang, total=total))
        _pipeline_logger.info(
            f"[task={task_id}] PHASE1 generate_table_columns — "
            f"query={params['query'][:100]}, patent_count={total}, "
            f"provider={flash_provider.model if hasattr(flash_provider, 'model') else 'flash'}"
        )
        columns = await generate_table_columns(
            query=params['query'],
            patent_count=total,
            provider=flash_provider,
            lang=batch_lang,
        )
        _pipeline_logger.info(
            f"[task={task_id}] PHASE1 columns_generated — "
            f"column_count={len(columns)}, columns={columns}"
        )
        update_task_status(task_id, 'generating_columns', 5,
                           f'分析维度：{" | ".join(columns[1:4])}...',
                           table_columns=columns)
        # Update MySQL long_tasks table after phase 1
        _update_mysql_progress(task_id, 'generating_columns', 5)

    # ==== Phase 2: Per-patent download -> analyze -> summarize ====
    _pipeline_logger.info(
        f"[task={task_id}] PHASE2 START — pending_count={len(pending)}, "
        f"total={total}"
    )
    for i, patent_id in enumerate(pending):
        patent_index = len(table_rows) + 1

        # ── Pause / Stop checkpoint ──────────────────────────────────────
        _uid = params.get('user_id', '')
        _completed = len(table_rows)

        _result = _handle_task_stop(task_id, _uid, _completed, total)
        if _result:
            if patent_file_refs:
                import shutil as _shutil
                upload_dir = os.path.dirname(patent_file_refs[0]['path'])
                try:
                    _shutil.rmtree(upload_dir, ignore_errors=True)
                except Exception:
                    pass
            return _result

        _result = _handle_task_pause(task_id, _uid, _completed, total, {
            'completed': [r.get('专利号', '') for r in table_rows if not r.get('_failed')],
            'current': patent_id,
            'pending': pending[i:],
            'completed_rows': table_rows,
            'failed': [r.get('专利号', '') for r in table_rows if r.get('_failed')],
            'columns': columns,
        })
        if _result:
            return _result

        try:
            # Use patent_index-based progress so resumed tasks show correct %
            completed_before = len(table_rows)
            update_task_status(task_id, 'analyzing',
                               progress_pct(completed_before + i, total),
                               _t('downloading_patent', batch_lang, current=patent_index, total=total),
                               table_rows=table_rows)

            # Try scene tool download first, fall back to hardcoded download
            # In file-upload mode, use pre-extracted text from conversation_history
            fallback_binary = None
            if is_file_upload_mode:
                patent_text = patent_texts.get(patent_id, '')
                _pipeline_logger.info(
                    f"[task={task_id}] PHASE2 patent[{patent_index}/{total}] — "
                    f"patent_id={patent_id}, using_uploaded_text, "
                    f"text_length={len(patent_text)}"
                )
            else:
                patent_text, fallback_binary = await _download_patent_via_scene_or_fallback(
                    patent_id=patent_id,
                    params=params,
                    scene_candidates=scene_candidates,
                    flash_provider=flash_provider,
                    download_patent_document=download_patent_document,
                    id_url_map=id_url_map,
                    id_pid_map=id_pid_map,
                )

            _pipeline_logger.info(
                f"[task={task_id}] PHASE2 patent[{patent_index}/{total}] download_done — "
                f"patent_id={patent_id}, "
                f"text_length={len(patent_text) if patent_text else 0}, "
                f"binary_cached={fallback_binary is not None}, "
                f"binary_len={len(fallback_binary) if fallback_binary else 0}"
            )

            update_task_status(task_id, 'analyzing',
                               progress_pct(completed_before + i, total),
                               _t('analyzing', batch_lang, current=patent_index, total=total),
                               table_rows=table_rows)

            # ── Text extraction may be incomplete for short patents or
            #     scanned/image PDFs (pypdf returns little/nothing).  Threshold
            #     at 10k chars ensures we have enough text for meaningful analysis.
            #     Strategy: try vision (MiniMax-M3) first, then OCR as fallback.
            #     When vision_enabled=false, skip straight to OCR.
            text_ok = patent_text and len(patent_text) >= 1000
            want_vision = vision_provider is not None

            if text_ok:
                row = await analyze_single_patent(
                    patent_id=patent_id, patent_text=patent_text,
                    columns=columns, query=params['query'],
                    provider=pro_provider, timeout=60, lang=batch_lang,
                )
                _pipeline_logger.info(
                    f"[task={task_id}] PHASE2 patent[{patent_index}/{total}] analyze_done — "
                    f"patent_id={patent_id}, row_keys={list(row.keys()) if row else 'None'}"
                )
            else:
                # Text insufficient — need binary for vision or OCR
                pdf_bytes = None
                if fallback_binary is not None:
                    # Reuse binary already downloaded during text extraction
                    pdf_bytes = fallback_binary
                    _pipeline_logger.info(
                        f"[task={task_id}] PHASE2 patent[{patent_index}/{total}] "
                        f"vision_using_cached_binary — patent_id={patent_id}, "
                        f"len={len(pdf_bytes)}"
                    )
                elif is_file_upload_mode:
                    # Read the uploaded file from disk for vision/OCR
                    ref = next(
                        (r for r in patent_file_refs
                         if r['filename'].rsplit('.', 1)[0] == patent_id),
                        None,
                    )
                    if ref:
                        try:
                            with open(ref['path'], 'rb') as fh:
                                pdf_bytes = fh.read()
                            _pipeline_logger.info(
                                f"[task={task_id}] PHASE2 patent[{patent_index}/{total}] "
                                f"file_upload_binary_read — patent_id={patent_id}, "
                                f"file={ref['filename']}, "
                                f"len={len(pdf_bytes)}"
                            )
                        except Exception as e:
                            _pipeline_logger.warning(
                                f"[task={task_id}] PHASE2 patent[{patent_index}/{total}] "
                                f"file_upload_binary_read_failed — patent_id={patent_id}, "
                                f"error={e}"
                            )
                elif params.get('patent_source') == 'uspto':
                    pdf_bytes = await _download_uspto_binary_for_vision(
                        patent_id, flash_provider,
                    )
                if not pdf_bytes:
                    row = build_failed_row(patent_id,
                        "PDF text extraction failed and could not download binary",
                        lang=batch_lang)
                elif want_vision:
                    # Path A: MiniMax-M3 vision → OCR fallback
                    from sources.long_task.patent_analyzer import analyze_patent_with_vision
                    row = await analyze_patent_with_vision(
                        pdf_bytes=pdf_bytes, patent_id=patent_id,
                        columns=columns, query=params['query'],
                        vision_provider=vision_provider, lang=batch_lang,
                    )
                    if row.get('_failed'):
                        _pipeline_logger.info(
                            f"[task={task_id}] PHASE2 patent[{patent_index}/{total}] "
                            f"vision_failed_fallback_to_ocr — patent_id={patent_id}"
                        )
                        ocr_text = _ocr_from_pdf_reader(pdf_bytes)
                        if ocr_text and len(ocr_text) >= 1000:
                            row = await analyze_single_patent(
                                patent_id=patent_id, patent_text=ocr_text,
                                columns=columns, query=params['query'],
                                provider=pro_provider, timeout=60, lang=batch_lang,
                            )
                else:
                    # Path B: Vision disabled — straight to OCR
                    _pipeline_logger.info(
                        f"[task={task_id}] PHASE2 patent[{patent_index}/{total}] "
                        f"vision_disabled_ocr — patent_id={patent_id}"
                    )
                    ocr_text = _ocr_from_pdf_reader(pdf_bytes)
                    if ocr_text and len(ocr_text) >= 1000:
                        row = await analyze_single_patent(
                            patent_id=patent_id, patent_text=ocr_text,
                            columns=columns, query=params['query'],
                            provider=pro_provider, timeout=60, lang=batch_lang,
                        )
                    else:
                        row = build_failed_row(patent_id,
                            f"Text extraction and OCR both failed ({len(ocr_text) if ocr_text else 0} chars)",
                            lang=batch_lang)

            row['_summary'] = await generate_patent_summary(
                patent_id=patent_id, row=row, query=params['query'],
                provider=pro_provider, lang=batch_lang,
            )
            _pipeline_logger.info(
                f"[task={task_id}] PHASE2 patent[{patent_index}/{total}] summary_done — "
                f"patent_id={patent_id}, "
                f"summary_length={len(row.get('_summary', '')) if row.get('_summary') else 0}"
            )

        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] PHASE2 patent[{patent_index}/{total}] FAILED — "
                f"patent_id={patent_id}, error={e}"
            )
            row = build_failed_row(patent_id, str(e), lang=batch_lang)

        table_rows.append(row)
        _pipeline_logger.info(
            f"[task={task_id}] PHASE2 table_updated — "
            f"completed={len(table_rows)}/{total}, "
            f"successful={len([r for r in table_rows if not r.get('_failed')])}, "
            f"failed={len([r for r in table_rows if r.get('_failed')])}"
        )
        save_checkpoint(task_id, {
            'completed': [r.get('专利号', patent_id) for r in table_rows if not r.get('_failed')],
            'current': patent_id,
            'pending': pending[i+1:],
            'completed_rows': table_rows,
            'failed': [r.get('专利号', patent_id) for r in table_rows if r.get('_failed')],
            'columns': columns,
        })

        update_task_status(task_id, 'analyzing',
                           progress_pct(completed_before + i + 1, total),
                           _t('analysis_completed_count', batch_lang, current=len(table_rows), total=total),
                           table_rows=table_rows)
    # Clean up temp upload directory after Phase 2 (vision fallback may have read files)
    if patent_file_refs:
        import shutil as _shutil
        upload_dir = os.path.dirname(patent_file_refs[0]['path'])
        try:
            _shutil.rmtree(upload_dir, ignore_errors=True)
            _pipeline_logger.info(
                f"[task={task_id}] FILE_CLEANUP — removed {upload_dir}"
            )
        except Exception:
            pass
    # Update MySQL long_tasks table after phase 2
    _pipeline_logger.info(
        f"[task={task_id}] PHASE2 COMPLETE — "
        f"total_rows={len(table_rows)}, table_columns={columns}"
    )
    _update_mysql_progress(task_id, 'analyzing', 75)

    # ==== Phase 3: Generate report (Pro, dynamic) ====
    _pipeline_logger.info(
        f"[task={task_id}] PHASE3 generate_report — "
        f"columns={columns}, table_rows_count={len(table_rows)}, "
        f"provider={pro_provider.model if hasattr(pro_provider, 'model') else 'pro'}"
    )

    from sources.long_task.status_manager import ThrottledSummaryUpdater
    summary_updater = ThrottledSummaryUpdater(
        task_id, progress=76, step_msg=_t('writing_summary', batch_lang),
    )
    report_title = _t('default_report_title', batch_lang)
    exec_heading = _t('default_exec_summary_heading', batch_lang)

    def _assemble_report(
        exec_summary: str | None,
        completed_parts: list[str],
        current_heading: str | None = None,
        current_text: str = '',
    ) -> str:
        parts = [f"# {report_title}\n\n"]
        if exec_summary:
            parts.append(f"## {exec_heading}\n\n{exec_summary}\n\n")
        parts.extend(completed_parts)
        if current_heading and current_text:
            parts.append(f"## {current_heading}\n\n{current_text}")
        return "".join(parts)

    # ── Executive Summary ──
    update_task_status(task_id, 'generating_report', 76,
                       _t('writing_summary', batch_lang))

    def _exec_chunk(partial: str) -> None:
        summary_updater.push(
            _assemble_report(partial, []),
            step_msg=_t('writing_summary', batch_lang),
        )

    try:
        exec_summary = await generate_executive_summary(
            table_rows=table_rows,
            columns=columns,
            query=params['query'],
            provider=pro_provider,
            lang=batch_lang,
            on_chunk=_exec_chunk,
        )
        _pipeline_logger.info(
            f"[task={task_id}] PHASE3 exec_summary — "
            f"length={len(exec_summary)}"
        )
        summary_updater.push(
            _assemble_report(exec_summary, []),
            progress=78,
            step_msg=_t('writing_summary', batch_lang),
            force=True,
        )
    except Exception as e:
        _pipeline_logger.error(
            f"[task={task_id}] PHASE3 exec_summary FAILED — {e}"
        )
        exec_summary = None

    # ── Outline ──
    update_task_status(task_id, 'generating_report', 80,
                       _t('generating_outline', batch_lang))
    try:
        outline = await generate_report_outline(
            query=params['query'], columns=columns,
            table_rows=table_rows, provider=pro_provider,
            lang=batch_lang,
        )
    except Exception as e:
        _pipeline_logger.error(
            f"[task={task_id}] PHASE3 outline FAILED — {e}, falling back to default"
        )
        outline = {
            'title': _t('default_report_title', batch_lang),
            'sections': [{'heading': _t('default_section_heading', batch_lang), 'description': ''}],
        }
    report_title = outline.get('title', report_title)
    _pipeline_logger.info(
        f"[task={task_id}] PHASE3 outline — "
        f"title={outline.get('title', '')}, "
        f"sections_count={len(outline.get('sections', []))}, "
        f"sections={[s.get('heading', '') for s in outline.get('sections', [])]}"
    )

    report_parts = []
    sections = outline.get('sections', [{'heading': _t('default_section_heading', batch_lang), 'description': ''}])
    for idx, section in enumerate(sections):
        sec_pct = 80 + int((idx + 1) / len(sections) * 10)
        step_msg = _t('writing_section', batch_lang, heading=section["heading"])
        update_task_status(task_id, 'generating_report', sec_pct, step_msg)
        summary_updater.progress = sec_pct
        summary_updater.step_msg = step_msg

        def _section_chunk(partial: str, _heading=section['heading']) -> None:
            summary_updater.push(
                _assemble_report(
                    exec_summary,
                    report_parts,
                    current_heading=_heading,
                    current_text=partial,
                ),
                step_msg=step_msg,
            )

        try:
            text = await generate_report_section(
                section=section, query=params['query'],
                columns=columns, table_rows=table_rows,
                provider=pro_provider,
                lang=batch_lang,
                on_chunk=_section_chunk,
            )
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] PHASE3 section[{idx+1}/{len(sections)}] FAILED — {e}"
            )
            text = _t('section_failed', batch_lang, heading=section['heading'])
        section_md = f"## {section['heading']}\n\n{text}"
        report_parts.append(section_md)
        summary_updater.push(
            _assemble_report(exec_summary, report_parts),
            progress=sec_pct,
            step_msg=step_msg,
            force=True,
        )
        _pipeline_logger.info(
            f"[task={task_id}] PHASE3 section[{idx+1}/{len(sections)}] — "
            f"heading={section['heading']}, text_length={len(text)}"
        )

    # ── Assemble report ──
    exec_section = (
        f"## {exec_heading}\n\n{exec_summary}\n\n"
        if exec_summary else ""
    )
    report_text = (
        f"# {outline.get('title', _t('default_report_title', batch_lang))}\n\n"
        + exec_section
        + "\n\n".join(report_parts)
    )
    _pipeline_logger.info(
        f"[task={task_id}] PHASE3 report_text — "
        f"total_length={len(report_text)}, sections_written={len(report_parts)}, "
        f"has_exec_summary={exec_summary is not None}"
    )
    update_task_status(task_id, 'generating_report', 90,
                       _t('report_writing_complete', batch_lang), result_summary=report_text)
    _update_mysql_progress(task_id, 'generating_report', 90, result_summary=report_text)

    # ==== Phase 4: Export files ====
    from sources.long_task.storage import get_storage_config
    storage_cfg = get_storage_config()
    try:
        storage = create_storage(storage_cfg)
    except Exception as e:
        _pipeline_logger.error(
            f"[task={task_id}] PHASE4 storage_init FAILED — {e}, falling back to local"
        )
        storage = create_storage({'report_storage_backend': 'local'})
    _pipeline_logger.info(
        f"[task={task_id}] PHASE4 export — "
        f"storage_backend={storage_cfg.get('report_storage_backend', 'local')}"
    )

    report_files = []
    local_storage = None

    async def _get_local_storage():
        """Lazy-init local storage for fallback."""
        nonlocal local_storage
        if local_storage is None:
            from sources.long_task.storage import create_storage as _create
            local_storage = _create({'report_storage_backend': 'local'})
        return local_storage

    # ── Generate and upload DOCX ──
    update_task_status(task_id, 'exporting', 90, _t('generating_word', batch_lang))
    docx_bytes = await export_docx_async(report_text, table_rows, columns, lang=batch_lang)
    try:
        await storage.put(task_id, 'report.docx', docx_bytes)
        _pipeline_logger.info(
            f"[task={task_id}] PHASE4 docx — size_bytes={len(docx_bytes)}"
        )
        report_files.append({'format': 'docx', 'filename': 'report.docx', 'size': len(docx_bytes)})
    except Exception as e:
        _pipeline_logger.error(
            f"[task={task_id}] PHASE4 docx upload FAILED — {e}, falling back to local"
        )
        try:
            await (await _get_local_storage()).put(task_id, 'report.docx', docx_bytes)
            _pipeline_logger.info(
                f"[task={task_id}] PHASE4 docx — local fallback OK, size_bytes={len(docx_bytes)}"
            )
            report_files.append({'format': 'docx', 'filename': 'report.docx', 'size': len(docx_bytes)})
        except Exception as e2:
            _pipeline_logger.error(
                f"[task={task_id}] PHASE4 docx local fallback ALSO FAILED — {e2}"
            )

    # ── Generate and upload PDF ──
    update_task_status(task_id, 'exporting', 95, _t('generating_pdf', batch_lang))
    pdf_bytes = await export_pdf_async(docx_bytes)
    try:
        await storage.put(task_id, 'report.pdf', pdf_bytes)
        _pipeline_logger.info(
            f"[task={task_id}] PHASE4 pdf — size_bytes={len(pdf_bytes)}"
        )
        report_files.append({'format': 'pdf', 'filename': 'report.pdf', 'size': len(pdf_bytes)})
    except Exception as e:
        _pipeline_logger.error(
            f"[task={task_id}] PHASE4 pdf upload FAILED — {e}, falling back to local"
        )
        try:
            await (await _get_local_storage()).put(task_id, 'report.pdf', pdf_bytes)
            _pipeline_logger.info(
                f"[task={task_id}] PHASE4 pdf — local fallback OK, size_bytes={len(pdf_bytes)}"
            )
            report_files.append({'format': 'pdf', 'filename': 'report.pdf', 'size': len(pdf_bytes)})
        except Exception as e2:
            _pipeline_logger.error(
                f"[task={task_id}] PHASE4 pdf local fallback ALSO FAILED — {e2}"
            )

    _pipeline_logger.info(
        f"[task={task_id}] COMPLETED — "
        f"patents_analyzed={len(table_rows)}, "
        f"report_files={[f['format'] for f in report_files]}"
    )
    _update_mysql_progress(task_id, 'exporting', 100)

    # Extract patent IDs from result table for the frontend to store in
    # conversation messages (hidden field), enabling follow-up conversation_refs
    # queries without relying solely on Redis.
    _completed_patent_ids = []
    _id_col = columns[0] if columns else None
    if _id_col:
        for _row in table_rows:
            _pid = str(_row.get(_id_col, '')).strip()
            if _pid and not _row.get('_failed'):
                _completed_patent_ids.append(_pid)

    set_task_completed(task_id, report_files,
                       patent_ids=_completed_patent_ids if _completed_patent_ids else None)

    # Store analyzed patent IDs so follow-up conversation_refs queries
    # (e.g. "哪些是AI相关的") can find them even when the previous query
    # was a long task rather than a general agent query with artifacts.
    _store_long_task_patent_ids(
        task_id=task_id,
        table_rows=table_rows,
        columns=columns,
        user_id=params.get('user_id', ''),
    )

    user_id_for_analytics = params.get('user_id', '')
    if user_id_for_analytics:
        track_event("long_task:complete", user_id=user_id_for_analytics,
                    task_id=task_id, patent_count=len(table_rows),
                    patent_source=params.get('patent_source', ''))

    # ── Per-user queue: dispatch next queued task if any ──
    user_id = params.get('user_id', '')
    if user_id:
        try:
            from sources.long_task.user_queue import complete_user_task
            next_task_id = complete_user_task(str(user_id), task_id)
            if next_task_id:
                _dispatch_queued_task(next_task_id, user_id)
        except Exception as e:
            import traceback
            _pipeline_logger.warning(
                f"[task={task_id}] QUEUE_DISPATCH_FAILED — {type(e).__name__}: {e}"
            )
            _pipeline_logger.warning(
                f"[task={task_id}] QUEUE_DISPATCH_TRACEBACK —\n{traceback.format_exc()}"
            )

    return {'status': 'completed', 'task_id': task_id}


# ═══════════════════════════════════════════════════════════════════════════════
# Patent Family Analysis Task (cross-jurisdiction, US first)
# ═══════════════════════════════════════════════════════════════════════════════


@app.task(bind=True, max_retries=2, default_retry_delay=60, time_limit=3600, soft_time_limit=3540)
def execute_family_analysis(self, task_id: str, params: dict):
    """Cross-jurisdiction patent family prosecution analysis.

    Phase 0: EPO OPS family lookup → extract US member → normalize app_number.
    Phase 1-4: Same US prosecution pipeline as ``execute_prosecution_analysis``.
    """
    import asyncio as _asyncio

    retry_count = self.request.retries
    patent_id = str(params.get('patent_id', '')).strip()
    query = params.get('query', '')
    lang = params.get('lang', 'zh')
    session_id = params.get('session_id', '')
    user_id = params.get('user_id', '')

    _pipeline_logger.info(
        f"[task={task_id}] FAMILY START — "
        f"patent_id={patent_id}, "
        f"query={query[:120]}, "
        f"lang={lang}, "
        f"session_id={session_id}, "
        f"retry={retry_count}/{self.max_retries}"
    )

    if retry_count >= self.max_retries:
        _pipeline_logger.error(
            f"[task={task_id}] FAMILY HARD_STOP — retry_count={retry_count} >= {self.max_retries}"
        )
        if user_id:
            from sources.long_task.user_queue import complete_user_task
            try:
                complete_user_task(str(user_id), task_id)
            except Exception:
                pass
        return {'status': 'failed', 'task_id': task_id,
                'error': f'Max retries ({self.max_retries}) exceeded'}

    if not patent_id:
        _pipeline_logger.error(f"[task={task_id}] FAMILY no patent_id provided")
        _update_mysql_progress(task_id, 'failed', 0)
        return {'status': 'failed', 'task_id': task_id, 'error': 'No patent_id'}

    # ── Imports ──────────────────────────────────────────────────────────────
    from sources.long_task.status_manager import (
        update_task_status, set_task_completed, set_task_failed,
    )
    from sources.long_task.config import (
        get_long_task_config, get_family_config, get_prosecution_config,
        get_sipop_config, get_jpo_config,
        DEFAULT_VISION_PROVIDER, DEFAULT_VISION_MODEL,
    )
    from sources.long_task.patent_family import EPOFamilyClient, EPOError
    from sources.long_task.family_member import PatentFamily
    from sources.long_task.prosecution_downloader import (
        classify_prosecution_documents,
        download_single_document,
    )
    from sources.long_task.storage import create_storage
    from sources.llm_provider import Provider

    # ── Provider setup ───────────────────────────────────────────────────────
    ltc = get_long_task_config()
    model_family = ltc['provider_family']

    if model_family == 'minimax':
        flash_provider = Provider(
            provider_name='minimax', model='MiniMax-M2.7-highspeed',
            server_address='', is_local=False,
        )
        pro_provider = Provider(
            provider_name='minimax', model='MiniMax-M3',
            server_address='', is_local=False,
        )
    else:
        flash_provider = Provider(
            provider_name='deepseek', model='deepseek-v4-flash',
            server_address='', is_local=False,
        )
        pro_provider = Provider(
            provider_name='deepseek', model='deepseek-v4-pro',
            server_address='', is_local=False,
        )

    ptc = get_prosecution_config()
    include_priority_2 = ptc.get('include_priority_2', True)
    vision_enabled = ltc.get('vision_enabled', True)
    vision_provider = None
    if vision_enabled:
        vision_cfg_provider = ltc.get('vision_provider', DEFAULT_VISION_PROVIDER)
        vision_cfg_model = ltc.get('vision_model', DEFAULT_VISION_MODEL)
        vision_provider = Provider(
            provider_name=vision_cfg_provider, model=vision_cfg_model,
            server_address='', is_local=False,
        )

    # ── Family config ────────────────────────────────────────────────────────
    fc = get_family_config()
    epo_key = fc.get('epo_consumer_key', '')
    epo_secret = fc.get('epo_consumer_secret', '')

    if not epo_key or not epo_secret:
        _pipeline_logger.error(
            f"[task={task_id}] FAMILY missing EPO credentials — "
            f"set [FAMILY] epo_consumer_key / epo_consumer_secret in config.ini"
        )
        set_task_failed(task_id, "EPO API credentials not configured")
        _update_mysql_progress(task_id, 'failed', 0)
        return {'status': 'failed', 'task_id': task_id,
                'error': 'EPO credentials not configured'}

    # ── Run pipeline ─────────────────────────────────────────────────────────
    async def _run():
        import os as _os
        # ═════════════════════════════════════════════════════════════════
        # Phase 0: EPO family lookup → extract US member
        # ═════════════════════════════════════════════════════════════════
        update_task_status(task_id, 'preparing', 1,
                           _t('family_lookup', lang),
                           analysis_type='family')

        epo_client = EPOFamilyClient(
            consumer_key=epo_key,
            consumer_secret=epo_secret,
        )

        try:
            family = await epo_client.lookup_family(patent_id)
        except EPOError as e:
            _pipeline_logger.error(
                f"[task={task_id}] FAMILY PHASE0 epo_error — {e}"
            )
            set_task_failed(task_id, f"EPO family lookup failed: {e}")
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': str(e)}

        _pipeline_logger.info(
            f"[task={task_id}] FAMILY PHASE0 epo_ok — "
            f"family_id={family.family_id}, "
            f"total_count={family.total_count}, "
            f"jurisdictions={family.jurisdictions}, "
            f"dedup_members={len(family.deduplicated_members)}"
        )

        # ── Extract US member ────────────────────────────────────────────
        us_member = family.get_representative('US')
        if not us_member:
            _pipeline_logger.warning(
                f"[task={task_id}] FAMILY PHASE0 no_us_member — "
                f"jurisdictions={family.jurisdictions}"
            )
            if lang == 'zh':
                msg = (
                    f"在EPO同族专利数据库中未找到 {patent_id} 的美国同族成员。"
                    f"该专利的同族覆盖: {', '.join(family.jurisdictions)}。"
                )
            else:
                msg = (
                    f"No US family member found for {patent_id} in the EPO "
                    f"family database. Jurisdictions found: {', '.join(family.jurisdictions)}."
                )
            set_task_failed(task_id, msg)
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': msg}

        us_pub_number = us_member.pub_number
        us_app_number = us_member.normalized_app_number

        _pipeline_logger.info(
            f"[task={task_id}] FAMILY PHASE0 us_member — "
            f"pub_number={us_pub_number}, "
            f"pub_kind={us_member.pub_kind}, "
            f"is_granted={us_member.is_granted}, "
            f"app_number_raw={us_member.app_number}, "
            f"app_number_normalized={us_app_number}, "
            f"title={us_member.title[:80] if us_member.title else '(none)'}"
        )

        if not us_app_number or len(us_app_number) < 8:
            _pipeline_logger.error(
                f"[task={task_id}] FAMILY PHASE0 invalid_app_number — "
                f"raw={us_member.app_number}, normalized={us_app_number}"
            )
            if lang == 'zh':
                msg = f"无法从EPO同族数据中提取有效的美国专利申请号 (原始: {us_member.app_number})"
            else:
                msg = f"Cannot extract valid US application number from EPO family data (raw: {us_member.app_number})"
            set_task_failed(task_id, msg)
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': msg}

        # Update status with family overview for frontend
        family_overview = _build_family_overview_status(family, lang)
        update_task_status(task_id, 'preparing', 5,
                           family_overview.get('step_msg', ''),
                           family_overview=family_overview,
                           analysis_type='family')

        # ── Initialize per-jurisdiction progress tracking ──
        # EP member states (DE, GB, FR, etc.) are merged into a single "EP" row
        # because they share the same EPO examination process.
        _EP_MEMBER_STATES = frozenset({
            'AT', 'BE', 'BG', 'CH', 'CY', 'CZ', 'DE', 'DK', 'EE', 'ES',
            'FI', 'FR', 'GB', 'GR', 'HR', 'HU', 'IE', 'IS', 'IT', 'LI',
            'LT', 'LU', 'LV', 'MC', 'MK', 'MT', 'NL', 'NO', 'PL', 'PT',
            'RO', 'SE', 'SI', 'SK', 'SM', 'TR',
        })
        _JUR_LABELS = {'US': '美国 (US)', 'CN': '中国 (CN)', 'JP': '日本 (JP)',
                       'EP': '欧洲 (EPO)', 'WO': 'PCT (WO)', 'KR': '韩国 (KR)'}
        _jurisdictions: list[dict] = []
        _seen_ep = False  # merge all EP member states into one row
        _priority_order = ['US', 'CN', 'JP', 'EP', 'WO', 'KR']
        _all_codes = list(family.jurisdictions)
        # Sort: priority countries first, then by code
        _all_codes.sort(key=lambda c: (
            _priority_order.index(c) if c in _priority_order else 99, c,
        ))
        for _code in _all_codes:
            if _code in _EP_MEMBER_STATES:
                if not _seen_ep:
                    _seen_ep = True
                    _jurisdictions.append({
                        'code': 'EP',
                        'label': _JUR_LABELS.get('EP', 'EP'),
                        'status': 'pending',
                        'progress': 0,
                        'detail': '',
                        'file_count': 0,
                        'files_done': 0,
                    })
                continue  # skip individual member states
            _jurisdictions.append({
                'code': _code,
                'label': _JUR_LABELS.get(_code, _code),
                'status': 'pending',
                'progress': 0,
                'detail': '',
                'file_count': 0,
                'files_done': 0,
            })

        def _push_jurisdictions(_current_progress: int, _current_phase: str,
                                _step_label: str):
            """Push jurisdiction statuses to Redis."""
            update_task_status(task_id, _current_phase, _current_progress,
                               _step_label,
                               jurisdictions=_jurisdictions,
                               analysis_type='family')

        # ═════════════════════════════════════════════════════════════════
        # Phase 0.3: China examination data (if CN member exists)
        # ═════════════════════════════════════════════════════════════════
        cn_exam_data: dict = {}  # stored for report merging later
        cn_member = family.get_representative('CN')
        if cn_member:
            cn_app_number = cn_member.app_number
            cn_app_number = __import__('re').sub(
                r'^CN\s*', '', cn_app_number or '', flags=__import__('re').IGNORECASE,
            ).replace('.', '').replace('-', '').replace(' ', '')
            _pipeline_logger.info(
                f"[task={task_id}] FAMILY PHASE0.3 cn_member — "
                f"cn_app={cn_app_number}, pub={cn_member.pub_number}"
            )

            sipop_cfg = get_sipop_config()
            sipop_key = sipop_cfg.get('app_key', '')
            sipop_secret = sipop_cfg.get('app_secret', '')
            if sipop_key and sipop_secret and cn_app_number:
                try:
                    # Progress update for frontend — show CN examination phase
                    update_task_status(task_id, 'preparing', 55,
                                       _t('china_fetching_for_family', lang,
                                          cn_app=cn_app_number))
                    from sources.sipop_client import SipopClient
                    sipop = SipopClient(app_key=sipop_key, app_secret=sipop_secret)
                    from sources.long_task.china_examination import (
                        fetch_examination_data,
                        build_examination_timeline,
                        format_claims_for_report,
                        build_legal_status_timeline,
                    )
                    cn_events, cn_law, cn_info, cn_claims, cn_legal = await fetch_examination_data(
                        cn_app_number, sipop,
                    )
                    cn_exam_data = {
                        'cn_app_number': cn_app_number,
                        'events': cn_events,
                        'law_state': cn_law,
                        'basic_info': cn_info,
                        'claims': cn_claims,
                        'legal_timeline': cn_legal,
                        'timeline_md': await build_examination_timeline(
                            cn_events, cn_law, cn_info, lang,
                        ) if cn_events else '',
                        'claims_md': format_claims_for_report(cn_claims, lang) if cn_claims else '',
                        'legal_md': build_legal_status_timeline(cn_legal, lang) if cn_legal else '',
                    }
                    _pipeline_logger.info(
                        f"[task={task_id}] FAMILY PHASE0.3 cn_data — "
                        f"events={len(cn_events)}, claims={len(cn_claims)}, "
                        f"legal_events={len(cn_legal)}"
                    )
                    # Update CN jurisdiction status
                    for _j in _jurisdictions:
                        if _j['code'] == 'CN':
                            if cn_events:
                                _j['status'] = 'done'
                                _j['progress'] = 100
                                _j['detail'] = _t('family_cn_done', lang,
                                                  events=len(cn_events),
                                                  claims=len(cn_claims))
                            else:
                                _j['status'] = 'done'  # basic data only
                                _j['progress'] = 70
                                _j['detail'] = _t('family_cn_basic', lang)
                except Exception as e:
                    _pipeline_logger.warning(
                        f"[task={task_id}] FAMILY PHASE0.3 cn_fetch_failed — {e}"
                    )
                    for _j in _jurisdictions:
                        if _j['code'] == 'CN':
                            _j['status'] = 'no_data'
                            _j['detail'] = str(e)[:60]

        # ═════════════════════════════════════════════════════════════════
        # Phase 0.4: Japan examination data (if JP member exists)
        # ═════════════════════════════════════════════════════════════════
        jp_exam_data: dict = {}  # stored for report merging later
        jp_member = family.get_representative('JP')
        if jp_member:
            jp_app_number = jp_member.app_number
            jp_app_number = __import__('re').sub(
                r'^JP\s*', '', jp_app_number or '', flags=__import__('re').IGNORECASE,
            ).replace('.', '').replace('-', '').replace(' ', '')
            _pipeline_logger.info(
                f"[task={task_id}] FAMILY PHASE0.4 jp_member — "
                f"jp_app={jp_app_number}, pub={jp_member.pub_number}"
            )

            jpo_cfg = get_jpo_config()
            jpo_user = jpo_cfg.get('username', '')
            jpo_pass = jpo_cfg.get('password', '')
            if jpo_user and jpo_pass and jp_app_number:
                try:
                    update_task_status(task_id, 'preparing', 6,
                                       _t('japan_fetching_for_family', lang,
                                          jp_app=jp_app_number))
                    from sources.jpo_client import JpoClient
                    from sources.long_task.japan_examination import (
                        fetch_examination_data,
                        build_examination_timeline,
                        build_registration_summary,
                        build_citations_summary,
                    )
                    jpo = JpoClient(username=jpo_user, password=jpo_pass)
                    jp_data = await fetch_examination_data(jp_app_number, jpo)

                    # ── Fallback: EPO DOCDB app_number format may not match
                    # what the JPO API expects.  If the primary fetch returned
                    # zero progress events, try alternative number formats.
                    if jp_data.get('progress_count', 0) == 0:
                        fallback_candidates = []

                        # Candidate 1: try different app_number formats
                        if len(jp_app_number) >= 10 and jp_app_number[:4].isdigit():
                            _base = __import__('re').sub(r'[^\d]', '', jp_app_number)
                            if len(_base) == 10:
                                fallback_candidates.append(
                                    ('app_fmt', f"{_base[:4]}-{_base[4:]}")
                                )  # 2019159764 → 2019-159764
                            fallback_candidates.append(
                                ('app_fmt', f"JP{_base}")
                            )  # → JP2019159764
                            fallback_candidates.append(
                                ('app_fmt', f"JP{_base[:4]}-{_base[4:]}")
                            )  # → JP2019-159764

                        # Candidate 2: resolve via publication number.
                        # JPO API expects full format: JP{pub_number}{kind}
                        if jp_member.pub_number:
                            jp_pub = str(jp_member.pub_number).strip()
                            jp_kind = str(getattr(jp_member, 'pub_kind', '') or '').strip()
                            # Try with kind code first, then without
                            fb_pubs = []
                            if jp_kind:
                                fb_pubs.append(f"JP{jp_pub}{jp_kind}")
                            fb_pubs.append(f"JP{jp_pub}")
                            for fb_pub in fb_pubs:
                                fallback_candidates.append(('pub_lookup', fb_pub))

                        for fb_type, fb_value in fallback_candidates:
                            _pipeline_logger.info(
                                f"[task={task_id}] FAMILY PHASE0.4 jp_fallback — "
                                f"type={fb_type}, value={fb_value}"
                            )
                            try:
                                if fb_type == 'pub_lookup':
                                    from sources.jpo_client import JpoAPIError as _JpoErr
                                    ref = await jpo.lookup_number_relation(
                                        'publication', fb_value,
                                    )
                                    fb_app = (
                                        ref.get('applicationNumber')
                                        or ref.get('applicationDocNum')
                                        or ref.get('appNumber')
                                        or ''
                                    )
                                    if fb_app:
                                        fb_app = str(fb_app).strip()
                                        _pipeline_logger.info(
                                            f"[task={task_id}] FAMILY PHASE0.4 "
                                            f"jp_fallback — pub_lookup resolved "
                                            f"app_number={fb_app}"
                                        )
                                        fallback_jp_data = await fetch_examination_data(
                                            fb_app, jpo,
                                        )
                                    else:
                                        _pipeline_logger.warning(
                                            f"[task={task_id}] FAMILY PHASE0.4 "
                                            f"jp_fallback — pub_lookup returned "
                                            f"no app_number, ref_keys={list(ref.keys())}"
                                        )
                                        continue
                                elif fb_type == 'app_fmt':
                                    fallback_jp_data = await fetch_examination_data(
                                        fb_value, jpo,
                                    )

                                if fallback_jp_data.get('progress_count', 0) > 0:
                                    jp_data = fallback_jp_data
                                    if fb_type == 'pub_lookup':
                                        jp_app_number = str(ref.get(
                                            'applicationNumber',
                                            ref.get('applicationDocNum', ''),
                                        )).strip()
                                    else:
                                        jp_app_number = fb_value
                                    _pipeline_logger.info(
                                        f"[task={task_id}] FAMILY PHASE0.4 "
                                        f"jp_fallback — SUCCESS with type={fb_type}, "
                                        f"events={jp_data.get('progress_count', 0)}"
                                    )
                                    break
                            except Exception as _fe:
                                _pipeline_logger.warning(
                                    f"[task={task_id}] FAMILY PHASE0.4 "
                                    f"jp_fallback_{fb_type}_failed — {_fe}"
                                )

                    jp_exam_data = {
                        'jp_app_number': jp_app_number,
                        'progress': jp_data.get('progress', []),
                        'registration': jp_data.get('registration'),
                        'citations': jp_data.get('citations'),
                        'refusal_reasons': jp_data.get('refusal_reasons'),
                        'amendments': jp_data.get('amendments'),
                        'timeline_md': build_examination_timeline(jp_data, lang),
                        'registration_md': build_registration_summary(jp_data, lang),
                        'citations_md': build_citations_summary(jp_data, lang),
                    }
                    _pipeline_logger.info(
                        f"[task={task_id}] FAMILY PHASE0.4 jp_data — "
                        f"events={jp_data.get('progress_count', 0)}, "
                        f"has_reg={jp_data.get('has_registration')}, "
                        f"has_cites={jp_data.get('has_citations')}, "
                        f"has_refusal={jp_data.get('has_refusal_reasons')}, "
                        f"has_amendments={jp_data.get('has_amendments')}"
                    )
                    for _j in _jurisdictions:
                        if _j['code'] == 'JP':
                            _pc = jp_data.get('progress_count', 0)
                            if _pc > 0 or jp_data.get('has_registration'):
                                _j['status'] = 'done'; _j['progress'] = 100
                                _j['detail'] = _t('family_jp_done', lang, events=_pc)
                            else:
                                _j['status'] = 'done'; _j['progress'] = 70
                                _j['detail'] = _t('family_jp_basic', lang)
                except Exception as e:
                    _pipeline_logger.warning(
                        f"[task={task_id}] FAMILY PHASE0.4 jp_fetch_failed — {e}"
                    )
                    for _j in _jurisdictions:
                        if _j['code'] == 'JP':
                            _j['status'] = 'no_data'
                            _j['detail'] = str(e)[:60]

        # ═════════════════════════════════════════════════════════════════
        # Phase 0.45: EPO examination data (if EP member exists)
        # ═════════════════════════════════════════════════════════════════
        ep_exam_data: dict = {}
        ep_member = family.get_representative('EP')
        if ep_member:
            ep_app_number = ep_member.app_number
            ep_app_number = __import__('re').sub(
                r'^EP\s*', '', ep_app_number or '', flags=__import__('re').IGNORECASE,
            ).replace('.', '').replace('-', '').replace(' ', '')
            _pipeline_logger.info(
                f"[task={task_id}] FAMILY PHASE0.45 ep_member — "
                f"ep_app={ep_app_number}, pub={ep_member.pub_number}"
            )

            ep_family_cfg = get_family_config()
            ep_ck = ep_family_cfg.get('epo_consumer_key', '')
            ep_cs = ep_family_cfg.get('epo_consumer_secret', '')
            # Also check env vars (same fallback as EPOFamilyClient)
            import os as _os_ep
            ep_ck = _os_ep.getenv('EPO_CONSUMER_KEY', ep_ck)
            ep_cs = _os_ep.getenv('EPO_CONSUMER_SECRET', ep_cs)
            if ep_ck and ep_cs and ep_app_number:
                try:
                    update_task_status(task_id, 'preparing', 7,
                                       _t('epo_fetching_for_family', lang,
                                          ep_app=ep_app_number))
                    from sources.epo_ops_client import EPOClient
                    epo_client = EPOClient(
                        consumer_key=ep_ck, consumer_secret=ep_cs,
                    )
                    from sources.long_task.epo_examination import (
                        fetch_examination_data,
                        build_examination_timeline,
                    )
                    ep_data = await fetch_examination_data(ep_app_number, epo_client)
                    ep_exam_data = {
                        'ep_app_number': ep_app_number,
                        'biblio': ep_data.get('biblio'),
                        'events': ep_data.get('events', []),
                        'procedural_steps': ep_data.get('procedural_steps', []),
                        'claims_text': ep_data.get('claims_text', ''),
                        'search_report_text': ep_data.get('search_report_text', ''),
                        'timeline_events': ep_data.get('timeline_events', []),
                        'timeline_md': build_examination_timeline(
                            ep_data.get('timeline_events', []),
                            ep_data.get('biblio'),
                            lang,
                        ),
                        'status': ep_data.get('status', 'UNKNOWN'),
                    }
                    _pipeline_logger.info(
                        f"[task={task_id}] FAMILY PHASE0.45 ep_data — "
                        f"events={len(ep_data.get('events', []))}, "
                        f"steps={len(ep_data.get('procedural_steps', []))}, "
                        f"timeline={len(ep_data.get('timeline_events', []))}, "
                        f"has_search_report={ep_data.get('has_search_report')}, "
                        f"status={ep_data.get('status')}"
                    )
                    for _j in _jurisdictions:
                        if _j['code'] == 'EP':
                            _tl = len(ep_data.get('timeline_events', []))
                            if _tl > 0 or ep_data.get('has_biblio'):
                                _j['status'] = 'done'; _j['progress'] = 100
                                _j['detail'] = _t('family_ep_done', lang, steps=_tl,
                                                   status=ep_data.get('status', ''))
                            else:
                                _j['status'] = 'done'; _j['progress'] = 70
                                _j['detail'] = _t('family_ep_basic', lang)
                except Exception as e:
                    _pipeline_logger.warning(
                        f"[task={task_id}] FAMILY PHASE0.45 ep_fetch_failed — {e}"
                    )
                    for _j in _jurisdictions:
                        if _j['code'] == 'EP':
                            _j['status'] = 'no_data'
                            _j['detail'] = str(e)[:60]

        # ═════════════════════════════════════════════════════════════════
        # Phase 0.5: Fetch USPTO document list + classify
        # ═════════════════════════════════════════════════════════════════
        update_task_status(task_id, 'preparing', 8,
                           _t('fetching_uspto', lang))

        headers = {'Accept': 'application/json'}
        uspto_key = _os.getenv('USPTO_API_KEY', '')
        if uspto_key:
            headers['X-API-Key'] = uspto_key

        doc_list_url = (
            f"https://api.uspto.gov/api/v1/patent/applications/"
            f"{us_app_number}/documents"
        )
        _pipeline_logger.info(
            f"[task={task_id}] FAMILY PHASE0 fetch_doc_list — "
            f"app_number={us_app_number}, url={doc_list_url}"
        )
        resp = await _uspto_get_with_retry(doc_list_url, headers, timeout=20)
        if resp.status_code != 200:
            _pipeline_logger.error(
                f"[task={task_id}] FAMILY PHASE0 doc_list_failed — status={resp.status_code}"
            )
            set_task_failed(task_id, f"USPTO API returned HTTP {resp.status_code}")
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id,
                    'error': f'USPTO API status {resp.status_code}'}

        doc_list = resp.json() if resp.text else {}
        documents = (
            doc_list.get('documentBag', [])
            if isinstance(doc_list, dict)
            else []
        )
        if not documents:
            if lang == 'zh':
                msg = f"USPTO未返回专利申请 {us_app_number} 的任何文件。可能原因：申请号不存在、无权访问、或尚未公开。"
            else:
                msg = f"USPTO returned no documents for application {us_app_number}. The application may not exist, may not be accessible, or may not yet be published."
            set_task_failed(task_id, msg)
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': msg}

        manifest = classify_prosecution_documents(documents)
        docs_to_download = manifest.must_download.copy()
        if include_priority_2:
            docs_to_download.extend(manifest.recommended)

        _pipeline_logger.info(
            f"[task={task_id}] FAMILY PHASE0 classified — "
            f"total_in_bag={len(documents)}, "
            f"must_download={len(manifest.must_download)}, "
            f"recommended={len(manifest.recommended)}, "
            f"skipped={len(manifest.skipped)}, "
            f"to_download={len(docs_to_download)}"
        )

        if not docs_to_download:
            if lang == 'zh':
                msg = "未找到可分析的审查文件（无 Office Action、Response 或 Amendment）。可能该专利尚未进入实质审查阶段。"
            else:
                msg = "No analyzable prosecution documents found (no Office Actions, Responses, or Amendments). The patent may not have entered substantive examination."
            set_task_failed(task_id, msg)
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': msg}

        # ═════════════════════════════════════════════════════════════════
        # Phase 1: Generate table columns (Flash LLM)
        # ═════════════════════════════════════════════════════════════════
        from sources.long_task.prosecution_analyzer import (
            generate_table_columns,
            analyze_single_document,
            generate_document_summary,
            build_failed_row,
            generate_prosecution_report as gen_report,
        )

        update_task_status(task_id, 'generating_columns', 10,
                           _t('prosecution_framework', lang, total=len(docs_to_download)))
        _pipeline_logger.info(
            f"[task={task_id}] FAMILY PHASE1 generate_table_columns — "
            f"query={query[:100]}, doc_count={len(docs_to_download)}"
        )
        columns = await generate_table_columns(
            query=query, doc_count=len(docs_to_download),
            provider=flash_provider, lang=lang,
        )
        _pipeline_logger.info(
            f"[task={task_id}] FAMILY PHASE1 columns_generated — "
            f"column_count={len(columns)}, columns={columns}"
        )
        update_task_status(task_id, 'generating_columns', 12,
                           f'分析维度：{" | ".join(columns[1:4])}...',
                           table_columns=columns,
                           analysis_type='family')
        _update_mysql_progress(task_id, 'generating_columns', 12)

        # ═════════════════════════════════════════════════════════════════
        # Phase 2: Sequential download + pipelined parallel analysis
        # ═════════════════════════════════════════════════════════════════
        # Each document is downloaded sequentially.  As soon as download
        # completes, analysis (OCR -> LLM analyze -> LLM summarize) is launched
        # immediately via asyncio.create_task().  All analysis steps run under
        # a Semaphore(100), so multiple OCR + LLM calls execute concurrently.
        # (Same pattern as execute_prosecution_analysis Phase 2.)
        total_dl = len(docs_to_download)
        _pipeline_logger.info(
            f"[task={task_id}] FAMILY PHASE2 START — "
            f"pending_count={total_dl}, total={total_dl}"
        )

        # ── Activate US jurisdiction for Phase 2 document analysis ──
        for _j in _jurisdictions:
            if _j['code'] == 'US':
                _j['status'] = 'analyzing'
                _j['detail'] = _t('prosecution_downloading', lang, current=0, total=total_dl)
                _j['file_count'] = total_dl
                _j['files_done'] = 0

        _push_jurisdictions(15, 'downloading',
                            _t('prosecution_downloading', lang, current=0, total=total_dl))

        async def _fetch_prosecution(url: str, hdrs: dict, timeout: int):
            return await _uspto_get_with_retry(url, hdrs, timeout)

        _table_rows: list[dict] = [None] * total_dl  # preserve order by index
        _analyze_sem = asyncio.Semaphore(100)  # caps concurrent OCR + LLM calls
        _pending_tasks = []  # list of asyncio.Task
        _downloaded = 0  # documents downloaded so far
        _analyzed = 0    # documents with completed (or skipped) analysis
        from typing import Any as _AnyFam  # for type annotation in inner function

        async def _analyze_one(_doc: _AnyFam, _i: int) -> None:
            nonlocal _analyzed, _downloaded
            async with _analyze_sem:
                doc_index = _i + 1
                # Update US jurisdiction progress
                for _j in _jurisdictions:
                    if _j['code'] == 'US':
                        _j['detail'] = _t('prosecution_analyzing', lang, current=_j['files_done'], total=_j['file_count'], desc=_doc.description[:30])
                        _j['progress'] = int((_j['files_done'] / max(_j['file_count'], 1)) * 90) + 5

                # ── Text extraction (local-first -> Vision fallback) ──
                if not _doc.text and _doc.binary:
                    from sources.long_task.text_extractor import extract_text_from_binary
                    _local_text = extract_text_from_binary(
                        _doc.binary, skip_pdf_extraction=False,
                    )
                    if _local_text and len(_local_text.strip()) > 50:
                        _doc.text = _local_text.strip()
                        _pipeline_logger.info(
                            f"[task={task_id}] FAMILY PHASE2 local_extract_ok — "
                            f"code={_doc.document_code}, idx={doc_index}/{total_dl}, "
                            f"chars={len(_doc.text)}"
                        )
                    elif vision_enabled:
                        try:
                            _text = await _extract_text_via_vision(
                                _doc.binary, _doc.description, vision_provider,
                            )
                            if _text and len(_text.strip()) > 50:
                                _doc.text = _text.strip()
                                _pipeline_logger.info(
                                    f"[task={task_id}] FAMILY PHASE2 vision_ok — "
                                    f"code={_doc.document_code}, idx={doc_index}/{total_dl}, "
                                    f"chars={len(_doc.text)}"
                                )
                        except Exception as _e:
                            _pipeline_logger.warning(
                                f"[task={task_id}] FAMILY PHASE2 vision_error — "
                                f"code={_doc.document_code}: {type(_e).__name__}: {_e}"
                            )

                if not _doc.text or len(_doc.text.strip()) < 50:
                    row = build_failed_row(_doc.document_code, "text extraction failed", columns, lang)
                    row["_failed"] = True
                    row["_summary"] = ""
                    _table_rows[_i] = row
                    _analyzed += 1
                    for _j in _jurisdictions:
                        if _j['code'] == 'US':
                            _j['files_done'] += 1
                            _j['detail'] = _t('prosecution_analyzing', lang, current=_j['files_done'], total=_j['file_count'], desc=_doc.description[:30])
                            _j['progress'] = int((_j['files_done'] / max(_j['file_count'], 1)) * 90) + 5
                    _push_jurisdictions(
                        int((_analyzed + _downloaded) / max(total_dl, 1) * 70) + 15,
                        'analyzing',
                        _t('prosecution_analyzing', lang, current=_analyzed, total=total_dl))
                    return

                # ── Analyze (LLM) ──
                try:
                    row = await analyze_single_document(
                        doc_text=_doc.text,
                        doc_code=_doc.document_code,
                        doc_desc=_doc.description,
                        doc_category=_doc.category,
                        columns=columns,
                        query=query,
                        provider=pro_provider,
                        lang=lang,
                    )
                except Exception as e:
                    _pipeline_logger.warning(
                        f"[task={task_id}] FAMILY PHASE2 analyze_error — "
                        f"code={_doc.document_code}, idx={doc_index}/{total_dl}: "
                        f"{type(e).__name__}: {e}"
                    )
                    row = build_failed_row(_doc.document_code, str(e), columns, lang)
                    row["_failed"] = True

                # ── Summarize (LLM) ──
                try:
                    summary = await generate_document_summary(
                        doc_text=_doc.text, row=row, query=query,
                        provider=pro_provider, lang=lang,
                    )
                except Exception as e:
                    _pipeline_logger.warning(
                        f"[task={task_id}] FAMILY PHASE2 summary_error — "
                        f"code={_doc.document_code}: {type(e).__name__}: {e}"
                    )
                    summary = ""
                row["_summary"] = summary
                _table_rows[_i] = row
                _analyzed += 1
                # Update US jurisdiction: increment files_done, update progress
                for _j in _jurisdictions:
                    if _j['code'] == 'US':
                        _j['files_done'] += 1
                        _j['progress'] = int((_j['files_done'] / max(_j['file_count'], 1)) * 90) + 5
                        if _j['files_done'] >= _j['file_count']:
                            _j['status'] = 'done'
                            _j['progress'] = 100
                        _j['detail'] = _t('prosecution_analyzing', lang, current=_j['files_done'], total=_j['file_count'])

                _p = max(_downloaded, _analyzed)
                update_task_status(
                    task_id, 'analyzing',
                    progress_pct(_p, total_dl),
                    _t('prosecution_analyzing', lang,
                       current=_analyzed, total=total_dl,
                       desc=_doc.description[:40]),
                    table_rows=[r for r in _table_rows if r is not None],
                    table_columns=columns,
                    jurisdictions=_jurisdictions,
                )
                _pipeline_logger.info(
                    f"[task={task_id}] FAMILY PHASE2 analysis_done — "
                    f"code={_doc.document_code}, idx={doc_index}/{total_dl}, "
                    f"analyzed={_analyzed}/{total_dl}"
                )

        # ── Main pipeline: download sequentially, launch analysis immediately ──
        for _i, _doc in enumerate(docs_to_download):
            doc_index = _i + 1

            _result = _handle_task_stop(task_id, user_id, _downloaded, total_dl)
            if _result:
                return _result
            _result = _handle_task_pause(task_id, user_id, _downloaded, total_dl, {
                'completed': [r.get(columns[0], '') for r in _table_rows if r is not None and not r.get('_failed')],
                'pending': [d.document_code for d in docs_to_download[_i:]],
                'completed_rows': [r for r in _table_rows if r is not None],
                'failed': [r.get(columns[0], '') for r in _table_rows if r is not None and r.get('_failed')],
                'columns': columns,
                'downloaded': _downloaded,
                'analyzed': _analyzed,
            })
            if _result:
                return _result

            _visible_progress = max(_downloaded, _analyzed)
            # Update US jurisdiction detail during download phase
            for _j in _jurisdictions:
                if _j['code'] == 'US':
                    _j['detail'] = _t('prosecution_downloading', lang, current=doc_index, total=total_dl)
                    _j['progress'] = int((_downloaded / max(total_dl, 1)) * 15) + 5  # 5-20% during download
            update_task_status(
                task_id, 'downloading',
                progress_pct(_visible_progress, total_dl),
                _t('prosecution_downloading', lang, current=doc_index, total=total_dl),
                jurisdictions=_jurisdictions,
                table_columns=columns,
            )
            await download_single_document(_doc, _fetch_prosecution, us_app_number, headers)
            _downloaded += 1

            _pipeline_logger.info(
                f"[task={task_id}] FAMILY PHASE2 download_done — "
                f"code={_doc.document_code}, idx={doc_index}/{total_dl}, "
                f"text_len={len(_doc.text) if _doc.text else 0}"
            )

            _pending_tasks.append(asyncio.create_task(_analyze_one(_doc, _i)))

        # ── Wait for all in-flight analyses to finish ──
        docs_with_text = [d for d in docs_to_download if d.text and len(d.text.strip()) >= 50]
        _pipeline_logger.info(
            f"[task={task_id}] FAMILY PHASE2 all_downloads_done — "
            f"with_text={len(docs_with_text)}/{total_dl}, "
            f"waiting_for_{len(_pending_tasks)}_analyses"
        )
        if _pending_tasks:
            await asyncio.gather(*_pending_tasks)

        # Reconstruct table_rows in original order
        table_rows = [r for r in _table_rows if r is not None]

        if not table_rows:
            if lang == 'zh':
                msg = "所有审查文件处理失败。文件可能是扫描件或加密PDF。"
            else:
                msg = "All prosecution documents failed processing. Files may be scanned images or encrypted PDFs."
            set_task_failed(task_id, msg)
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': msg}

        # ═════════════════════════════════════════════════════════════════
        # Phase 3: Generate report
        # ═════════════════════════════════════════════════════════════════
        _pipeline_logger.info(
            f"[task={task_id}] FAMILY PHASE3 generate_report — "
            f"columns={columns}, table_rows_count={len(table_rows)}, "
            f"has_cn_data={bool(cn_exam_data and cn_exam_data.get('events') or cn_exam_data and cn_exam_data.get('timeline_md'))}"
        )
        update_task_status(task_id, 'generating_report', 80,
                           _t('writing_summary', lang))

        from sources.long_task.status_manager import ThrottledSummaryUpdater
        from sources.long_task.prosecution_analyzer import generate_family_prosecution_report as _gen_family_report
        summary_updater = ThrottledSummaryUpdater(
            task_id,
            progress=80,
            step_msg=_t('writing_summary', lang),
        )

        # Use the cross-jurisdiction report generator when CN data exists;
        # falls back to US-only report internally when no CN data.
        report_text = await _gen_family_report(
            table_rows=table_rows,
            columns=columns,
            query=query,
            patent_id=patent_id,
            flash_provider=flash_provider,
            pro_provider=pro_provider,
            lang=lang,
            cn_exam_data=cn_exam_data if cn_exam_data else None,
            jp_exam_data=jp_exam_data if jp_exam_data else None,
            ep_exam_data=ep_exam_data if ep_exam_data else None,
            family_overview=family_overview,
            summary_updater=summary_updater,
        )

        _pipeline_logger.info(
            f"[task={task_id}] FAMILY PHASE3 report_generated — "
            f"total_chars={len(report_text)}"
        )
        update_task_status(task_id, 'generating_report', 90,
                           _t('prosecution_report_complete', lang),
                           result_summary=report_text)
        _update_mysql_progress(task_id, 'generating_report', 90, result_summary=report_text)

        # ═════════════════════════════════════════════════════════════════
        # Phase 4: Export files (DOCX + PDF)
        # ═════════════════════════════════════════════════════════════════
        from sources.long_task.storage import get_storage_config

        storage_cfg = get_storage_config()
        _pipeline_logger.info(
            f"[task={task_id}] FAMILY PHASE4 export — "
            f"storage_backend={storage_cfg.get('report_storage_backend', 'local')}"
        )
        try:
            storage = create_storage(storage_cfg)
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] FAMILY PHASE4 storage_init FAILED — {e}, falling back to local"
            )
            storage = create_storage({'report_storage_backend': 'local'})

        report_files = []
        local_storage = None

        async def _get_local_storage():
            nonlocal local_storage
            if local_storage is None:
                from sources.long_task.storage import create_storage as _create
                local_storage = _create({'report_storage_backend': 'local'})
            return local_storage

        update_task_status(task_id, 'exporting', 90, _t('generating_word', lang))
        docx_bytes = await export_docx_async(report_text, table_rows, columns, lang=lang)
        try:
            await storage.put(task_id, 'report.docx', docx_bytes)
            _pipeline_logger.info(
                f"[task={task_id}] FAMILY PHASE4 docx — size_bytes={len(docx_bytes)}"
            )
            report_files.append({
                'format': 'docx', 'filename': 'report.docx', 'size': len(docx_bytes),
            })
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] FAMILY PHASE4 docx upload FAILED — {e}, falling back to local"
            )
            try:
                await (await _get_local_storage()).put(task_id, 'report.docx', docx_bytes)
                _pipeline_logger.info(
                    f"[task={task_id}] FAMILY PHASE4 docx — local fallback OK, size_bytes={len(docx_bytes)}"
                )
                report_files.append({
                    'format': 'docx', 'filename': 'report.docx', 'size': len(docx_bytes),
                })
            except Exception as e2:
                _pipeline_logger.error(
                    f"[task={task_id}] FAMILY PHASE4 docx local fallback ALSO FAILED — {e2}"
                )

        update_task_status(task_id, 'exporting', 95, _t('generating_pdf', lang))
        pdf_bytes = await export_pdf_async(docx_bytes)
        try:
            await storage.put(task_id, 'report.pdf', pdf_bytes)
            _pipeline_logger.info(
                f"[task={task_id}] FAMILY PHASE4 pdf — size_bytes={len(pdf_bytes)}"
            )
            report_files.append({
                'format': 'pdf', 'filename': 'report.pdf', 'size': len(pdf_bytes),
            })
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] FAMILY PHASE4 pdf upload FAILED — {e}, falling back to local"
            )
            try:
                await (await _get_local_storage()).put(task_id, 'report.pdf', pdf_bytes)
                _pipeline_logger.info(
                    f"[task={task_id}] FAMILY PHASE4 pdf — local fallback OK, size_bytes={len(pdf_bytes)}"
                )
                report_files.append({
                    'format': 'pdf', 'filename': 'report.pdf', 'size': len(pdf_bytes),
                })
            except Exception as e2:
                _pipeline_logger.error(
                    f"[task={task_id}] FAMILY PHASE4 pdf local fallback ALSO FAILED — {e2}"
                )

        # ── Complete ──
        _pipeline_logger.info(
            f"[task={task_id}] FAMILY COMPLETED — "
            f"patent_id={patent_id}, "
            f"us_pub={us_pub_number}, "
            f"docs_analyzed={len(table_rows)}, "
            f"report_chars={len(report_text)}, "
            f"report_files={[f['format'] for f in report_files]}"
        )
        _update_mysql_progress(task_id, 'exporting', 100)
        set_task_completed(task_id, report_files)

        if user_id:
            from sources.long_task.user_queue import complete_user_task
            try:
                next_task_id = complete_user_task(str(user_id), task_id)
                if next_task_id:
                    _dispatch_queued_task(next_task_id, user_id)
            except Exception as e:
                _pipeline_logger.warning(
                    f"[task={task_id}] FAMILY QUEUE_DISPATCH_FAILED — {e}"
                )

        return {'status': 'completed', 'task_id': task_id}

    loop = _asyncio.new_event_loop()
    try:
        return loop.run_until_complete(_run())
    except Exception as e:
        import traceback
        _pipeline_logger.error(
            f"[task={task_id}] FAMILY UNHANDLED_ERROR — "
            f"{type(e).__name__}: {e}\n{traceback.format_exc()}"
        )
        if user_id:
            try:
                from sources.long_task.user_queue import complete_user_task
                complete_user_task(str(user_id), task_id)
            except Exception:
                pass
        set_task_failed(task_id, f"{type(e).__name__}: {e}")
        _update_mysql_progress(task_id, 'failed', 0)
        return {'status': 'failed', 'task_id': task_id,
                'error': f'{type(e).__name__}: {e}'}
    finally:
        loop.close()


# ── Family helpers ─────────────────────────────────────────────────────────────


def _build_family_overview_status(family, lang: str) -> dict:
    """Build a family_overview dict for the frontend progress card."""
    members_data = []
    for m in family.deduplicated_members:
        members_data.append({
            'country': m.country,
            'pub_number': m.pub_number,
            'kind': m.pub_kind,
            'is_granted': m.is_granted,
            'title': m.title,
        })

    # Merge EP member states (DE, GB, FR, AT, …) into a single "EP" entry
    # for the frontend jurisdiction badges — same logic as the progress table.
    _EP_MEMBER_CODES = frozenset({
        'AT', 'BE', 'BG', 'CH', 'CY', 'CZ', 'DE', 'DK', 'EE', 'ES',
        'FI', 'FR', 'GB', 'GR', 'HR', 'HU', 'IE', 'IS', 'IT', 'LI',
        'LT', 'LU', 'LV', 'MC', 'MK', 'MT', 'NL', 'NO', 'PL', 'PT',
        'RO', 'SE', 'SI', 'SK', 'SM', 'TR',
    })
    display_jurisdictions: list[str] = []
    _seen_ep = False
    for _c in family.jurisdictions:
        if _c in _EP_MEMBER_CODES:
            if not _seen_ep:
                display_jurisdictions.append('EP')
                _seen_ep = True
        else:
            display_jurisdictions.append(_c)

    if lang == 'zh':
        step_msg = (
            f"查到 {family.query_pub_number} 的同族专利，"
            f"共 {len(family.deduplicated_members)} 个成员，"
            f"涉及 {len(display_jurisdictions)} 个司法辖区: "
            f"{', '.join(display_jurisdictions)}"
        )
    else:
        step_msg = (
            f"Found {len(family.deduplicated_members)} family members for "
            f"{family.query_pub_number} across {len(display_jurisdictions)} "
            f"jurisdictions: {', '.join(display_jurisdictions)}"
        )

    return {
        'family_id': family.family_id,
        'total_count': family.total_count,
        'jurisdictions': display_jurisdictions,
        'members': members_data,
        'step_msg': step_msg,
    }


# ── China Examination helpers ──────────────────────────────────────────────────


def _build_markdown_table(
    table_rows: list[dict], columns: list[str], lang: str = "zh",
) -> str:
    """Build a markdown table from analysis rows and columns."""
    if not table_rows or not columns:
        return ""

    lines: list[str] = []
    # Header
    header = "| " + " | ".join(columns) + " |"
    separator = "|" + "|".join("------" for _ in columns) + "|"
    lines.append(header)
    lines.append(separator)

    for row in table_rows:
        cells = []
        for col in columns:
            val = row.get(col, "")
            val = str(val).replace("|", "\\|").replace("\n", " ")
            if len(val) > 120:
                val = val[:117] + "..."
            cells.append(val)
        lines.append("| " + " | ".join(cells) + " |")

    if lang == "zh":
        lines.insert(0, "## 审查决定分析表\n")
    else:
        lines.insert(0, "## Examination Decision Analysis Table\n")

    return "\n".join(lines)


def _build_events_summary_text(
    events: list, lang: str = "zh",
) -> str:
    """Build a condensed text summary of all examination events for the LLM."""
    from sources.long_task.china_examination import (
        ExaminationEvent, _format_event_for_llm,
    )
    parts: list[str] = []
    for i, evt in enumerate(events):
        if not isinstance(evt, ExaminationEvent):
            continue
        label = evt.decision_label_zh if lang == "zh" else evt.decision_label_en
        parts.append(
            f"### {i+1}. {label} — {evt.decision_number} ({evt.decision_date})\n\n"
            f"{_format_event_for_llm(evt)[:4000]}"
        )
    return "\n\n".join(parts)


def _build_basic_info_timeline(
    basic_info: dict, law_state: dict, lang: str = "zh",
) -> str:
    """Build a basic info + law-state summary markdown block."""
    lines: list[str] = []
    if lang == "zh":
        lines.append("## 专利基本信息\n")
    else:
        lines.append("## Patent Basic Information\n")

    title = basic_info.get("title", "")
    app_num = basic_info.get("applicationDocNum", "")
    app_date = basic_info.get("applicationDate", "")
    pub_num = basic_info.get("publicationDocNum", "")
    pub_date = basic_info.get("publicationDate", "")
    applicants = basic_info.get("applicant", [])
    applicant_str = ", ".join(applicants) if isinstance(applicants, list) else str(applicants or "")
    inventors = basic_info.get("inventor", [])
    inventor_str = ", ".join(inventors) if isinstance(inventors, list) else str(inventors or "")
    ipc_main = basic_info.get("ipcMain", "")
    abstract = basic_info.get("abstractDesc", "")

    if lang == "zh":
        if title:
            lines.append(f"**专利名称**：{title}")
        if app_num:
            lines.append(f"**申请号**：{app_num}")
        if app_date:
            lines.append(f"**申请日**：{app_date}")
        if pub_num:
            lines.append(f"**公开号**：{pub_num}")
        if pub_date:
            lines.append(f"**公开日**：{pub_date}")
        if applicant_str:
            lines.append(f"**申请人**：{applicant_str}")
        if inventor_str:
            lines.append(f"**发明人**：{inventor_str}")
        if ipc_main:
            lines.append(f"**主分类号**：{ipc_main}")
        if abstract:
            lines.append(f"\n**摘要**：{abstract}")

        law_status = law_state.get("lawStatus", "")
        law_date = law_state.get("date", "")
        if law_status or law_date:
            lines.append(f"\n**当前法律状态**：{law_status}（{law_date}）")
    else:
        if title:
            lines.append(f"**Title**: {title}")
        if app_num:
            lines.append(f"**Application Number**: {app_num}")
        if app_date:
            lines.append(f"**Filing Date**: {app_date}")
        if pub_num:
            lines.append(f"**Publication Number**: {pub_num}")
        if pub_date:
            lines.append(f"**Publication Date**: {pub_date}")
        if applicant_str:
            lines.append(f"**Applicant**: {applicant_str}")
        if inventor_str:
            lines.append(f"**Inventor**: {inventor_str}")
        if ipc_main:
            lines.append(f"**IPC Main**: {ipc_main}")
        if abstract:
            lines.append(f"\n**Abstract**: {abstract}")

        law_status = law_state.get("lawStatus", "")
        law_date = law_state.get("date", "")
        if law_status or law_date:
            lines.append(f"\n**Current Legal Status**: {law_status} ({law_date})")

    return "\n".join(lines) + "\n"


# ═══════════════════════════════════════════════════════════════════════════════
# China Patent Examination History Analysis Task (single patent)
# ═══════════════════════════════════════════════════════════════════════════════


@app.task(bind=True, max_retries=2, default_retry_delay=60, time_limit=1800, soft_time_limit=1770)
def execute_china_examination_analysis(self, task_id: str, params: dict):
    """Analyze Chinese patent examination history via sipop.cn open data platform.

    Phase 0: US/EP patent → CN application number via EPO family
    Phase 1: Fetch examination data via SIPOP API (queryPatentReview + lawState + basicInfo)
    Phase 2: AI analysis of examination events (Flash columns → Pro per-event analysis)
    Phase 3: Generate report (DOCX + PDF)

    Unlike USPTO prosecution analysis which downloads documents, this task works
    with *structured* review decision data returned by the SIPOP API.
    """
    import asyncio as _asyncio
    import os as _os

    retry_count = self.request.retries
    patent_id = str(params.get('patent_id', '')).strip()
    query = params.get('query', '')
    lang = params.get('lang', 'zh')
    session_id = params.get('session_id', '')
    user_id = params.get('user_id', '')

    _pipeline_logger.info(
        f"[task={task_id}] CHINA_EXAM START — "
        f"patent_id={patent_id}, query={query[:120]}, lang={lang}, "
        f"retry={retry_count}/{self.max_retries}"
    )

    if retry_count >= self.max_retries:
        _pipeline_logger.error(
            f"[task={task_id}] CHINA_EXAM HARD_STOP — retry_count={retry_count}"
        )
        if user_id:
            from sources.long_task.user_queue import complete_user_task
            try:
                complete_user_task(str(user_id), task_id)
            except Exception:
                pass
        return {'status': 'failed', 'task_id': task_id,
                'error': f'Max retries ({self.max_retries}) exceeded'}

    if not patent_id:
        _pipeline_logger.error(f"[task={task_id}] CHINA_EXAM no patent_id provided")
        return {'status': 'failed', 'task_id': task_id, 'error': 'No patent_id'}

    # ── Imports ──────────────────────────────────────────────────────────────
    from sources.long_task.status_manager import (
        update_task_status, set_task_completed, set_task_failed,
    )
    from sources.long_task.config import (
        get_long_task_config, get_sipop_config, get_family_config,
    )
    from sources.long_task.patent_family import EPOFamilyClient, EPOError
    from sources.long_task.china_examination import (
        resolve_cn_application_number,
        fetch_examination_data,
        generate_table_columns,
        analyze_single_event,
        generate_event_summary,
        build_examination_timeline,
        format_claims_for_report,
        build_legal_status_timeline,
    )
    from sources.sipop_client import SipopClient, SipopError
    from sources.long_task.storage import create_storage
    from sources.llm_provider import Provider

    # ── Provider setup ───────────────────────────────────────────────────────
    ltc = get_long_task_config()
    model_family = ltc['provider_family']

    if model_family == 'minimax':
        flash_provider = Provider(
            provider_name='minimax', model='MiniMax-M2.7-highspeed',
            server_address='', is_local=False,
        )
        pro_provider = Provider(
            provider_name='minimax', model='MiniMax-M3',
            server_address='', is_local=False,
        )
    else:
        flash_provider = Provider(
            provider_name='deepseek', model='deepseek-v4-flash',
            server_address='', is_local=False,
        )
        pro_provider = Provider(
            provider_name='deepseek', model='deepseek-v4-pro',
            server_address='', is_local=False,
        )

    # ── SIPOP config ─────────────────────────────────────────────────────────
    sipop_cfg = get_sipop_config()
    sipop_app_key = sipop_cfg.get('app_key', '')
    sipop_app_secret = sipop_cfg.get('app_secret', '')

    if not sipop_app_key or not sipop_app_secret:
        _pipeline_logger.error(
            f"[task={task_id}] CHINA_EXAM missing SIPOP credentials"
        )
        set_task_failed(task_id, "SIPOP API credentials not configured")
        _update_mysql_progress(task_id, 'failed', 0)
        return {'status': 'failed', 'task_id': task_id,
                'error': 'SIPOP credentials not configured'}

    sipop_client = SipopClient(app_key=sipop_app_key, app_secret=sipop_app_secret)

    # ── Family config (for EPO family lookup) ────────────────────────────────
    fc = get_family_config()
    epo_key = fc.get('epo_consumer_key', '')
    epo_secret = fc.get('epo_consumer_secret', '')
    epo_client = EPOFamilyClient(
        consumer_key=epo_key, consumer_secret=epo_secret,
    ) if epo_key and epo_secret else None

    # ── Run pipeline ─────────────────────────────────────────────────────────
    async def _run():
        # ═════════════════════════════════════════════════════════════════
        # Phase 0: Resolve CN application number
        # ═════════════════════════════════════════════════════════════════
        update_task_status(task_id, 'preparing', 1,
                           _t('china_resolving', lang),
                           analysis_type='prosecution')

        family_context: dict = {}
        try:
            cn_app_number, family_context = await resolve_cn_application_number(
                patent_id, epo_client,
            )
        except ValueError as e:
            _pipeline_logger.error(
                f"[task={task_id}] CHINA_EXAM PHASE0 resolve_failed — {e}"
            )
            set_task_failed(task_id, str(e))
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': str(e)}

        _pipeline_logger.info(
            f"[task={task_id}] CHINA_EXAM PHASE0 resolved — "
            f"input={patent_id}, cn_app={cn_app_number}, "
            f"direct_cn={family_context.get('direct_cn', False)}, "
            f"jurisdictions={family_context.get('jurisdictions', [])}"
        )

        # If we used EPO family, build an overview for the frontend
        if not family_context.get('direct_cn'):
            jurisdictions = family_context.get('jurisdictions', [])
            update_task_status(
                task_id, 'preparing', 5,
                _t('china_family_overview', lang,
                   input_id=patent_id, cn_app=cn_app_number,
                   n_juris=len(jurisdictions),
                   juris=', '.join(jurisdictions)),
                family_overview={
                    'input_id': patent_id,
                    'cn_app_number': cn_app_number,
                    'jurisdictions': jurisdictions,
                    'members': family_context.get('members', []),
                },
            )

        # ═════════════════════════════════════════════════════════════════
        # Phase 1: Fetch examination data from SIPOP API
        # ═════════════════════════════════════════════════════════════════
        update_task_status(task_id, 'preparing', 10,
                           _t('china_fetching_review', lang))

        try:
            events, law_state, basic_info, claims, legal_timeline = await fetch_examination_data(
                cn_app_number, sipop_client,
            )
        except SipopError as e:
            _pipeline_logger.error(
                f"[task={task_id}] CHINA_EXAM PHASE1 fetch_failed — {e}"
            )
            set_task_failed(task_id, f"SIPOP API error: {e}")
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': str(e)}

        total_events = len(events)
        _pipeline_logger.info(
            f"[task={task_id}] CHINA_EXAM PHASE1 data_fetched — "
            f"events={total_events}, claims={len(claims)}, "
            f"legal_events={len(legal_timeline)}, "
            f"has_law_state={bool(law_state)}, "
            f"has_basic_info={bool(basic_info)}"
        )

        if not events:
            # No examination review decisions — but we may still have claims
            # and legal status data to build a useful report.
            title = basic_info.get("title", "") or patent_id
            report_title = _t('china_report_title', lang, patent_id=title)
            claims_md = format_claims_for_report(claims, lang) if claims else ""
            legal_md = build_legal_status_timeline(legal_timeline, lang) if legal_timeline else ""

            # Build timeline from basic info
            basic_timeline = _build_basic_info_timeline(basic_info, law_state, lang)

            if not claims_md and not legal_md:
                # Truly nothing — fail gracefully
                if lang == 'zh':
                    msg = (
                        f"未找到中国专利申请 {cn_app_number} 的审查决定、权利要求"
                        f"或法律状态数据。可能该专利数据尚未录入平台。"
                    )
                else:
                    msg = (
                        f"No examination review, claims, or legal status data "
                        f"found for CN application {cn_app_number}."
                    )
                set_task_failed(task_id, msg)
                _update_mysql_progress(task_id, 'failed', 0)
                return {'status': 'failed', 'task_id': task_id, 'error': msg}

            if lang == 'zh':
                exec_text = (
                    f"中国专利申请 {cn_app_number} 暂无复审、无效或异议审查决定记录。"
                    f"以下为专利的基本信息、权利要求及法律状态时间线。"
                )
            else:
                exec_text = (
                    f"No reexamination, invalidation, or opposition decisions "
                    f"found for CN application {cn_app_number}. "
                    f"Below are the patent's basic information, claims, and "
                    f"legal status timeline."
                )

            exec_heading = _t('default_exec_summary_heading', lang)
            report_text = (
                f"# {report_title}\n\n"
                f"## {exec_heading}\n\n{exec_text}\n\n"
                f"{basic_timeline}\n\n"
                f"{claims_md}\n\n"
                f"{legal_md}\n"
            )

            # Export
            table_rows = []
            columns = ["决定号"]
            _update_mysql_progress(task_id, 'generating_report', 90, result_summary=report_text)

            from sources.long_task.storage import get_storage_config
            storage_cfg = get_storage_config()
            try:
                storage = create_storage(storage_cfg)
            except Exception:
                storage = create_storage({'report_storage_backend': 'local'})

            report_files = []
            docx_bytes = await export_docx_async(report_text, table_rows, columns, lang=lang)
            try:
                await storage.put(task_id, 'report.docx', docx_bytes)
                report_files.append({'format': 'docx', 'filename': 'report.docx', 'size': len(docx_bytes)})
            except Exception:
                pass

            pdf_bytes = await export_pdf_async(docx_bytes)
            try:
                await storage.put(task_id, 'report.pdf', pdf_bytes)
                report_files.append({'format': 'pdf', 'filename': 'report.pdf', 'size': len(pdf_bytes)})
            except Exception:
                pass

            set_task_completed(task_id, report_files)
            if user_id:
                from sources.long_task.user_queue import complete_user_task
                try:
                    next_task_id = complete_user_task(str(user_id), task_id)
                    if next_task_id:
                        _dispatch_queued_task(next_task_id, user_id)
                except Exception:
                    pass
            return {'status': 'completed', 'task_id': task_id}

        # ═════════════════════════════════════════════════════════════════
        # Phase 2: Generate table columns (Flash LLM)
        # ═════════════════════════════════════════════════════════════════
        update_task_status(
            task_id, 'generating_columns', 15,
            _t('china_analyzing_decisions', lang, current=0, total=total_events),
        )

        _pipeline_logger.info(
            f"[task={task_id}] CHINA_EXAM PHASE2 generate_table_columns"
        )
        columns = await generate_table_columns(
            query=query, event_count=total_events,
            provider=flash_provider, lang=lang,
        )
        _pipeline_logger.info(
            f"[task={task_id}] CHINA_EXAM PHASE2 columns_generated — "
            f"columns={columns}"
        )
        _update_mysql_progress(task_id, 'generating_columns', 20)

        # ═════════════════════════════════════════════════════════════════
        # Phase 2b: Per-event analysis (Pro LLM)
        # ═════════════════════════════════════════════════════════════════
        table_rows: list[dict] = []
        for i, event in enumerate(events):
            idx = i + 1
            pct = 20 + int((i / total_events) * 40)
            update_task_status(
                task_id, 'analyzing', pct,
                _t('china_analyzing_decisions', lang, current=idx, total=total_events),
                table_rows=table_rows,
            )

            # Pause / Stop handling
            _uid = params.get('user_id', '')
            _completed = len(table_rows)
            _result = _handle_task_stop(task_id, _uid, _completed, total_events)
            if _result:
                return _result
            _result = _handle_task_pause(task_id, _uid, _completed, total_events, {
                'completed': [r.get('决定号', '') for r in table_rows],
                'pending': [e.decision_number for e in events[i:]],
                'completed_rows': table_rows,
                'columns': columns,
            })
            if _result:
                return _result

            _pipeline_logger.info(
                f"[task={task_id}] CHINA_EXAM PHASE2 event[{idx}/{total_events}] — "
                f"decision={event.decision_number}, type={event.decision}"
            )

            try:
                row = await analyze_single_event(
                    event=event, columns=columns, query=query,
                    provider=pro_provider, lang=lang,
                )
            except Exception as e:
                _pipeline_logger.warning(
                    f"[task={task_id}] CHINA_EXAM PHASE2 event[{idx}] analyze_error — {e}"
                )
                row = {"决定号": event.decision_number, "_failed": True,
                       "_failure_reason": str(e)}
                for col in columns:
                    if col not in row:
                        row[col] = "分析失败" if lang == "zh" else "Analysis failed"

            try:
                summary = await generate_event_summary(
                    event=event, row=row, query=query,
                    provider=pro_provider, lang=lang,
                )
            except Exception as e:
                _pipeline_logger.warning(
                    f"[task={task_id}] CHINA_EXAM PHASE2 event[{idx}] summary_error — {e}"
                )
                summary = ""
            row["_summary"] = summary
            table_rows.append(row)

        _pipeline_logger.info(
            f"[task={task_id}] CHINA_EXAM PHASE2 complete — "
            f"table_rows={len(table_rows)}, "
            f"failed={sum(1 for r in table_rows if r.get('_failed'))}"
        )

        # ═════════════════════════════════════════════════════════════════
        # Phase 3: Generate report
        # ═════════════════════════════════════════════════════════════════
        update_task_status(task_id, 'generating_report', 65,
                           _t('china_generating_timeline', lang))

        # Build timeline
        timeline_md = await build_examination_timeline(
            events=events, law_state=law_state,
            basic_info=basic_info, lang=lang,
        )

        # Executive summary via AI
        update_task_status(task_id, 'generating_report', 70,
                           _t('writing_summary', lang))

        from sources.long_task.status_manager import ThrottledSummaryUpdater
        summary_updater = ThrottledSummaryUpdater(
            task_id, progress=70, step_msg=_t('writing_summary', lang),
        )

        # Assemble report: Timeline + Analysis Table + AI-generated sections
        title = basic_info.get("title", "") or patent_id
        report_title = _t('china_report_title', lang, patent_id=title)

        # Build the analysis table
        table_md = _build_markdown_table(table_rows, columns, lang)

        # Build a text summary of all events for the AI to generate report
        events_summary_text = _build_events_summary_text(events, lang)

        if lang == "zh":
            exec_prompt = (
                "你是一个中国专利审查历史分析专家。基于以下审查决定的综合分析数据，"
                "撰写一份\"核心审查洞察\"摘要（300-500字）。\n\n"
                "要求：\n"
                "- 总结该专利在中国经历的关键审查程序\n"
                "- 提炼出最关键的审查决定及其影响\n"
                "- 如果涉及多个程序（如先复审再无效），分析程序之间的关系\n"
                "- 对专利最终的法律状态给出判断\n"
                "- 用 bullet points 组织，每点 1-2 句话\n"
            )
        else:
            exec_prompt = (
                "You are a China patent examination history analysis expert. "
                "Based on the comprehensive analysis of the following examination "
                "decisions, write a \"Key Examination Insights\" summary (300-500 words).\n\n"
                "Requirements:\n"
                "- Summarize the key examination procedures this patent underwent in China\n"
                "- Extract the most critical decisions and their impact\n"
                "- If multiple procedures are involved (e.g. reexamination then invalidation), "
                "analyze the relationships between them\n"
                "- Provide a judgment on the patent's final legal status\n"
                "- Use bullet points, 1-2 sentences each\n"
                "Write ALL content in English."
            )

        exec_text = ""
        try:
            exec_result = await pro_provider.complete_json(
                exec_prompt,
                f"专利：{title}（申请号：{cn_app_number}）\n\n"
                f"{timeline_md}\n\n{events_summary_text}",
            )
            if isinstance(exec_result, dict):
                exec_text = exec_result.get("summary", "") or str(exec_result)
            else:
                exec_text = str(exec_result) if exec_result else ""
        except Exception as e:
            _pipeline_logger.warning(
                f"[task={task_id}] CHINA_EXAM PHASE3 exec_summary failed — {e}"
            )
            if lang == "zh":
                exec_text = f"本报告分析了中国专利申请 {cn_app_number} 的 {total_events} 个审查决定。详细分析见下方审查决定分析表。"
            else:
                exec_text = f"This report analyzes {total_events} examination decisions for CN application {cn_app_number}. See detailed analysis in the table below."

        exec_heading = _t('default_exec_summary_heading', lang)

        # Build claims section
        claims_md = format_claims_for_report(claims, lang) if claims else ""

        # Build legal status timeline
        legal_md = build_legal_status_timeline(legal_timeline, lang) if legal_timeline else ""

        report_text = (
            f"# {report_title}\n\n"
            f"## {exec_heading}\n\n{exec_text}\n\n"
            f"{timeline_md}\n\n"
            f"{table_md}\n\n"
        )

        # If there are events with substantial content, add per-decision analysis
        for i, (event, row) in enumerate(zip(events, table_rows)):
            summary = row.get("_summary", "")
            if summary:
                if lang == "zh":
                    heading = f"决定 {event.decision_number} — {event.decision_label_zh}"
                else:
                    heading = f"Decision {event.decision_number} — {event.decision_label_en}"
                report_text += f"### {heading}\n\n{summary}\n\n"

        # Append claims and legal status at the end
        if claims_md:
            report_text += f"\n{claims_md}\n"
        if legal_md:
            report_text += f"\n{legal_md}\n"

        _pipeline_logger.info(
            f"[task={task_id}] CHINA_EXAM PHASE3 report — "
            f"total_chars={len(report_text)}"
        )
        update_task_status(task_id, 'generating_report', 90,
                           _t('report_writing_complete', lang),
                           result_summary=report_text)
        _update_mysql_progress(task_id, 'generating_report', 90, result_summary=report_text)

        # ═════════════════════════════════════════════════════════════════
        # Phase 4: Export DOCX + PDF
        # ═════════════════════════════════════════════════════════════════
        from sources.long_task.storage import get_storage_config

        storage_cfg = get_storage_config()
        _pipeline_logger.info(
            f"[task={task_id}] CHINA_EXAM PHASE4 export — "
            f"backend={storage_cfg.get('report_storage_backend', 'local')}"
        )
        try:
            storage = create_storage(storage_cfg)
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] CHINA_EXAM PHASE4 storage_init FAILED — {e}"
            )
            storage = create_storage({'report_storage_backend': 'local'})

        report_files = []
        local_storage = None

        async def _get_local_storage():
            nonlocal local_storage
            if local_storage is None:
                from sources.long_task.storage import create_storage as _create
                local_storage = _create({'report_storage_backend': 'local'})
            return local_storage

        # DOCX
        update_task_status(task_id, 'exporting', 90, _t('generating_word', lang))
        docx_bytes = await export_docx_async(report_text, table_rows, columns, lang=lang)
        try:
            await storage.put(task_id, 'report.docx', docx_bytes)
            _pipeline_logger.info(
                f"[task={task_id}] CHINA_EXAM PHASE4 docx — size={len(docx_bytes)}"
            )
            report_files.append({'format': 'docx', 'filename': 'report.docx', 'size': len(docx_bytes)})
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] CHINA_EXAM PHASE4 docx upload FAILED — {e}"
            )
            try:
                await (await _get_local_storage()).put(task_id, 'report.docx', docx_bytes)
                report_files.append({'format': 'docx', 'filename': 'report.docx', 'size': len(docx_bytes)})
            except Exception as e2:
                _pipeline_logger.error(
                    f"[task={task_id}] CHINA_EXAM PHASE4 docx local fallback ALSO FAILED — {e2}"
                )

        # PDF
        update_task_status(task_id, 'exporting', 95, _t('generating_pdf', lang))
        pdf_bytes = await export_pdf_async(docx_bytes)
        try:
            await storage.put(task_id, 'report.pdf', pdf_bytes)
            _pipeline_logger.info(
                f"[task={task_id}] CHINA_EXAM PHASE4 pdf — size={len(pdf_bytes)}"
            )
            report_files.append({'format': 'pdf', 'filename': 'report.pdf', 'size': len(pdf_bytes)})
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] CHINA_EXAM PHASE4 pdf upload FAILED — {e}"
            )
            try:
                await (await _get_local_storage()).put(task_id, 'report.pdf', pdf_bytes)
                report_files.append({'format': 'pdf', 'filename': 'report.pdf', 'size': len(pdf_bytes)})
            except Exception as e2:
                _pipeline_logger.error(
                    f"[task={task_id}] CHINA_EXAM PHASE4 pdf local fallback ALSO FAILED — {e2}"
                )

        # ── Complete ──
        _pipeline_logger.info(
            f"[task={task_id}] CHINA_EXAM COMPLETED — "
            f"patent_id={patent_id}, cn_app={cn_app_number}, "
            f"decisions_analyzed={len(table_rows)}, "
            f"report_chars={len(report_text)}, "
            f"files={[f['format'] for f in report_files]}"
        )
        _update_mysql_progress(task_id, 'exporting', 100)
        set_task_completed(task_id, report_files)

        if user_id:
            from sources.long_task.user_queue import complete_user_task
            try:
                next_task_id = complete_user_task(str(user_id), task_id)
                if next_task_id:
                    _dispatch_queued_task(next_task_id, user_id)
            except Exception as e:
                _pipeline_logger.warning(
                    f"[task={task_id}] CHINA_EXAM QUEUE_DISPATCH_FAILED — {e}"
                )

        return {'status': 'completed', 'task_id': task_id}

    loop = _asyncio.new_event_loop()
    try:
        return loop.run_until_complete(_run())
    except Exception as e:
        import traceback
        _pipeline_logger.error(
            f"[task={task_id}] CHINA_EXAM UNHANDLED_ERROR — "
            f"{type(e).__name__}: {e}\n{traceback.format_exc()}"
        )
        if user_id:
            try:
                from sources.long_task.user_queue import complete_user_task
                complete_user_task(str(user_id), task_id)
            except Exception:
                pass
        set_task_failed(task_id, f"{type(e).__name__}: {e}")
        _update_mysql_progress(task_id, 'failed', 0)
        return {'status': 'failed', 'task_id': task_id,
                'error': f'{type(e).__name__}: {e}'}
    finally:
        loop.close()


# ── EPO Examination History Analysis Task (single patent)
# ═══════════════════════════════════════════════════════════════════════════════


@app.task(bind=True, max_retries=2, default_retry_delay=60, time_limit=1800, soft_time_limit=1770)
def execute_epo_examination_analysis(self, task_id: str, params: dict):
    """Analyze European patent examination history via EPO OPS Register API.

    Phase 0: Patent ID → EP application number via EPO family
    Phase 1: Fetch Register data + Published Data (search report text)
    Phase 2: AI analysis of search opinion + procedural timeline
    Phase 3: Generate report (DOCX + PDF)

    Unlike USPTO which downloads full documents, EPO analysis works with
    a mix of full-text (search opinion) and structured metadata (procedural
    steps, legal events).
    """
    import asyncio as _asyncio
    import os as _os

    retry_count = self.request.retries
    patent_id = str(params.get('patent_id', '')).strip()
    query = params.get('query', '')
    lang = params.get('lang', 'zh')
    session_id = params.get('session_id', '')
    user_id = params.get('user_id', '')

    _pipeline_logger.info(
        f"[task={task_id}] EPO_EXAM START — "
        f"patent_id={patent_id}, query={query[:120]}, lang={lang}, "
        f"retry={retry_count}/{self.max_retries}"
    )

    if retry_count >= self.max_retries:
        _pipeline_logger.error(
            f"[task={task_id}] EPO_EXAM HARD_STOP — retry_count={retry_count} >= {self.max_retries}"
        )
        if user_id:
            from sources.long_task.user_queue import complete_user_task
            try:
                complete_user_task(str(user_id), task_id)
            except Exception:
                pass
        return {'status': 'failed', 'task_id': task_id,
                'error': f'Max retries ({self.max_retries}) exceeded'}

    if not patent_id:
        _pipeline_logger.error(f"[task={task_id}] EPO_EXAM no patent_id provided")
        _update_mysql_progress(task_id, 'failed', 0)
        return {'status': 'failed', 'task_id': task_id, 'error': 'No patent_id'}

    from sources.long_task.status_manager import (
        update_task_status, set_task_completed, set_task_failed,
    )
    from sources.long_task.config import (
        get_long_task_config, get_family_config,
    )
    from sources.long_task.storage import create_storage
    from sources.llm_provider import Provider

    # ── Provider setup ──
    ltc = get_long_task_config()
    model_family = ltc['provider_family']

    if model_family == 'minimax':
        flash_provider = Provider(
            provider_name='minimax', model='MiniMax-M2.7-highspeed',
            server_address='', is_local=False,
        )
        pro_provider = Provider(
            provider_name='minimax', model='MiniMax-M3',
            server_address='', is_local=False,
        )
    else:
        flash_provider = Provider(
            provider_name='deepseek', model='deepseek-v4-flash',
            server_address='', is_local=False,
        )
        pro_provider = Provider(
            provider_name='deepseek', model='deepseek-v4-pro',
            server_address='', is_local=False,
        )

    async def _run():
        # ── Phase 0: Resolve EP application number ──────────────────
        update_task_status(task_id, 'preparing', 5,
                           _t('preparing', lang, total=1),
                           analysis_type='prosecution')

        ep_family_cfg = get_family_config()
        ep_ck = ep_family_cfg.get('epo_consumer_key', '')
        ep_cs = ep_family_cfg.get('epo_consumer_secret', '')
        # Env var overrides
        ep_ck = _os.getenv('EPO_CONSUMER_KEY', ep_ck)
        ep_cs = _os.getenv('EPO_CONSUMER_SECRET', ep_cs)
        if not ep_ck or not ep_cs:
            msg = "EPO API credentials not configured (set EPO_CONSUMER_KEY/EPO_CONSUMER_SECRET)."
            set_task_failed(task_id, msg)
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': msg}

        from sources.epo_ops_client import EPOClient
        epo_client = EPOClient(consumer_key=ep_ck, consumer_secret=ep_cs)

        from sources.long_task.epo_examination import (
            resolve_ep_application_number,
            fetch_examination_data,
            build_examination_timeline,
            generate_table_columns,
            analyze_single_timeline_event,
            generate_event_summary,
        )

        try:
            ep_app_number, family_ctx = await resolve_ep_application_number(
                patent_id, epo_client,
            )
        except ValueError as e:
            _pipeline_logger.error(
                f"[task={task_id}] EPO_EXAM PHASE0 resolve_failed — {e}"
            )
            set_task_failed(task_id, str(e))
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': str(e)}

        _pipeline_logger.info(
            f"[task={task_id}] EPO_EXAM PHASE0 resolved — "
            f"input={patent_id}, ep_app={ep_app_number}"
        )

        # ── Phase 1: Fetch examination data ──────────────────────────
        update_task_status(task_id, 'preparing', 10,
                           f"Fetching EPO Register data for EP{ep_app_number}...")
        ep_data = await fetch_examination_data(ep_app_number, epo_client)

        if not ep_data.get('has_events') and not ep_data.get('has_steps'):
            _pipeline_logger.warning(
                f"[task={task_id}] EPO_EXAM PHASE1 no_data — app={ep_app_number}"
            )
            set_task_failed(task_id, _t('epo_no_review_data', lang))
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id,
                    'error': _t('epo_no_review_data', lang)}

        _pipeline_logger.info(
            f"[task={task_id}] EPO_EXAM PHASE1 data_fetched — "
            f"events={len(ep_data.get('events', []))}, "
            f"steps={len(ep_data.get('procedural_steps', []))}, "
            f"timeline={len(ep_data.get('timeline_events', []))}, "
            f"has_search_report={ep_data.get('has_search_report')}, "
            f"status={ep_data.get('status')}"
        )

        # ── Phase 2: AI analysis ─────────────────────────────────────
        timeline_events = ep_data.get('timeline_events', [])
        search_report_text = ep_data.get('search_report_text', '')
        claims_text = ep_data.get('claims_text', '')

        # Generate table columns
        update_task_status(task_id, 'generating_columns', 30,
                           _t('china_generating_columns', lang))
        columns = await generate_table_columns(
            query=query, event_count=len(timeline_events),
            provider=flash_provider, lang=lang,
        )
        _pipeline_logger.info(
            f"[task={task_id}] EPO_EXAM PHASE2 columns — {columns}"
        )

        # Analyze each timeline event
        total_events = len(timeline_events)
        table_rows: list[dict] = []
        for i, evt in enumerate(timeline_events):
            update_task_status(
                task_id, 'analyzing', 35 + int(40 * i / max(total_events, 1)),
                _t('china_analyzing_decisions', lang, current=i + 1, total=total_events),
            )
            try:
                row = await analyze_single_timeline_event(
                    evt, search_report_text, claims_text,
                    columns, query, pro_provider, lang,
                )
                summary = await generate_event_summary(
                    evt, row, query, pro_provider, lang,
                )
                row['_summary'] = summary
                row['_event_type'] = evt.event_type
                row['_event_date'] = evt.date
                table_rows.append(row)
                _pipeline_logger.info(
                    f"[task={task_id}] EPO_EXAM PHASE2 analyzed — "
                    f"{i + 1}/{total_events} type={evt.event_type} "
                    f"summary={summary[:80]}"
                )
            except Exception as e:
                _pipeline_logger.warning(
                    f"[task={task_id}] EPO_EXAM PHASE2 analysis_error — "
                    f"event={i + 1}/{total_events}: {e}"
                )
                row = {col: "Analysis error" for col in columns}
                row['_summary'] = f"Analysis failed: {e}"
                row['_event_type'] = evt.event_type
                row['_event_date'] = evt.date
                table_rows.append(row)

        # ── Phase 3: Build report ─────────────────────────────────────
        update_task_status(task_id, 'generating_report', 80,
                           _t('writing_summary', lang))

        # Build markdown report
        biblio = ep_data.get('biblio')
        report_title = _t('epo_report_title', lang, patent_id=patent_id)

        lines = [f"# {report_title}\n"]

        # Executive summary via AI
        exec_heading = "核心审查洞察" if lang == "zh" else "Key Examination Insights"
        lines.append(f"## {exec_heading}\n")

        # Build a data summary for the executive summary AI call
        timeline_md = build_examination_timeline(timeline_events, biblio, lang)
        rows_text = "\n".join(
            f"- {r.get('_event_date', '?')} [{r.get('_event_type', '?')}]: "
            f"{r.get('_summary', '')[:200]}"
            for r in table_rows
        )

        if lang == "zh":
            exec_prompt = (
                "你是一个欧洲专利审查历史分析专家。基于以下审查数据，"
                "撰写一段 3-5 句话的核心审查洞察总结。\n"
                "包括：检索意见的核心发现、审查过程中主要的争议点、"
                "权利要求的变化方向、以及最终的审查结论。"
            )
        else:
            exec_prompt = (
                "You are an EPO patent examination analysis expert. "
                "Write a 3-5 sentence executive summary based on the "
                "examination data below. Cover: key findings from the "
                "search opinion, main issues during examination, direction "
                "of claim amendments, and final outcome."
            )

        try:
            exec_json = await pro_provider.complete_json(
                exec_prompt,
                f"Timeline:\n{timeline_md}\n\n"
                f"Analysis rows:\n{rows_text}\n\n"
                f"Search Report excerpt:\n{search_report_text[:3000]}\n\n"
                f'Return JSON: {{"summary": "..."}}'
            )
            exec_summary = exec_json.get("summary", "")
        except Exception:
            exec_summary = ""
        if exec_summary:
            lines.append(f"{exec_summary}\n")

        # Timeline
        lines.append(timeline_md)

        # Analysis table
        ana_heading = "审查事件分析表" if lang == "zh" else "Examination Event Analysis Table"
        lines.append(f"## {ana_heading}\n")

        if lang == "zh":
            header = "| " + " | ".join(columns) + " | 摘要 |"
            sep = "|" + "|".join("------" for _ in range(len(columns) + 1)) + "|"
        else:
            header = "| " + " | ".join(columns) + " | Summary |"
            sep = "|" + "|".join("------" for _ in range(len(columns) + 1)) + "|"
        lines.append(header)
        lines.append(sep)

        for row in table_rows:
            vals = [str(row.get(c, ""))[:200] for c in columns]
            vals.append(str(row.get("_summary", ""))[:300])
            lines.append("| " + " | ".join(vals) + " |")
        lines.append("")

        # Search opinion analysis (if available)
        if search_report_text:
            so_heading = "检索意见/书面意见" if lang == "zh" else "Search Opinion / Written Opinion"
            lines.append(f"## {so_heading}\n")
            lines.append(search_report_text[:5000])
            lines.append("")

        report_text = "\n".join(lines)

        _pipeline_logger.info(
            f"[task={task_id}] EPO_EXAM PHASE3 report — chars={len(report_text)}"
        )

        # ── Phase 4: Export ────────────────────────────────────────────
        from sources.result_export import export_docx_md_async as _export_docx
        from sources.result_export import export_pdf_async

        update_task_status(task_id, 'exporting', 90, _t('generating_word', lang))
        docx_bytes = await _export_docx(report_text, task_id=task_id)
        storage = create_storage()
        report_files: list[dict] = []

        try:
            await storage.put(task_id, 'report.docx', docx_bytes)
            _pipeline_logger.info(
                f"[task={task_id}] EPO_EXAM PHASE4 docx — size={len(docx_bytes)}"
            )
            report_files.append(
                {'format': 'docx', 'filename': 'report.docx', 'size': len(docx_bytes)}
            )
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] EPO_EXAM PHASE4 docx upload FAILED — {e}"
            )

        update_task_status(task_id, 'exporting', 95, _t('generating_pdf', lang))
        pdf_bytes = await export_pdf_async(docx_bytes)
        try:
            await storage.put(task_id, 'report.pdf', pdf_bytes)
            _pipeline_logger.info(
                f"[task={task_id}] EPO_EXAM PHASE4 pdf — size={len(pdf_bytes)}"
            )
            report_files.append(
                {'format': 'pdf', 'filename': 'report.pdf', 'size': len(pdf_bytes)}
            )
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] EPO_EXAM PHASE4 pdf upload FAILED — {e}"
            )

        # ── Complete ──
        _pipeline_logger.info(
            f"[task={task_id}] EPO_EXAM COMPLETED — "
            f"patent_id={patent_id}, ep_app={ep_app_number}, "
            f"events_analyzed={len(table_rows)}, "
            f"report_chars={len(report_text)}, "
            f"files={[f['format'] for f in report_files]}"
        )
        _update_mysql_progress(task_id, 'exporting', 100)
        set_task_completed(task_id, report_files)

        if user_id:
            from sources.long_task.user_queue import complete_user_task
            try:
                next_task_id = complete_user_task(str(user_id), task_id)
                if next_task_id:
                    _dispatch_queued_task(next_task_id, user_id)
            except Exception as e:
                _pipeline_logger.warning(
                    f"[task={task_id}] EPO_EXAM QUEUE_DISPATCH_FAILED — {e}"
                )

        return {'status': 'completed', 'task_id': task_id}

    loop = _asyncio.new_event_loop()
    try:
        return loop.run_until_complete(_run())
    except Exception as e:
        import traceback
        _pipeline_logger.error(
            f"[task={task_id}] EPO_EXAM UNHANDLED_ERROR — "
            f"{type(e).__name__}: {e}\n{traceback.format_exc()}"
        )
        if user_id:
            try:
                from sources.long_task.user_queue import complete_user_task
                complete_user_task(str(user_id), task_id)
            except Exception:
                pass
        set_task_failed(task_id, f"{type(e).__name__}: {e}")
        _update_mysql_progress(task_id, 'failed', 0)
        return {'status': 'failed', 'task_id': task_id,
                'error': f'{type(e).__name__}: {e}'}
    finally:
        loop.close()


# ── Japan Examination History Analysis Task (single patent)
# ═══════════════════════════════════════════════════════════════════════════════


@app.task(bind=True, max_retries=2, default_retry_delay=60, time_limit=1800, soft_time_limit=1770)
def execute_japan_examination_analysis(self, task_id: str, params: dict):
    """Analyze Japanese patent examination history via JPO IP Data Platform.

    Phase 0: Patent ID → JP application number via EPO family
    Phase 1: Fetch progress data from JPO API
    Phase 2: AI analysis of examination progress events
    Phase 3: Generate report (DOCX + PDF)
    """
    import asyncio as _asyncio
    import os as _os

    retry_count = self.request.retries
    patent_id = str(params.get('patent_id', '')).strip()
    query = params.get('query', '')
    lang = params.get('lang', 'zh')
    session_id = params.get('session_id', '')
    user_id = params.get('user_id', '')

    _pipeline_logger.info(
        f"[task={task_id}] JAPAN_EXAM START — "
        f"patent_id={patent_id}, query={query[:120]}, lang={lang}, "
        f"retry={retry_count}/{self.max_retries}"
    )

    if retry_count >= self.max_retries:
        _pipeline_logger.error(
            f"[task={task_id}] JAPAN_EXAM HARD_STOP — retry_count={retry_count} >= {self.max_retries}"
        )
        if user_id:
            from sources.long_task.user_queue import complete_user_task
            try:
                complete_user_task(str(user_id), task_id)
            except Exception:
                pass
        return {'status': 'failed', 'task_id': task_id,
                'error': f'Max retries ({self.max_retries}) exceeded'}

    if not patent_id:
        _pipeline_logger.error(f"[task={task_id}] JAPAN_EXAM no patent_id provided")
        _update_mysql_progress(task_id, 'failed', 0)
        return {'status': 'failed', 'task_id': task_id, 'error': 'No patent_id'}

    from sources.long_task.status_manager import (
        update_task_status, set_task_completed, set_task_failed,
    )
    from sources.long_task.config import (
        get_long_task_config, get_family_config, get_jpo_config,
    )
    from sources.long_task.storage import create_storage
    from sources.llm_provider import Provider

    # ── Provider setup ──
    ltc = get_long_task_config()
    model_family = ltc['provider_family']

    if model_family == 'minimax':
        flash_provider = Provider(
            provider_name='minimax', model='MiniMax-M2.7-highspeed',
            server_address='', is_local=False,
        )
        pro_provider = Provider(
            provider_name='minimax', model='MiniMax-M3',
            server_address='', is_local=False,
        )
    else:
        flash_provider = Provider(
            provider_name='deepseek', model='deepseek-v4-flash',
            server_address='', is_local=False,
        )
        pro_provider = Provider(
            provider_name='deepseek', model='deepseek-v4-pro',
            server_address='', is_local=False,
        )

    async def _run():
        # ── Phase 0: Resolve JP application number ──────────────────
        update_task_status(task_id, 'preparing', 5,
                           _t('preparing', lang, total=1),
                           analysis_type='prosecution')

        family_cfg = get_family_config()
        ep_ck = family_cfg.get('epo_consumer_key', '')
        ep_cs = family_cfg.get('epo_consumer_secret', '')
        ep_ck = _os.getenv('EPO_CONSUMER_KEY', ep_ck)
        ep_cs = _os.getenv('EPO_CONSUMER_SECRET', ep_cs)

        from sources.long_task.patent_family import EPOFamilyClient

        epo_client = None
        if ep_ck and ep_cs:
            epo_client = EPOFamilyClient(
                consumer_key=ep_ck, consumer_secret=ep_cs,
            )

        from sources.long_task.japan_examination import (
            resolve_jp_application_number,
            fetch_examination_data,
            build_examination_timeline,
            build_registration_summary,
            build_citations_summary,
            generate_table_columns,
            analyze_single_event,
            generate_event_summary,
        )

        jp_app_number, family_ctx = await resolve_jp_application_number(
            patent_id, epo_client,
        )
        if not jp_app_number:
            msg = _t('japan_no_review_data', lang)
            set_task_failed(task_id, msg)
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': msg}

        _pipeline_logger.info(
            f"[task={task_id}] JAPAN_EXAM PHASE0 resolved — "
            f"input={patent_id}, jp_app={jp_app_number}"
        )

        # ── Phase 1: Fetch examination data ──────────────────────────
        jpo_cfg = get_jpo_config()
        jpo_user = jpo_cfg.get('username', '')
        jpo_pass = jpo_cfg.get('password', '')
        if not jpo_user or not jpo_pass:
            msg = "JPO API credentials not configured (set JPO_USERNAME/JPO_PASSWORD env vars or [JPO] in config.ini)."
            set_task_failed(task_id, msg)
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': msg}

        from sources.jpo_client import JpoClient
        jpo = JpoClient(username=jpo_user, password=jpo_pass)

        update_task_status(task_id, 'preparing', 10,
                           f"Fetching JPO examination data for {jp_app_number}...")
        jp_data = await fetch_examination_data(jp_app_number, jpo)

        progress = jp_data.get('progress', [])
        if not progress:
            _pipeline_logger.warning(
                f"[task={task_id}] JAPAN_EXAM PHASE1 no_data — app={jp_app_number}"
            )
            set_task_failed(task_id, _t('japan_no_review_data', lang))
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id,
                    'error': _t('japan_no_review_data', lang)}

        _pipeline_logger.info(
            f"[task={task_id}] JAPAN_EXAM PHASE1 data_fetched — "
            f"events={jp_data.get('progress_count', 0)}, "
            f"has_reg={jp_data.get('has_registration')}, "
            f"has_cites={jp_data.get('has_citations')}"
        )

        # ── Phase 2: AI analysis ─────────────────────────────────────
        update_task_status(task_id, 'generating_columns', 30,
                           _t('china_generating_columns', lang))
        columns = await generate_table_columns(
            query=query, event_count=len(progress),
            provider=flash_provider, lang=lang,
        )
        _pipeline_logger.info(
            f"[task={task_id}] JAPAN_EXAM PHASE2 columns — {columns}"
        )

        total_events = len(progress)
        table_rows: list[dict] = []
        for i, evt in enumerate(progress):
            update_task_status(
                task_id, 'analyzing', 35 + int(40 * i / max(total_events, 1)),
                _t('china_analyzing_decisions', lang, current=i + 1, total=total_events),
            )
            try:
                row = await analyze_single_event(
                    evt, columns, query, pro_provider, lang,
                )
                summary = await generate_event_summary(
                    evt, row, query, pro_provider, lang,
                )
                row['_summary'] = summary
                row['_event_name'] = evt.get('event', '')
                row['_event_date'] = evt.get('event_date', '')
                table_rows.append(row)
                _pipeline_logger.info(
                    f"[task={task_id}] JAPAN_EXAM PHASE2 analyzed — "
                    f"{i + 1}/{total_events} event={evt.get('event', '?')[:40]}"
                )
            except Exception as e:
                _pipeline_logger.warning(
                    f"[task={task_id}] JAPAN_EXAM PHASE2 analysis_error — "
                    f"event={i + 1}/{total_events}: {e}"
                )
                row = {col: "Analysis error" for col in columns}
                row['_summary'] = f"Analysis failed: {e}"
                row['_event_name'] = evt.get('event', '')
                row['_event_date'] = evt.get('event_date', '')
                table_rows.append(row)

        # ── Phase 3: Build report ─────────────────────────────────────
        update_task_status(task_id, 'generating_report', 80,
                           _t('writing_summary', lang))

        report_title = _t('japan_report_title', lang, patent_id=patent_id)
        lines = [f"# {report_title}\n"]

        exec_heading = "核心审查洞察" if lang == "zh" else "Key Examination Insights"
        lines.append(f"## {exec_heading}\n")

        # Build context for executive summary
        timeline_md = build_examination_timeline(jp_data, lang)
        rows_text = "\n".join(
            f"- {r.get('_event_date', '?')} [{r.get('_event_name', '?')[:40]}]: "
            f"{r.get('_summary', '')[:200]}"
            for r in table_rows
        )

        if lang == "zh":
            exec_prompt = (
                "你是一个日本专利审查历史分析专家。基于以下审查数据，"
                "撰写一段 3-5 句话的核心审查洞察总结。\n"
                "包括：审查过程中的关键转折点、申请人的应对策略、"
                "权利要求的变化方向、以及最终的审查结论。"
            )
        else:
            exec_prompt = (
                "You are a Japan patent examination analysis expert. "
                "Write a 3-5 sentence executive summary based on the "
                "examination data below. Cover: key turning points in "
                "examination, applicant's strategy, direction of claim "
                "amendments, and final outcome."
            )

        try:
            exec_json = await pro_provider.complete_json(
                exec_prompt,
                f"Timeline:\n{timeline_md}\n\n"
                f"Analysis rows:\n{rows_text}\n\n"
                f'Return JSON: {{"summary": "..."}}'
            )
            exec_summary = exec_json.get("summary", "")
        except Exception:
            exec_summary = ""
        if exec_summary:
            lines.append(f"{exec_summary}\n")

        # Timeline
        lines.append(timeline_md)

        # Registration
        reg_md = build_registration_summary(jp_data, lang)
        if reg_md:
            lines.append(reg_md)

        # Citations
        cites_md = build_citations_summary(jp_data, lang)
        if cites_md:
            lines.append(cites_md)

        # Analysis table
        ana_heading = "审查事件分析表" if lang == "zh" else "Examination Event Analysis Table"
        lines.append(f"## {ana_heading}\n")

        if lang == "zh":
            header = "| " + " | ".join(columns) + " | 摘要 |"
            sep = "|" + "|".join("------" for _ in range(len(columns) + 1)) + "|"
        else:
            header = "| " + " | ".join(columns) + " | Summary |"
            sep = "|" + "|".join("------" for _ in range(len(columns) + 1)) + "|"
        lines.append(header)
        lines.append(sep)

        for row in table_rows:
            vals = [str(row.get(c, ""))[:200] for c in columns]
            vals.append(str(row.get("_summary", ""))[:300])
            lines.append("| " + " | ".join(vals) + " |")
        lines.append("")

        report_text = "\n".join(lines)

        _pipeline_logger.info(
            f"[task={task_id}] JAPAN_EXAM PHASE3 report — chars={len(report_text)}"
        )

        # ── Phase 4: Export ────────────────────────────────────────────
        from sources.result_export import export_docx_md_async as _export_docx
        from sources.result_export import export_pdf_async

        update_task_status(task_id, 'exporting', 90, _t('generating_word', lang))
        docx_bytes = await _export_docx(report_text, task_id=task_id)
        storage = create_storage()
        report_files: list[dict] = []

        try:
            await storage.put(task_id, 'report.docx', docx_bytes)
            _pipeline_logger.info(
                f"[task={task_id}] JAPAN_EXAM PHASE4 docx — size={len(docx_bytes)}"
            )
            report_files.append(
                {'format': 'docx', 'filename': 'report.docx', 'size': len(docx_bytes)}
            )
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] JAPAN_EXAM PHASE4 docx upload FAILED — {e}"
            )

        update_task_status(task_id, 'exporting', 95, _t('generating_pdf', lang))
        pdf_bytes = await export_pdf_async(docx_bytes)
        try:
            await storage.put(task_id, 'report.pdf', pdf_bytes)
            _pipeline_logger.info(
                f"[task={task_id}] JAPAN_EXAM PHASE4 pdf — size={len(pdf_bytes)}"
            )
            report_files.append(
                {'format': 'pdf', 'filename': 'report.pdf', 'size': len(pdf_bytes)}
            )
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] JAPAN_EXAM PHASE4 pdf upload FAILED — {e}"
            )

        # ── Complete ──
        _pipeline_logger.info(
            f"[task={task_id}] JAPAN_EXAM COMPLETED — "
            f"patent_id={patent_id}, jp_app={jp_app_number}, "
            f"events_analyzed={len(table_rows)}, "
            f"report_chars={len(report_text)}, "
            f"files={[f['format'] for f in report_files]}"
        )
        _update_mysql_progress(task_id, 'exporting', 100)
        set_task_completed(task_id, report_files)

        if user_id:
            from sources.long_task.user_queue import complete_user_task
            try:
                next_task_id = complete_user_task(str(user_id), task_id)
                if next_task_id:
                    _dispatch_queued_task(next_task_id, user_id)
            except Exception as e:
                _pipeline_logger.warning(
                    f"[task={task_id}] JAPAN_EXAM QUEUE_DISPATCH_FAILED — {e}"
                )

        return {'status': 'completed', 'task_id': task_id}

    loop = _asyncio.new_event_loop()
    try:
        return loop.run_until_complete(_run())
    except Exception as e:
        import traceback
        _pipeline_logger.error(
            f"[task={task_id}] JAPAN_EXAM UNHANDLED_ERROR — "
            f"{type(e).__name__}: {e}\n{traceback.format_exc()}"
        )
        if user_id:
            try:
                from sources.long_task.user_queue import complete_user_task
                complete_user_task(str(user_id), task_id)
            except Exception:
                pass
        set_task_failed(task_id, f"{type(e).__name__}: {e}")
        _update_mysql_progress(task_id, 'failed', 0)
        return {'status': 'failed', 'task_id': task_id,
                'error': f'{type(e).__name__}: {e}'}
    finally:
        loop.close()


# ── Prosecution History Analysis Task (single patent)
# ═══════════════════════════════════════════════════════════════════════════════


@app.task(bind=True, max_retries=2, default_retry_delay=30, time_limit=1800, soft_time_limit=1770)
def execute_prosecution_analysis(self, task_id: str, params: dict):
    """Analyze the prosecution history of a SINGLE USPTO patent application.

    Downloads Office Actions, Applicant Responses, Amendments, and Notice of
    Allowance from the USPTO Documents API, then uses AI to produce a
    comprehensive prosecution history report.

    Simpler than the batch pipeline: no Phase 0 search, no Phase 1 columns,
    no per-patent iteration.  Single patent, single report.
    """
    import asyncio
    import os as _os
    from typing import Any

    # ── Version marker: proves this is the new code (2026-07-28, streaming_provider) ──
    _pipeline_logger.info(
        f"[task={task_id}] BOOTSTRAP — celery_worker version=2026-07-28-streaming"
    )

    from sources.long_task.status_manager import (
        update_task_status, set_task_completed, set_task_failed,
    )
    from sources.long_task.config import (
        get_long_task_config, get_prosecution_config,
        DEFAULT_VISION_PROVIDER, DEFAULT_VISION_MODEL,
    )
    from sources.long_task.prosecution_downloader import (
        classify_prosecution_documents,
        download_single_document,
    )
    from sources.long_task.storage import create_storage
    from sources.llm_provider import Provider

    patent_id = str(params.get('patent_id', '')).strip()
    query = params.get('query', '')
    lang = params.get('lang', 'zh')
    session_id = params.get('session_id', '')
    user_id = params.get('user_id', '')

    _pipeline_logger.info(
        f"[task={task_id}] START — "
        f"patent_id={patent_id}, "
        f"query={query[:120]}, "
        f"lang={lang}, "
        f"session_id={session_id}"
    )

    if not patent_id:
        _pipeline_logger.error(f"[task={task_id}] PROSECUTION no patent_id provided")
        set_task_failed(task_id, "No patent application number provided")
        _update_mysql_progress(task_id, 'failed', 0)
        return {'status': 'failed', 'task_id': task_id, 'error': 'No patent_id'}

    # ── Provider setup ──
    ltc = get_long_task_config()
    model_family = ltc['provider_family']

    if model_family == 'minimax':
        flash_provider = Provider(
            provider_name='minimax', model='MiniMax-M2.7-highspeed',
            server_address='', is_local=False,
        )
        pro_provider = Provider(
            provider_name='minimax', model='MiniMax-M3',
            server_address='', is_local=False,
        )
    else:
        flash_provider = Provider(
            provider_name='deepseek', model='deepseek-v4-flash',
            server_address='', is_local=False,
        )
        pro_provider = Provider(
            provider_name='deepseek', model='deepseek-v4-pro',
            server_address='', is_local=False,
        )

    ptc = get_prosecution_config()
    include_priority_2 = ptc.get('include_priority_2', True)

    # ── Streaming provider (report summary + section writing) ──
    # When configured in [PROSECUTION], overrides pro_provider for the
    # streaming output phases so users can use their chat model (e.g.
    # openrouter/openai/gpt-5.4-mini) for the visible streaming text.
    streaming_provider = None
    _stream_cfg_provider = ptc.get('streaming_provider')
    _stream_cfg_model = ptc.get('streaming_model')
    _pipeline_logger.info(
        f"[task={task_id}] CONFIG — prosecution_config="
        f"streaming_provider={_stream_cfg_provider!r}, "
        f"streaming_model={_stream_cfg_model!r}, "
        f"ptc_keys={list(ptc.keys())}"
    )
    if _stream_cfg_provider and _stream_cfg_model:
        streaming_provider = Provider(
            provider_name=_stream_cfg_provider, model=_stream_cfg_model,
            server_address='', is_local=False,
        )
        _pipeline_logger.info(
            f"[task={task_id}] CONFIG — streaming_provider CREATED: "
            f"{_stream_cfg_provider}/{_stream_cfg_model}"
        )
    else:
        _pipeline_logger.warning(
            f"[task={task_id}] CONFIG — streaming_provider SKIPPED: "
            f"provider={_stream_cfg_provider!r}, model={_stream_cfg_model!r}"
        )

    vision_enabled = ltc.get('vision_enabled', True)
    vision_provider = None
    if vision_enabled:
        vision_cfg_provider = ltc.get('vision_provider', DEFAULT_VISION_PROVIDER)
        vision_cfg_model = ltc.get('vision_model', DEFAULT_VISION_MODEL)
        vision_provider = Provider(
            provider_name=vision_cfg_provider, model=vision_cfg_model,
            server_address='', is_local=False,
        )

    # ── Run pipeline ──
    async def _run():
        # ═════════════════════════════════════════════════════════════════
        # Phase 0: Resolve patent identifier → application number, then
        # fetch USPTO document list + classify by priority.
        # ═════════════════════════════════════════════════════════════════
        update_task_status(task_id, 'preparing', 1,
                           _t('fetching_uspto', lang),
                           analysis_type='prosecution')

        raw_patent_id = patent_id.strip()
        digits_only = ''.join(c for c in raw_patent_id if c.isdigit())
        id_type = params.get('patent_id_type', 'unknown')
        app_number = None

        # ── Resolve grant / publication number → application number ──
        # Also resolve when id_type is unknown or digits don't look like an
        # application number (e.g. 7-digit grant numbers like 8388852).
        _needs_resolution = (
            id_type in ('grant_number', 'publication_number')
            or id_type == 'unknown'
            or len(digits_only) < 8
        )
        if _needs_resolution:
            _pipeline_logger.info(
                f"[task={task_id}] PHASE0 resolving — "
                f"patent_id={patent_id}, id_type={id_type}"
            )
            try:
                import json as _json, re as _re
                _lucene_special = _re.compile(r'(["\\+\-!(){}[\]^~*?:/]|&&|\|\|)')
                def _escape_lucene(s: str) -> str:
                    return _lucene_special.sub(r'\\\1', s)

                _digits_esc = _escape_lucene(digits_only)
                _raw_esc = _escape_lucene(raw_patent_id)
                pub_candidates = [_raw_esc]
                if digits_only and digits_only != raw_patent_id:
                    pub_candidates.append(_digits_esc)
                pub_clauses = ' OR '.join(
                    f'applicationMetaData.earliestPublicationNumber:"{c}"'
                    for c in pub_candidates
                )
                # Search strategy by id_type:
                # - grant_number: search by patentNumber (most precise)
                # - publication_number: search by earliestPublicationNumber
                # - unknown: search ALL fields to maximize chance of finding it
                # Always include applicationNumberText as fallback.
                if id_type == 'grant_number':
                    search_q = f'applicationMetaData.patentNumber:"{_digits_esc}"'
                elif id_type == 'publication_number':
                    search_q = pub_clauses
                else:
                    # unknown or other — search both patentNumber and publicationNumber
                    search_q = (
                        f'applicationMetaData.patentNumber:"{_digits_esc}"'
                        f' OR {pub_clauses}'
                    )
                search_q += f' OR applicationNumberText:"{_digits_esc}"'

                search_body = {
                    "q": search_q,
                    "pagination": {"offset": 0, "limit": 3},
                    "fields": [
                        "applicationNumberText",
                        "applicationMetaData.patentNumber",
                        "applicationMetaData.earliestPublicationNumber",
                        "applicationMetaData.inventionTitle",
                    ],
                }
                uspto_key = _os.getenv('USPTO_API_KEY', '')
                search_headers = {'Content-Type': 'application/json', 'Accept': 'application/json'}
                if uspto_key:
                    search_headers['X-API-Key'] = uspto_key
                search_url = "https://api.uspto.gov/api/v1/patent/applications/search"

                # Use outbound_http for retry support on USPTO API calls.
                from sources.http_outbound import outbound_http as _outbound_http
                search_resp = await asyncio.to_thread(
                    _outbound_http.post, search_url,
                    purpose="patent_download",
                    headers=search_headers,
                    json=search_body,
                    timeout=30,
                )
                if search_resp.status_code == 200 and search_resp.text:
                    search_data = _json.loads(search_resp.text)
                    # USPTO returns different wrapper keys; scan for the first
                    # list value in the response dict (same pattern as
                    # _extract_raw_items in dynamic_tool_params.py).
                    raw_items = None
                    if isinstance(search_data, list):
                        raw_items = search_data
                    elif isinstance(search_data, dict):
                        for _v in search_data.values():
                            if isinstance(_v, list) and _v:
                                raw_items = _v
                                break
                    if raw_items:
                        _first = raw_items[0]
                        if isinstance(_first, dict):
                            app_number = (
                                _first.get('applicationNumberText') or
                                _first.get('applicationNumber') or
                                _first.get('appNumberText') or
                                _first.get('appNumber')
                            )
                            if app_number:
                                app_number = ''.join(c for c in str(app_number) if c.isdigit())
                if app_number:
                    _pipeline_logger.info(
                        f"[task={task_id}] PHASE0 resolved — "
                        f"input={patent_id} ({id_type}) → app_number={app_number}"
                    )
                else:
                    # Log response preview for debugging.
                    _resp_preview = (
                        search_resp.text[:500] if search_resp.text
                        else "(empty)"
                    )
                    _pipeline_logger.warning(
                        f"[task={task_id}] PHASE0 resolve_failed — "
                        f"status={search_resp.status_code}, "
                        f"response_preview={_resp_preview}, "
                        f"falling back to digits_only={digits_only}"
                    )
            except Exception as _resolve_err:
                _pipeline_logger.warning(
                    f"[task={task_id}] PHASE0 resolve_error — {_resolve_err}"
                )

        # ── Fallback: use digits_only as application number ──
        if not app_number:
            app_number = digits_only

        if not app_number or len(app_number) < 8:
            set_task_failed(task_id, f"Invalid patent application number: {patent_id}")
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id,
                    'error': f'Invalid patent_id: {patent_id}'}

        headers = {'Accept': 'application/json'}
        uspto_key = _os.getenv('USPTO_API_KEY', '')
        if uspto_key:
            headers['X-API-Key'] = uspto_key

        doc_list_url = (
            f"https://api.uspto.gov/api/v1/patent/applications/"
            f"{app_number}/documents"
        )
        _pipeline_logger.info(
            f"[task={task_id}] PHASE0 fetch_doc_list — "
            f"patent_id={patent_id}, url={doc_list_url}"
        )
        resp = await _uspto_get_with_retry(doc_list_url, headers, timeout=20)
        if resp.status_code != 200:
            _pipeline_logger.error(
                f"[task={task_id}] PHASE0 doc_list_failed — status={resp.status_code}"
            )
            set_task_failed(task_id, f"USPTO API returned HTTP {resp.status_code}")
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id,
                    'error': f'USPTO API status {resp.status_code}'}

        doc_list = resp.json() if resp.text else {}
        documents = (
            doc_list.get('documentBag', [])
            if isinstance(doc_list, dict)
            else []
        )
        if not documents:
            if lang == 'zh':
                msg = f"USPTO未返回专利申请 {app_number} 的任何文件。可能原因：申请号不存在、无权访问、或尚未公开。"
            else:
                msg = f"USPTO returned no documents for application {app_number}. The application may not exist, may not be accessible, or may not yet be published."
            set_task_failed(task_id, msg)
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': msg}

        # Classify documents by priority (deterministic, no LLM)
        manifest = classify_prosecution_documents(documents)
        docs_to_download = manifest.must_download.copy()
        if include_priority_2:
            docs_to_download.extend(manifest.recommended)

        _pipeline_logger.info(
            f"[task={task_id}] PHASE0 classified — "
            f"total_in_bag={len(documents)}, "
            f"must_download={len(manifest.must_download)}, "
            f"recommended={len(manifest.recommended)}, "
            f"skipped={len(manifest.skipped)}, "
            f"to_download={len(docs_to_download)}"
        )

        if not docs_to_download:
            if lang == 'zh':
                msg = "未找到可分析的审查文件（无 Office Action、Response 或 Amendment）。可能该专利尚未进入实质审查阶段。"
            else:
                msg = "No analyzable prosecution documents found (no Office Actions, Responses, or Amendments). The patent may not have entered substantive examination."
            set_task_failed(task_id, msg)
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': msg}

        # ═════════════════════════════════════════════════════════════════
        # Phase 1: Generate table columns (Flash LLM)
        # ═════════════════════════════════════════════════════════════════
        from sources.long_task.prosecution_analyzer import (
            generate_table_columns,
            analyze_single_document,
            generate_document_summary,
            build_failed_row,
            generate_prosecution_report as gen_report,
        )

        update_task_status(task_id, 'generating_columns', 5,
                           _t('prosecution_framework', lang, total=len(docs_to_download)))
        _pipeline_logger.info(
            f"[task={task_id}] PHASE1 generate_table_columns — "
            f"query={query[:100]}, doc_count={len(docs_to_download)}, "
            f"provider={flash_provider.model if hasattr(flash_provider, 'model') else 'flash'}"
        )
        columns = await generate_table_columns(
            query=query, doc_count=len(docs_to_download),
            provider=flash_provider, lang=lang,
        )
        _pipeline_logger.info(
            f"[task={task_id}] PHASE1 columns_generated — "
            f"column_count={len(columns)}, columns={columns}"
        )
        update_task_status(task_id, 'generating_columns', 5,
                           f'分析维度：{" | ".join(columns[1:4])}...',
                           table_columns=columns,
                           analysis_type='prosecution')
        _update_mysql_progress(task_id, 'generating_columns', 5)

        # ═════════════════════════════════════════════════════════════════
        # Phase 2: Sequential download + pipelined parallel analysis
        # ═════════════════════════════════════════════════════════════════
        # Each document is downloaded sequentially.  As soon as download
        # completes, analysis (OCR → LLM analyze → LLM summarize) is launched
        # immediately via asyncio.create_task().  All analysis steps run under
        # a Semaphore(5), so multiple OCR + LLM calls execute concurrently.
        total_dl = len(docs_to_download)
        _pipeline_logger.info(
            f"[task={task_id}] PHASE2 START — "
            f"pending_count={total_dl}, total={total_dl}"
        )
        update_task_status(task_id, 'downloading', 5,
                           _t('prosecution_downloading', lang, current=0, total=total_dl))

        async def _fetch_prosecution(url: str, hdrs: dict, timeout: int):
            return await _uspto_get_with_retry(url, hdrs, timeout)

        _table_rows: list[dict] = [None] * total_dl  # preserve order by index
        _analyze_sem = asyncio.Semaphore(100)  # caps concurrent OCR + LLM calls
        _pending_tasks = []  # list of asyncio.Task
        _downloaded = 0  # documents downloaded so far
        _analyzed = 0    # documents with completed (or skipped) analysis

        async def _analyze_one(_doc: Any, _i: int) -> None:
            nonlocal _analyzed, _downloaded
            async with _analyze_sem:
                doc_index = _i + 1

                # ── Text extraction (parallel, local-first → Vision fallback) ──
                if not _doc.text and _doc.binary:
                    # Step 1: try local qpdf+pdftotext first (fast, no API call)
                    _local_text = extract_text_from_binary(
                        _doc.binary, skip_pdf_extraction=False,
                    )
                    if _local_text and len(_local_text.strip()) > 50:
                        _doc.text = _local_text.strip()
                        _pipeline_logger.info(
                            f"[task={task_id}] PHASE2 local_extract_ok — "
                            f"code={_doc.document_code}, idx={doc_index}/{total_dl}, "
                            f"chars={len(_doc.text)}"
                        )
                    elif vision_enabled:
                        # Step 2: Vision LLM OCR (slower but handles scanned PDFs)
                        try:
                            _text = await _extract_text_via_vision(
                                _doc.binary, _doc.description, vision_provider,
                            )
                            if _text and len(_text.strip()) > 50:
                                _doc.text = _text.strip()
                                _pipeline_logger.info(
                                    f"[task={task_id}] PHASE2 vision_ok — "
                                    f"code={_doc.document_code}, idx={doc_index}/{total_dl}, "
                                    f"chars={len(_doc.text)}"
                                )
                        except Exception as _e:
                            _pipeline_logger.warning(
                                f"[task={task_id}] PHASE2 vision_error — "
                                f"code={_doc.document_code}: {type(_e).__name__}: {_e}"
                            )

                if not _doc.text or len(_doc.text.strip()) < 50:
                    row = build_failed_row(_doc.document_code, "text extraction failed", columns, lang)
                    row["_failed"] = True
                    row["_summary"] = ""
                    _table_rows[_i] = row
                    _analyzed += 1  # done with this doc (even though it failed)
                    return

                # ── Analyze (LLM) ──
                try:
                    row = await analyze_single_document(
                        doc_text=_doc.text,
                        doc_code=_doc.document_code,
                        doc_desc=_doc.description,
                        doc_category=_doc.category,
                        columns=columns,
                        query=query,
                        provider=pro_provider,
                        lang=lang,
                    )
                except Exception as e:
                    _pipeline_logger.warning(
                        f"[task={task_id}] PHASE2 analyze_error — "
                        f"code={_doc.document_code}, idx={doc_index}/{total_dl}: "
                        f"{type(e).__name__}: {e}"
                    )
                    row = build_failed_row(_doc.document_code, str(e), columns, lang)
                    row["_failed"] = True

                _pipeline_logger.info(
                    f"[task={task_id}] PHASE2 analyze_done — "
                    f"code={_doc.document_code}, idx={doc_index}/{total_dl}"
                )

                # ── Summarize (LLM) ──
                try:
                    summary = await generate_document_summary(
                        doc_text=_doc.text, row=row, query=query,
                        provider=pro_provider, lang=lang,
                    )
                except Exception as e:
                    _pipeline_logger.warning(
                        f"[task={task_id}] PHASE2 summary_error — "
                        f"code={_doc.document_code}: {type(e).__name__}: {e}"
                    )
                    summary = ""
                row["_summary"] = summary
                _table_rows[_i] = row
                _analyzed += 1

                # Always use max(downloaded, analyzed) so the bar never dips
                _p = max(_downloaded, _analyzed)
                _visible_rows = [r for r in _table_rows if r is not None]
                update_task_status(
                    task_id, 'analyzing',
                    progress_pct(_p, total_dl),
                    _t('prosecution_analyzing', lang,
                       current=_analyzed, total=total_dl,
                       desc=_doc.description[:40]),
                    table_rows=_visible_rows,
                    table_columns=columns,
                )

                _pipeline_logger.info(
                    f"[task={task_id}] PHASE2 analysis_complete — "
                    f"code={_doc.document_code}, idx={doc_index}/{total_dl}, "
                    f"summary_chars={len(summary)}, "
                    f"analyzed={_analyzed}/{total_dl}"
                )

        # ── Main pipeline: download sequentially, launch analysis immediately ──
        for _i, _doc in enumerate(docs_to_download):
            doc_index = _i + 1

            # ── Stop / Pause checkpoint ──
            _result = _handle_task_stop(task_id, user_id, _downloaded, total_dl)
            if _result:
                return _result
            _result = _handle_task_pause(task_id, user_id, _downloaded, total_dl, {
                'completed': [],
                'pending': [d.document_code for d in docs_to_download[_i:]],
                'completed_rows': [r for r in _table_rows if r is not None],
                'failed': [],
                'columns': columns,
                'downloaded': _downloaded,
                'analyzed': _analyzed,
            })
            if _result:
                return _result

            # ── Download (sequential) ──
            # Use max(downloaded, analyzed) so progress never drops when
            # an analysis completes while downloads are still running.
            _visible_progress = max(_downloaded, _analyzed)
            update_task_status(
                task_id, 'downloading',
                progress_pct(_visible_progress, total_dl),
                _t('prosecution_downloading', lang, current=doc_index, total=total_dl),
            )
            await download_single_document(_doc, _fetch_prosecution, app_number, headers)
            _downloaded += 1

            _pipeline_logger.info(
                f"[task={task_id}] PHASE2 download_done — "
                f"code={_doc.document_code}, idx={doc_index}/{total_dl}, "
                f"text_len={len(_doc.text) if _doc.text else 0}"
            )

            # ── Launch analysis immediately (OCR + LLM, parallel via semaphore) ──
            _pending_tasks.append(asyncio.create_task(_analyze_one(_doc, _i)))

        # ── Wait for all in-flight analyses to finish ──
        docs_with_text = [d for d in docs_to_download if d.text and len(d.text.strip()) >= 50]
        _pipeline_logger.info(
            f"[task={task_id}] PHASE2 all_downloads_done — "
            f"with_text={len(docs_with_text)}/{total_dl}, "
            f"waiting_for_{len(_pending_tasks)}_analyses"
        )
        if _pending_tasks:
            await asyncio.gather(*_pending_tasks)

        # Reconstruct table_rows in original order
        table_rows = [r for r in _table_rows if r is not None]

        _pipeline_logger.info(
            f"[task={task_id}] PHASE2 COMPLETE — "
            f"with_text={len(docs_with_text)}/{total_dl}, "
            f"table_rows={len(table_rows)}, "
            f"failed_rows={sum(1 for r in table_rows if r.get('_failed'))}"
        )

        if not table_rows:
            if lang == 'zh':
                msg = "所有审查文件处理失败。文件可能是扫描件或加密PDF。"
            else:
                msg = "All prosecution documents failed processing. Files may be scanned images or encrypted PDFs."
            set_task_failed(task_id, msg)
            _update_mysql_progress(task_id, 'failed', 0)
            return {'status': 'failed', 'task_id': task_id, 'error': msg}

        # ═════════════════════════════════════════════════════════════════
        # Phase 3: Generate report (Pro, dynamic outline + sections)
        # ═════════════════════════════════════════════════════════════════
        _pipeline_logger.info(
            f"[task={task_id}] PHASE3 generate_report — "
            f"columns={columns}, table_rows_count={len(table_rows)}, "
            f"provider={pro_provider.model if hasattr(pro_provider, 'model') else 'pro'}"
        )
        update_task_status(task_id, 'generating_report', 80,
                           _t('writing_summary', lang))

        from sources.long_task.status_manager import ThrottledSummaryUpdater
        summary_updater = ThrottledSummaryUpdater(
            task_id,
            progress=80,
            step_msg=_t('writing_summary', lang),
        )

        report_text = await gen_report(
            table_rows=table_rows,
            columns=columns,
            query=query,
            patent_id=patent_id,
            flash_provider=flash_provider,
            pro_provider=pro_provider,
            lang=lang,
            summary_updater=summary_updater,
            streaming_provider=streaming_provider,
        )
        _pipeline_logger.info(
            f"[task={task_id}] PHASE3 report_generated — "
            f"total_chars={len(report_text)}"
        )
        update_task_status(task_id, 'generating_report', 90,
                           _t('prosecution_report_complete', lang),
                           result_summary=report_text)
        _update_mysql_progress(task_id, 'generating_report', 90, result_summary=report_text)

        # ═════════════════════════════════════════════════════════════════
        # Phase 4: Export files (DOCX + PDF)
        # ═════════════════════════════════════════════════════════════════
        from sources.long_task.storage import get_storage_config

        storage_cfg = get_storage_config()
        _pipeline_logger.info(
            f"[task={task_id}] PHASE4 export — "
            f"storage_backend={storage_cfg.get('report_storage_backend', 'local')}"
        )
        try:
            storage = create_storage(storage_cfg)
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] PHASE4 storage_init FAILED — {e}, falling back to local"
            )
            storage = create_storage({'report_storage_backend': 'local'})

        report_files = []
        local_storage = None

        async def _get_local_storage():
            nonlocal local_storage
            if local_storage is None:
                from sources.long_task.storage import create_storage as _create
                local_storage = _create({'report_storage_backend': 'local'})
            return local_storage

        # ── DOCX (with table) ──
        update_task_status(task_id, 'exporting', 90, _t('generating_word', lang))
        docx_bytes = await export_docx_async(report_text, table_rows, columns, lang=lang)
        try:
            await storage.put(task_id, 'report.docx', docx_bytes)
            _pipeline_logger.info(
                f"[task={task_id}] PHASE4 docx — size_bytes={len(docx_bytes)}"
            )
            report_files.append({
                'format': 'docx', 'filename': 'report.docx', 'size': len(docx_bytes),
            })
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] PHASE4 docx upload FAILED — {e}, falling back to local"
            )
            try:
                await (await _get_local_storage()).put(task_id, 'report.docx', docx_bytes)
                _pipeline_logger.info(
                    f"[task={task_id}] PHASE4 docx — local fallback OK, size_bytes={len(docx_bytes)}"
                )
                report_files.append({
                    'format': 'docx', 'filename': 'report.docx', 'size': len(docx_bytes),
                })
            except Exception as e2:
                _pipeline_logger.error(
                    f"[task={task_id}] PHASE4 docx local fallback ALSO FAILED — {e2}"
                )

        # ── PDF ──
        update_task_status(task_id, 'exporting', 95, _t('generating_pdf', lang))
        pdf_bytes = await export_pdf_async(docx_bytes)
        try:
            await storage.put(task_id, 'report.pdf', pdf_bytes)
            _pipeline_logger.info(
                f"[task={task_id}] PHASE4 pdf — size_bytes={len(pdf_bytes)}"
            )
            report_files.append({
                'format': 'pdf', 'filename': 'report.pdf', 'size': len(pdf_bytes),
            })
        except Exception as e:
            _pipeline_logger.error(
                f"[task={task_id}] PHASE4 pdf upload FAILED — {e}, falling back to local"
            )
            try:
                await (await _get_local_storage()).put(task_id, 'report.pdf', pdf_bytes)
                _pipeline_logger.info(
                    f"[task={task_id}] PHASE4 pdf — local fallback OK, size_bytes={len(pdf_bytes)}"
                )
                report_files.append({
                    'format': 'pdf', 'filename': 'report.pdf', 'size': len(pdf_bytes),
                })
            except Exception as e2:
                _pipeline_logger.error(
                    f"[task={task_id}] PHASE4 pdf local fallback ALSO FAILED — {e2}"
                )

        # ── Complete ──
        _pipeline_logger.info(
            f"[task={task_id}] COMPLETED — "
            f"patent_id={patent_id}, "
            f"docs_analyzed={len(table_rows)}, "
            f"report_chars={len(report_text)}, "
            f"report_files={[f['format'] for f in report_files]}"
        )
        _update_mysql_progress(task_id, 'exporting', 100)
        set_task_completed(task_id, report_files)

        if user_id:
            from sources.long_task.user_queue import complete_user_task
            try:
                next_task_id = complete_user_task(str(user_id), task_id)
                if next_task_id:
                    _dispatch_queued_task(next_task_id, user_id)
            except Exception as e:
                import traceback
                _pipeline_logger.warning(
                    f"[task={task_id}] QUEUE_DISPATCH_FAILED — "
                    f"{type(e).__name__}: {e}"
                )

        return {'status': 'completed', 'task_id': task_id}

    loop = asyncio.new_event_loop()
    try:
        return loop.run_until_complete(_run())
    except Exception as e:
        import traceback
        _pipeline_logger.error(
            f"[task={task_id}] UNHANDLED_ERROR — "
            f"{type(e).__name__}: {e}\n{traceback.format_exc()}"
        )
        # Safety net: release user lock so queued tasks can proceed
        if user_id:
            try:
                from sources.long_task.user_queue import complete_user_task
                complete_user_task(str(user_id), task_id)
            except Exception:
                pass
        set_task_failed(task_id, f"{type(e).__name__}: {e}")
        _update_mysql_progress(task_id, 'failed', 0)
        return {'status': 'failed', 'task_id': task_id,
                'error': f'{type(e).__name__}: {e}'}
    finally:
        loop.close()


# ── Phase 0 helpers ─────────────────────────────────────────────────────────

def _get_max_patents_for_source(
    patent_source: str,
    default_max: int,
    max_cnipa: int = 10,
    max_uspto: int = 50,
) -> int:
    """Return the per-source patent count cap.

    CNIPA (China) → configurable (default 10).
    USPTO (US) → configurable (default 50).
    Unknown / other → conservative CNIPA limit.
    """
    if patent_source == 'uspto':
        return min(default_max, max_uspto)
    return min(default_max, max_cnipa)


def _infer_source_from_tool(tool_info, default: str = 'cnipa') -> str:
    """Infer patent source (cnipa/uspto) from a tool's URL."""
    url = getattr(tool_info, 'url', '') or ''
    if 'uspto' in url.lower():
        return 'uspto'
    if 'zldsj' in url.lower():
        return 'cnipa'
    return default


async def _download_patent_via_scene_or_fallback(
    patent_id: str,
    params: dict,
    scene_candidates: list | None,
    flash_provider,
    download_patent_document,
    id_url_map: dict | None = None,
    id_pid_map: dict | None = None,
) -> tuple[str | None, bytes | None]:
    """Download patent text via scene tool, direct USPTO API, or fallback.

    Returns (text, binary):
      - (text, None)          — text extracted successfully
      - (None, binary_bytes)  — binary cached for vision/OCR fallback
      - (None, None)          — download failed entirely
    """
    doc_url = (id_url_map or {}).get(patent_id, '')
    patent_source = params.get('patent_source', 'cnipa')
    pid = (id_pid_map or {}).get(patent_id, '')

    # ── Step 1: Try scene tool download ──
    if scene_candidates:
        from sources.long_task.scene_tools import select_tool, execute_tool

        context = f'专利来源: {patent_source}\npatent_id={patent_id}'
        if doc_url:
            context += f', document_url={doc_url}'
        if pid:
            context += f', pid={pid}'

        selected = await select_tool(
            'download patent document',
            context,
            scene_candidates,
            flash_provider,
        )
        if selected:
            tool_obj = selected.get('tool')
            _pipeline_logger.info(
                f"[download] scene_tool_selected — patent_id={patent_id}, "
                f"tool={getattr(tool_obj, 'title', '?') if tool_obj else '?'}"
            )
            result = await execute_tool(selected['tool'], selected['params'])
            from sources.long_task.scene_tools import extract_document_text
            text = extract_document_text(result)
            _pipeline_logger.info(
                f"[download] scene_tool_result — patent_id={patent_id}, "
                f"text_found={text is not None}, "
                f"text_length={len(text) if text else 0}"
            )
            # If the result looks like a document list (JSON metadata, not patent
            # text), go to direct USPTO download to fetch the actual specification.
            is_doc_list = (
                text and len(text) > 50
                and ('documentBag' in text or 'documentTypeCode' in text
                     or 'downloadOptionBag' in text
                     or 'document_code' in text.lower())
            )
            if text and len(text) > 50 and not is_doc_list:
                return (text, None)
            if is_doc_list:
                _pipeline_logger.info(
                    f"[download] scene_tool_returned_doc_list "
                    f"({len(text)} chars), "
                    f"trying direct USPTO API for specification"
                )
            else:
                _pipeline_logger.info(
                    f"[download] scene_tool_short_text "
                    f"({len(text) if text else 0} chars), "
                    f"trying direct USPTO API"
                )

    # ── Step 2: Direct USPTO API for US patents ──
    if patent_source == 'uspto':
        uspto_text, uspto_binary = await _download_uspto_patent_direct(patent_id, flash_provider)
        if uspto_text and len(uspto_text) > 100:
            return (uspto_text, None)
        if uspto_binary is not None:
            _pipeline_logger.info(
                f"[download] uspto_text_extraction_failed_but_binary_cached — "
                f"patent_id={patent_id}, binary_len={len(uspto_binary)}, "
                f"passing to vision/OCR fallback"
            )
            return (None, uspto_binary)
        _pipeline_logger.info(
            f"[download] uspto_text_extraction_failed — patent_id={patent_id}, "
            f"no binary cached, falling through to vision/OCR re-download"
        )
        return (None, None)

    # ── Step 3: Hardcoded download (CNIPA or other sources) ──
    _pipeline_logger.info(
        f"[download] fallback — patent_id={patent_id}, "
        f"patent_source={patent_source}"
    )
    text = await download_patent_document(patent_id, patent_source)
    return (text, None) if text else (None, None)


async def _download_uspto_patent_direct(patent_id: str, flash_provider=None) -> tuple[str | None, bytes | None]:
    """Download USPTO patent document text directly (two-step).

    Step 1: GET /api/v1/patent/applications/{appNumber}/documents → document list
    Step 2: LLM picks the specification → GET its download URL → return text

    Returns (text, binary):
      - (text, None)          — text extracted successfully
      - (None, binary_bytes)  — all specs failed text extraction, but binary cached
      - (None, None)          — download failed entirely
    """
    import asyncio
    import json as _json

    try:
        from sources.http_outbound import outbound_http

        # Normalize: strip commas, slashes, non-digits (US app numbers are pure digits)
        app_number = patent_id.strip().replace(',', '').replace('/', '')
        app_number = ''.join(c for c in app_number if c.isdigit())
        if not app_number or len(app_number) < 8:
            _pipeline_logger.warning(
                f"[download] uspto_invalid_app_number — patent_id={patent_id}"
            )
            return (None, None)
        headers = {'Accept': 'application/json'}
        uspto_key = os.getenv('USPTO_API_KEY', '')
        if uspto_key:
            headers['X-API-Key'] = uspto_key

        # Step 1: Get document list
        doc_list_url = (
            f"https://api.uspto.gov/api/v1/patent/applications/"
            f"{app_number}/documents"
        )
        _pipeline_logger.info(
            f"[download] uspto_step1 — url={doc_list_url}"
        )
        resp = await _uspto_get_with_retry(doc_list_url, headers, timeout=20)
        if resp.status_code != 200:
            _pipeline_logger.warning(
                f"[download] uspto_step1_failed — status={resp.status_code}"
            )
            return (None, None)

        doc_list = resp.json() if resp.text else {}
        documents = (
            doc_list.get('documentBag', [])
            if isinstance(doc_list, dict)
            else []
        )
        # Keep a compact one-line summary only
        if documents:
            spec_count = sum(
                1 for d in documents
                if isinstance(d, dict) and d.get('documentCode') == 'SPEC'
            )
            _pipeline_logger.info(
                f"[download] uspto_step1_done — doc_count={len(documents)}, "
                f"spec_count={spec_count}"
            )
        if not documents:
            _pipeline_logger.warning(
                f"[download] uspto_no_documents — patent_id={app_number}"
            )
            return (None, None)

        # Step 2: Collect ALL specification documents (there may be multiple)
        spec_docs: list[dict] = []

        # Find all SPEC documents heuristically first (fast, no LLM)
        for doc in documents:
            if not isinstance(doc, dict):
                continue
            code = str(doc.get('documentCode', '') or doc.get('documentTypeCode', ''))
            desc = str(doc.get('documentCodeDescriptionText', '') or doc.get('documentTypeName', ''))
            if 'SPEC' in code.upper() or 'specification' in desc.lower():
                spec_docs.append(doc)

        # If the LLM gave us a preferred index, move that one to the front
        if flash_provider and len(documents) > 1:
            doc_lines = []
            for i, doc in enumerate(documents):
                if not isinstance(doc, dict):
                    continue
                doc_lines.append(_json.dumps({
                    'index': i,
                    'code': doc.get('documentCode') or doc.get('documentTypeCode', ''),
                    'description': doc.get('documentCodeDescriptionText') or doc.get('documentTypeName', ''),
                    'pageCount': doc.get('pageTotalQuantity') or doc.get('pageCount', ''),
                    'hasDownload': bool(doc.get('downloadOptionBag')),
                }, ensure_ascii=False))
            try:
                selection = await flash_provider.complete_json(
                    "You are a patent document classifier. From a list of USPTO "
                    "patent application documents, identify the specification "
                    "(说明书). The specification is typically:\n"
                    "- code = 'SPEC' or description containing 'Specification'\n"
                    "- The main detailed description of the invention\n"
                    "- NOT the abstract, claims-only sequence listing, or drawings\n"
                    "- NOT administrative documents like Power of Attorney, "
                    "Fee Payment, Notice of Allowance\n\n"
                    "Return JSON: {\"selected_index\": <index of specification>, "
                    "\"reason\": \"<brief explanation>\"}",
                    f"Patent application: {app_number}\n"
                    f"Available documents:\n" + "\n".join(doc_lines),
                )
                if selection and isinstance(selection, dict):
                    idx = selection.get('selected_index')
                    if isinstance(idx, int) and 0 <= idx < len(documents):
                        preferred = documents[idx]
                        _pipeline_logger.info(
                            f"[download] llm_selected_spec — index={idx}, "
                            f"code={preferred.get('documentCode')}, "
                            f"reason={selection.get('reason', '')[:100]}"
                        )
                        # Move preferred to front (deduplicate if already in list)
                        spec_docs = [preferred] + [d for d in spec_docs if d is not preferred]
            except Exception as e:
                _pipeline_logger.warning(
                    f"[download] llm_spec_selection_failed: {e}"
                )

        if not spec_docs:
            _pipeline_logger.warning(
                f"[download] uspto_no_spec_found — patent_id={app_number}"
            )
            return (None, None)

        _pipeline_logger.info(
            f"[download] uspto_spec_candidates — count={len(spec_docs)}, "
            f"indices={[documents.index(d) for d in spec_docs]}"
        )

        # Step 3: Download ALL SPEC documents and concatenate their text.
        # A single patent application may have multiple SPEC files; downloading
        # all of them gives the most complete specification.
        all_parts: list[str] = []
        first_binary_fallback: bytes | None = None
        for attempt, spec_doc in enumerate(spec_docs):
            spec_code = spec_doc.get('documentCode') or spec_doc.get('documentTypeCode', '?')

            spec_url = get_download_url_from_doc(spec_doc)
            format_label = _guess_format_from_url(spec_url)
            _pipeline_logger.info(
                f"[download] uspto_spec[{attempt+1}/{len(spec_docs)}] — "
                f"code={spec_code}, format={format_label}, "
                f"url={spec_url[:100]}"
            )

            text, binary = await _download_uspto_spec_with_redirect(
                spec_doc, app_number, headers,
            )
            chars = len(text.strip()) if text else 0
            if chars > 200:
                _pipeline_logger.info(
                    f"[download] uspto_spec[{attempt+1}] ok — "
                    f"format={format_label}, chars={chars}"
                )
                all_parts.append(text.strip())
            else:
                _pipeline_logger.info(
                    f"[download] uspto_spec[{attempt+1}] skipped "
                    f"({chars} chars)"
                )
                if binary is not None and first_binary_fallback is None:
                    first_binary_fallback = binary
                    _pipeline_logger.info(
                        f"[download] uspto_spec[{attempt+1}] binary_cached "
                        f"for vision fallback — len={len(binary)}"
                    )

        if all_parts:
            combined = "\n\n".join(all_parts)
            _pipeline_logger.info(
                f"[download] uspto_all_specs_done — "
                f"parts={len(all_parts)}, total_chars={len(combined)}"
            )
            return (combined, None)

        if first_binary_fallback is not None:
            _pipeline_logger.info(
                f"[download] uspto_all_text_failed_but_binary_cached — "
                f"patent_id={app_number}, binary_len={len(first_binary_fallback)}"
            )
            return (None, first_binary_fallback)

        _pipeline_logger.warning(
            f"[download] uspto_all_specs_failed — patent_id={app_number}, "
            f"tried={len(spec_docs)}"
        )
        return (None, None)
    except Exception as e:
        _pipeline_logger.warning(
            f"[download] uspto_direct_error — patent_id={patent_id}, error={e}"
        )
        return (None, None)


# ── Vision extraction for scanned prosecution documents ───────────────────────


async def _extract_text_via_vision(
    pdf_bytes: bytes,
    doc_description: str,
    vision_provider,
    max_pages: int = 3,
) -> str | None:
    """Use vision LLM to extract text from a scanned/imaged PDF.

    Converts the first few pages to JPEG images and sends them to a vision
    model for OCR-like text extraction.  Only the first max_pages pages are
    processed to keep cost reasonable.

    Args:
        pdf_bytes: Raw PDF file bytes.
        doc_description: Human-readable document description for the prompt.
        vision_provider: Provider instance with OpenAI-compatible vision API.
        max_pages: Max pages to process (default 3).

    Returns:
        Extracted text, or None on failure.
    """
    try:
        import asyncio
        from sources.long_task.patent_analyzer import _pdf_to_base64_images

        # Run CPU-bound PDF→image conversion in thread to avoid blocking event loop
        loop = asyncio.get_running_loop()
        images = await loop.run_in_executor(
            None, _pdf_to_base64_images, pdf_bytes,
        )
        # Limit pages to control cost
        if len(images) > max_pages:
            images = images[:max_pages]
        if not images:
            _pipeline_logger.warning(
                f"[vision] no_images — desc={doc_description[:60]}"
            )
            return None

        _pipeline_logger.info(
            f"[vision] processing — desc={doc_description[:60]}, "
            f"pages={len(images)}"
        )

        client = vision_provider._get_raw_openai_client()
        model = vision_provider.model

        system_prompt = (
            "You are an OCR assistant. Extract ALL text from these patent "
            "prosecution document page images. Output the text verbatim — "
            "do not summarize, do not add commentary. Preserve section "
            "headings, numbered paragraphs, and claim language exactly as "
            "they appear."
        )

        user_content = [
            {
                "type": "text",
                "text": (
                    f"Extract all text from this USPTO document: "
                    f"\"{doc_description}\". "
                    f"Output the full text verbatim."
                ),
            },
        ]
        for img in images:
            user_content.append({
                "type": "image_url",
                "image_url": {"url": img, "detail": "high"},
            })

        import openai
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            max_tokens=4096,
            temperature=0.0,
        )
        text = response.choices[0].message.content or ""
        _pipeline_logger.info(
            f"[vision] done — desc={doc_description[:60]}, "
            f"chars={len(text)}"
        )
        return text.strip() if text.strip() else None

    except Exception as e:
        _pipeline_logger.warning(
            f"[vision] error — desc={doc_description[:60]}, "
            f"error={type(e).__name__}: {e}"
        )
        return None



def _guess_format_from_url(url: str) -> str:
    """Guess the file format from a download URL."""
    url_lower = url.lower()
    if url_lower.endswith('.docx') or 'ms_word' in url_lower:
        return 'DOCX'
    if url_lower.endswith('.xml') or 'xmlarchive' in url_lower:
        return 'XML'
    if url_lower.endswith('.pdf'):
        return 'PDF'
    return 'UNKNOWN'


def progress_pct(completed: int, total: int) -> int:
    """Map completed/total to progress percentage in [5, 75]."""
    if total == 0:
        return 5
    return 5 + min(70, int(completed / total * 70))


# ═══════════════════════════════════════════════════════════════════════════════
# Shared long-task control helpers (stop / pause)
# ═══════════════════════════════════════════════════════════════════════════════


def _handle_task_stop(
    task_id: str,
    user_id: str,
    completed: int,
    total: int,
) -> dict | None:
    """Check and handle task stop signal.  Shared by all long tasks.

    Returns {'status': 'cancelled', 'task_id': task_id} if stopped, None otherwise.
    """
    from sources.long_task.status_manager import is_task_stopped, set_task_cancelled

    if not is_task_stopped(task_id):
        return None

    _pipeline_logger.info(
        f"[task={task_id}] STOPPED_BY_USER — "
        f"completed={completed}/{total}"
    )
    set_task_cancelled(task_id)
    if user_id:
        try:
            from sources.long_task.user_queue import complete_user_task
            complete_user_task(str(user_id), task_id)
        except Exception:
            pass
    return {'status': 'cancelled', 'task_id': task_id}


def _handle_task_pause(
    task_id: str,
    user_id: str,
    completed: int,
    total: int,
    checkpoint: dict,
) -> dict | None:
    """Check and handle task pause signal.  Shared by all long tasks.

    Saves checkpoint, releases the per-user lock so the next queued task can
    start, and dispatches that next task.

    Returns {'status': 'paused', 'task_id': task_id} if paused, None otherwise.
    """
    from sources.long_task.status_manager import (
        is_task_paused, save_checkpoint, update_task_status,
    )

    if not is_task_paused(task_id):
        return None

    _pipeline_logger.info(
        f"[task={task_id}] PAUSED_BY_USER — "
        f"completed={completed}/{total}"
    )
    save_checkpoint(task_id, checkpoint)
    update_task_status(
        task_id, 'paused', progress_pct(completed, total),
        f'已暂停（{completed}/{total}），点击继续可恢复',
        status='paused',
    )
    if user_id:
        try:
            from sources.long_task.user_queue import complete_user_task
            next_id = complete_user_task(str(user_id), task_id)
            if next_id:
                _dispatch_queued_task(next_id, user_id)
        except Exception as e:
            _pipeline_logger.warning(
                f"[task={task_id}] PAUSE_DISPATCH_FAILED — {e}"
            )
    return {'status': 'paused', 'task_id': task_id}


def _dispatch_queued_task(next_task_id: str, user_id: str) -> None:
    """Read the next queued task from MySQL and dispatch it to Celery."""
    import json as _json
    from sources.knowledge.knowledge import get_db_connection as _gdc

    conn = _gdc()
    try:
        with conn.cursor() as cur:
            cur.execute(
                """SELECT input_params, session_id, scene_id, task_type
                   FROM long_tasks WHERE task_id = %s""",
                (next_task_id,),
            )
            row = cur.fetchone()
        if not row:
            return
        input_params = row.get('input_params')
        stored = _json.loads(input_params) if isinstance(input_params, str) else input_params
        next_params = {
            'query': stored.get('query', ''),
            'patent_ids': stored.get('patent_ids', []),
            'patent_source': stored.get('patent_source', 'auto'),
            'session_id': row.get('session_id') or '',
            'scene_id': row.get('scene_id'),
            'conversation_history': stored.get('conversation_history', []),
            'patent_file_refs': stored.get('patent_file_refs', []),
            'user_id': str(user_id),
        }
        if stored.get('patent_texts'):
            next_params['patent_texts'] = stored['patent_texts']

        task_type = row.get('task_type', 'patent_analysis')
        if task_type == 'family_analysis':
            next_params['patent_id'] = stored.get('patent_id', '')
            next_params['lang'] = stored.get('lang', 'zh')
            execute_family_analysis.delay(
                task_id=next_task_id, params=next_params,
            )
        elif task_type == 'prosecution_analysis':
            next_params['patent_id'] = stored.get('patent_id', '')
            next_params['lang'] = stored.get('lang', 'zh')
            execute_prosecution_analysis.delay(
                task_id=next_task_id, params=next_params,
            )
        else:
            execute_patent_analysis.delay(
                task_id=next_task_id, params=next_params,
            )
        _pipeline_logger.info(
            f"[queue] dispatched_after_pause — "
            f"next_task_id={next_task_id}, type={task_type}"
        )
    finally:
        conn.close()


async def export_pdf_async(docx_bytes: bytes) -> bytes:
    """Convert already-generated DOCX to PDF via LibreOffice headless."""
    import asyncio
    import tempfile
    import shutil

    def _convert():
        tmpdir = tempfile.mkdtemp(prefix='pdfconv_')
        try:
            docx_path = os.path.join(tmpdir, 'report.docx')
            with open(docx_path, 'wb') as f:
                f.write(docx_bytes)
            # LibreOffice headless conversion
            import subprocess
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf',
                 '--outdir', tmpdir, docx_path],
                capture_output=True, text=True, timeout=60,
            )
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
            pdf_path = os.path.join(tmpdir, 'report.pdf')
            with open(pdf_path, 'rb') as f:
                return f.read()
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)

    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(None, _convert)


async def export_docx_async(report_text: str, table_rows: list, columns: list, lang: str = 'zh') -> bytes:
    """Export report as DOCX using python-docx, run in executor."""
    import asyncio
    import io
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from lxml import etree
    import re as _re

    # Cross-platform font config for Word/WPS compatibility
    LATIN_FONT = 'Arial'
    CJK_FONT = 'Microsoft YaHei'  # 微软雅黑, pre-installed on Windows
    FALLBACK_FONT = 'SimSun'      # 宋体, universal fallback

    def _set_run_font(run, bold=False, italic=False):
        """Set cross-platform fonts on a run (Latin + East Asian)."""
        run.font.name = LATIN_FONT
        run.bold = bold
        run.italic = italic
        # Access XML to set East Asian font
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:ascii'), LATIN_FONT)
        rFonts.set(qn('w:hAnsi'), LATIN_FONT)
        rFonts.set(qn('w:eastAsia'), CJK_FONT)
        rFonts.set(qn('w:cs'), LATIN_FONT)

    def _set_paragraph_font(paragraph):
        """Set font on paragraph-level runs if they exist."""
        for run in paragraph.runs:
            _set_run_font(run)

    def _md_to_docx(doc, markdown_text: str):
        """Convert Markdown text to Word document with proper formatting."""
        lines = markdown_text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]

            # Heading: ### ...
            m = _re.match(r'^(#{1,4})\s+(.+)', line)
            if m:
                level = min(len(m.group(1)), 4)
                h = doc.add_heading(m.group(2).strip(), level=level)
                _set_heading_font(h)
                i += 1
                continue

            # Table: | col1 | col2 | ...
            if line.strip().startswith('|') and line.strip().endswith('|') and '---' not in line:
                # Collect table rows
                table_rows_md = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    row_line = lines[i].strip()
                    # A separator row (|:---:|:---|) becomes empty after removing
                    # pipe, hyphen, colon, and whitespace characters.
                    cleaned = row_line.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')
                    if cleaned:
                        cells = [c.strip() for c in row_line.strip('|').split('|')]
                        table_rows_md.append(cells)
                    i += 1
                if table_rows_md:
                    num_cols = max(len(r) for r in table_rows_md)
                    if num_cols > 0:
                        tbl = doc.add_table(rows=len(table_rows_md), cols=num_cols)
                        tbl.style = 'Table Grid'
                        for ri, row_cells in enumerate(table_rows_md):
                            for ci, cell_text in enumerate(row_cells):
                                if ci < num_cols:
                                    cell = tbl.rows[ri].cells[ci]
                                    _set_cell_text(cell, cell_text)
                        doc.add_paragraph('')  # spacing
                continue

            # Bullet list: - or * item
            m = _re.match(r'^\s*[-*]\s+(.+)', line)
            if m:
                p = doc.add_paragraph(style='List Bullet')
                _add_formatted_text(p, m.group(1).strip())
                _set_paragraph_font(p)
                i += 1
                continue

            # Numbered list: 1. item
            m = _re.match(r'^\s*\d+[.)]\s+(.+)', line)
            if m:
                p = doc.add_paragraph(style='List Number')
                _add_formatted_text(p, m.group(1).strip())
                _set_paragraph_font(p)
                i += 1
                continue

            # Separator
            if line.strip() in ('---', '***', '___'):
                doc.add_paragraph('─' * 60)
                i += 1
                continue

            # Regular paragraph
            if line.strip():
                p = doc.add_paragraph()
                _add_formatted_text(p, line.strip())
            else:
                doc.add_paragraph('')
            i += 1

    def _add_formatted_text(paragraph, text: str):
        """Add text with **bold** and *italic* formatting to a paragraph."""
        parts = _re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                _set_run_font(run, bold=True)
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                _set_run_font(run, italic=True)
            else:
                run = paragraph.add_run(part)
                _set_run_font(run)

    def _set_cell_text(cell, text: str):
        """Set cell text with cross-platform fonts."""
        cell.text = ''
        p = cell.paragraphs[0]
        _add_formatted_text(p, text)
        _set_paragraph_font(p)

    def _set_heading_font(heading):
        """Apply cross-platform fonts to a heading paragraph."""
        for run in heading.runs:
            _set_run_font(run, bold=True)

    def _sync_export():
        doc = Document()

        # Set document default fonts for cross-platform compatibility
        style = doc.styles['Normal']
        style.font.name = LATIN_FONT
        style.font.size = Pt(11)
        rPr = style.element.get_or_add_rPr()
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
        rFonts.set(qn('w:ascii'), LATIN_FONT)
        rFonts.set(qn('w:hAnsi'), LATIN_FONT)
        rFonts.set(qn('w:eastAsia'), CJK_FONT)
        rFonts.set(qn('w:cs'), LATIN_FONT)

        _md_to_docx(doc, report_text)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    loop = asyncio.get_running_loop()
    return await loop.run_in_executor(None, _sync_export)


def _build_report_html(report_text: str, table_rows: list, columns: list) -> str:
    """Build HTML for PDF export."""
    import html as html_mod
    rows_html = ""
    if table_rows and columns:
        header = "<tr>" + "".join(f"<th>{html_mod.escape(str(c))}</th>" for c in columns) + "</tr>"
        body = ""
        for r in table_rows:
            body += "<tr>" + "".join(f"<td>{html_mod.escape(str(r.get(c, '')))}</td>" for c in columns) + "</tr>"
        rows_html = f"<table>{header}{body}</table>"

    text_html = report_text.replace('\n', '<br>')
    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<style>
body{{
  font-family: 'Noto Sans CJK SC', 'Noto Sans SC', 'Noto Sans',
               'WenQuanYi Micro Hei', 'Microsoft YaHei',
               'PingFang SC', 'Hiragino Sans GB',
               'Arial', 'Helvetica Neue', sans-serif;
  max-width:900px;margin:0 auto;padding:20px;
  font-size:14px;line-height:1.8;color:#222;
}}
h1{{font-size:22px;font-weight:700;border-bottom:2px solid #10A37F;padding-bottom:8px;}}
h2{{font-size:18px;font-weight:700;margin-top:24px;}}
h3{{font-size:15px;font-weight:600;}}
table{{border-collapse:collapse;width:100%;margin:16px 0;}}
th,td{{border:1px solid #ddd;padding:8px;font-size:12px;text-align:left;}}
th{{background:#f5f5f5;font-weight:600;}}
tr:nth-child(even){{background:#fafafa;}}
</style></head>
<body>{text_html}{rows_html}</body></html>"""


# ── USPTO download helpers ────────────────────────────────────────────────────

import re as _re

# Text extraction functions moved to shared module
from sources.long_task.text_extractor import (
    USPTO_PREFERRED_MIME_ORDER,
    USPTO_PDF_PREFERRED_MIME_ORDER,
    get_download_url_from_doc,
    extract_text_from_pdf,
    extract_text_from_docx,
    extract_text_from_binary,
    _ocr_from_pdf_reader,
)

# Backward-compatible aliases used within this file
_USPTO_PREFERRED_MIME_ORDER = USPTO_PREFERRED_MIME_ORDER
_USPTO_PDF_PREFERRED_MIME_ORDER = USPTO_PDF_PREFERRED_MIME_ORDER
_get_download_url_from_doc = get_download_url_from_doc
_extract_text_from_pdf = extract_text_from_pdf
_extract_text_from_docx = extract_text_from_docx
_extract_text_from_binary = extract_text_from_binary


async def _uspto_get_with_retry(url: str, headers: dict, timeout: int = 30):
    """GET *url* via outbound_http.

    Retry on 400 / 429 is handled centrally by OutboundHttpClient for all
    USPTO API URLs.  This wrapper keeps the calling code unchanged.
    """
    import asyncio as _asyncio
    from sources.http_outbound import outbound_http

    return await _asyncio.to_thread(
        outbound_http.get, url, purpose="patent_download",
        headers=headers, timeout=timeout,
    )


async def _download_uspto_spec_with_redirect(
    spec_doc: dict,
    app_number: str,
    headers: dict,
) -> tuple[str | None, bytes | None]:
    """Download USPTO specification, following redirect URLs if needed.

    USPTO download URLs may return a text/JSON body containing another URL
    (e.g. "Please use redirect URL: https://...").  We follow at most one
    redirect.  Pattern taken from uspto_download.py.

    Returns (text, binary):
      - (text, None)          — text extracted successfully
      - (None, binary_bytes)  — binary downloaded but text extraction failed
      - (None, None)          — download failed entirely
    """
    import asyncio

    from sources.http_outbound import outbound_http

    spec_url = _get_download_url_from_doc(spec_doc)
    if not spec_url:
        _pipeline_logger.warning(
            f"[download] uspto_spec_no_url — app={app_number}, "
            f"spec_doc_keys={list(spec_doc.keys()) if spec_doc else 'N/A'}"
        )
        return (None, None)
    for hop in range(2):  # max 1 redirect
        resp = await _uspto_get_with_retry(spec_url, headers, timeout=30)
        if resp.status_code != 200:
            _pipeline_logger.warning(
                f"[download] uspto_spec_hop{hop}_failed — status={resp.status_code}"
            )
            return (None, None)

        content_type = resp.headers.get('Content-Type', '').lower()
        content = resp.text or ''

        # xmlarchive URLs always deliver tar binaries regardless of Content-Type
        # (USPTO may label them application/xml even though they are tar files).
        force_binary = 'xmlarchive' in spec_url.lower()

        # If response looks like a file (not text/JSON), extract text properly
        if force_binary or (content_type and not any(t in content_type for t in ('text/', 'json', 'xml', 'html'))):
            _pipeline_logger.info(
                f"[download] uspto_spec_binary — type={content_type}, "
                f"len={len(resp.content)}"
            )
            extracted = _extract_text_from_binary(
                resp.content, content_type, spec_url,
                skip_pdf_extraction=True,
            )
            if extracted and len(extracted) > 100:
                return (extracted, None)
            _pipeline_logger.warning(
                f"[download] uspto_spec_extract_empty — "
                f"type={content_type}, len={len(resp.content)}"
            )
            return (None, resp.content)

        # Check if the text response contains a redirect URL
        stripped = content.strip()
        if not stripped:
            _pipeline_logger.warning(
                f"[download] uspto_spec_empty — app={app_number}"
            )
            return (None, None)

        # Try to extract a redirect URL from the response
        from sources.dynamic_tool_params import _extract_first_url
        redirect_url = _extract_first_url(stripped)
        if redirect_url and redirect_url != spec_url:
            _pipeline_logger.info(
                f"[download] uspto_spec_redirect — to={redirect_url[:120]}"
            )
            spec_url = redirect_url
            continue

        # No redirect: this IS the content
        _pipeline_logger.info(
            f"[download] uspto_spec_done — len={len(stripped)}"
        )
        return (stripped, None)

    return (None, None)


# ── Binary download for vision fallback ──────────────────────────────────────

async def _download_uspto_binary_for_vision(
    patent_id: str,
    flash_provider,
) -> bytes | None:
    """Download USPTO specification as raw PDF bytes for vision LLM analysis.

    Mirrors the text-extraction path (_download_uspto_patent_direct) but
    returns the raw binary content instead of extracted text.
    """
    import asyncio

    from sources.dynamic_tool_params import _extract_first_url
    from sources.http_outbound import outbound_http

    uspto_headers = {
        "User-Agent": "Mozilla/5.0 (compatible; PatentAnalysis/1.0)",
        "Accept": "application/pdf, application/vnd.openxmlformats-officedocument.wordprocessingml.document, application/octet-stream, */*",
    }

    app_number = patent_id.strip()
    if not app_number.isdigit():
        app_number = "".join(c for c in app_number if c.isdigit())
        if not app_number:
            _pipeline_logger.warning(
                f"[vision_dl] uspto_invalid_app_number — patent_id={patent_id}"
            )
            return None

    # Step 1: Document list API
    docs_url = (
        f"https://api.uspto.gov/api/v1/patent/applications/{app_number}/documents"
    )
    try:
        resp = await _uspto_get_with_retry(docs_url, uspto_headers, timeout=30)
        if resp.status_code != 200:
            _pipeline_logger.warning(
                f"[vision_dl] uspto_doc_list_failed — status={resp.status_code}"
            )
            return None
        data = resp.json()
    except Exception as e:
        _pipeline_logger.warning(f"[vision_dl] uspto_doc_list_error — {e}")
        return None

    documents = data.get("documentBag", [])
    if not documents:
        return None

    # Step 2: Find specification (SPEC) documents
    spec_docs = [
        d for d in documents
        if d.get("documentCode") == "SPEC"
        or d.get("documentTypeCode") == "SPEC"
        or "SPEC" in str(d.get("applicationTypeCode", ""))
    ]
    if not spec_docs:
        # If LLM classification is available, use it
        if flash_provider and len(documents) > 1:
            try:
                doc_list = [
                    {
                        "idx": i,
                        "code": d.get("documentCode") or d.get("documentTypeCode", "?"),
                        "desc": (d.get("documentCodeDescriptionText") or d.get("documentTypeName", "")),
                        "has_docx": any(
                            "docx" in str(o.get("mimeTypeIdentifier", "")).lower()
                            or "ms_word" in str(o.get("mimeTypeIdentifier", "")).lower()
                            for o in d.get("downloadOptionBag", [])
                        ),
                    }
                    for i, d in enumerate(documents)
                ]
                prompt = (
                    "Which document index contains the patent specification "
                    "(detailed description)? Choose the one with code SPEC or "
                    "the most detailed description. Return JSON: "
                    '{"index": <number>, "code": "<code>", "reason": "<brief>"}'
                )
                selection = await flash_provider.complete_json(
                    "You select specification documents from patent document lists. "
                    "Return only valid JSON.",
                    f"Documents:\n{doc_list}\n\n{prompt}",
                )
                idx = int(selection.get("index", -1))
                if 0 <= idx < len(documents):
                    chosen = documents[idx]
                    spec_docs = [chosen]
                    _pipeline_logger.info(
                        f"[vision_dl] llm_selected_spec — index={idx}, "
                        f"code={chosen.get('documentCode', '?')}"
                    )
            except Exception:
                pass

    if not spec_docs:
        _pipeline_logger.warning(
            f"[vision_dl] uspto_no_spec — patent_id={app_number}"
        )
        return None

    # Step 3: Download binary from first successful spec URL.
    # Use PDF-first MIME order since vision LLM needs PDF page images.
    for attempt, spec_doc in enumerate(spec_docs):
        spec_url = get_download_url_from_doc(
            spec_doc, mime_order=_USPTO_PDF_PREFERRED_MIME_ORDER,
        )
        if not spec_url:
            continue

        _pipeline_logger.info(
            f"[vision_dl] attempt[{attempt+1}/{len(spec_docs)}] — "
            f"url={spec_url[:100]}"
        )

        for hop in range(2):
            try:
                resp = await _uspto_get_with_retry(spec_url, uspto_headers, timeout=30)
                if resp.status_code != 200:
                    break

                ct = resp.headers.get("Content-Type", "").lower()
                if ct and not any(t in ct for t in ("text/", "json", "xml", "html")):
                    _pipeline_logger.info(
                        f"[vision_dl] binary_downloaded — type={ct}, "
                        f"len={len(resp.content)}"
                    )
                    return resp.content

                # Check for redirect
                stripped = (resp.text or "").strip()
                redirect_url = _extract_first_url(stripped)
                if redirect_url and redirect_url != spec_url:
                    spec_url = redirect_url
                    continue

                break
            except Exception as e:
                _pipeline_logger.warning(f"[vision_dl] hop{hop}_error — {e}")
                break

    return None


# ── Conversation patent ID storage (for follow-up conversation_refs queries) ──────

_CONV_PATENT_IDS_TTL = 3600  # 1 hour


def _store_long_task_patent_ids(
    task_id: str,
    table_rows: list,
    columns: list,
    user_id: str,
) -> None:
    """Store patent IDs after a long task completes.

    Enables follow-up conversation_refs queries (e.g. "哪些是AI相关的")
    to find patent IDs even when the previous query was a long task rather
    than a general agent query with downloadable artifacts.
    """
    if not user_id or not table_rows:
        return

    # Use the first column (typically '专利号' / 'patent_number') as the ID source.
    id_column = columns[0] if columns else None
    if not id_column:
        return

    patent_ids = []
    for row in table_rows:
        pid = str(row.get(id_column, '')).strip()
        if pid and not row.get('_failed'):
            patent_ids.append(pid)

    if not patent_ids:
        return

    try:
        from sources.knowledge.knowledge import get_redis_connection
        r = get_redis_connection()
        key = f"lt:conv:{user_id}:patent_ids"
        r.set(key, json.dumps(patent_ids, ensure_ascii=False),
              ex=_CONV_PATENT_IDS_TTL)
        _pipeline_logger.info(
            f"[task={task_id}] stored_patent_ids_for_followup — "
            f"count={len(patent_ids)}, key={key}"
        )
    except Exception as e:
        _pipeline_logger.warning(
            f"[task={task_id}] store_patent_ids_failed — {e}"
        )
